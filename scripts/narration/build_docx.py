"""Phase 3: assemble Panduan Narasi - PKK Gr3.docx from pages_narration.json."""
from __future__ import annotations
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

HERE = Path(__file__).parent
NARRATION = HERE / "pages_narration.json"
OUT = Path(
    r"d:/DZAKI/S2/Sem. 1/Pelaporan Keuangan Korporat/"
    r"Pelaporan Keuangan Korporat Gr. 3/Panduan Narasi - PKK Gr3.docx"
)

ROLE_LABEL_ID = {
    "cover":   "Sampul",
    "agenda":  "Daftar Isi",
    "divider": "Pemisah Bagian",
    "content": "Konten",
    "case":    "Studi Kasus",
    "chart":   "Grafik",
    "table":   "Tabel",
    "qanda":   "Tanya Jawab",
    "ending":  "Penutup",
}

def add_cover(doc: Document, total: int) -> None:
    doc.add_heading("Panduan Narasi", level=0)
    p = doc.add_paragraph()
    p.add_run("Pelaporan Keuangan Korporat — Kelompok 3").bold = True
    doc.add_paragraph("Studi Kasus: FASB Conceptual Framework × Indofood (INDF) 2024")
    doc.add_paragraph(f"Total halaman deck: {total}")
    doc.add_paragraph("Tanggal: 2026-04-28")
    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run(
        "Dokumen ini berisi panduan narasi lisan untuk setiap halaman deck presentasi. "
        "Suara presenter yang digunakan adalah suara tim mahasiswa pascasarjana "
        '("kami"). Istilah teknis dijelaskan secara inline di dalam narasi, '
        "sedangkan istilah yang membutuhkan definisi lebih panjang dikumpulkan pada "
        'blok "Istilah Kunci" di bawah narasi setiap halaman.'
    )
    doc.paragraphs[-1].runs[0].add_break(WD_BREAK.PAGE)

def add_page_block(doc: Document, entry: dict) -> None:
    page, total = entry["page"], entry["total"]
    role = entry["role"]
    title = entry["title"] or "(tanpa judul)"

    doc.add_heading(f"Halaman {page}/{total} — {title}", level=1)
    role_p = doc.add_paragraph()
    role_p.add_run("Peran: ").bold = True
    role_p.add_run(f"{role} ({ROLE_LABEL_ID.get(role, role)})")

    nar_label = doc.add_paragraph()
    r = nar_label.add_run("\U0001F399️  Narasi")
    r.bold = True
    r.font.size = Pt(12)

    doc.add_paragraph(entry["narration"])

    if entry["key_terms"]:
        kt_label = doc.add_paragraph()
        rk = kt_label.add_run("\U0001F4D6  Istilah Kunci")
        rk.bold = True
        rk.font.size = Pt(12)
        for term in entry["key_terms"]:
            bullet = doc.add_paragraph(style="List Bullet")
            run_en = bullet.add_run(term["en"])
            run_en.bold = True
            bullet.add_run(f" ({term['id']}) — {term['def']}")

    last = doc.add_paragraph()
    last.add_run().add_break(WD_BREAK.PAGE)

def build():
    entries = json.loads(NARRATION.read_text(encoding="utf-8"))
    assert len(entries) == 69, f"expected 69 entries, got {len(entries)}"
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_cover(doc, total=entries[0]["total"])
    for entry in entries:
        add_page_block(doc, entry)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")

if __name__ == "__main__":
    build()
