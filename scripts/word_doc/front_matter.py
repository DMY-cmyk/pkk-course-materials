from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from scripts.word_doc.styles import (
    add_body_paragraph, add_page_break, set_run_font,
)


def _add_toc_field(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    for tag, text in [("begin", None), (None, ' TOC \\o "1-3" \\h \\z \\u '), ("end", None)]:
        if tag:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
            run._r.append(el)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = text
            run._r.append(instr)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(6)
    run_note = note.add_run(
        "[Catatan: Klik kanan pada area ini → Update Field untuk memperbarui nomor halaman]"
    )
    set_run_font(run_note, 9, italic=True, color_rgb=(100, 100, 100))


def build_front_matter(doc) -> None:
    """Adds Daftar Isi (TOC) page then Kata Pengantar page."""
    p_toc_title = doc.add_paragraph()
    p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_toc_title.paragraph_format.space_after = Pt(12)
    run = p_toc_title.add_run("DAFTAR ISI")
    set_run_font(run, 14, bold=True)

    _add_toc_field(doc)
    add_page_break(doc)

    p_kp_title = doc.add_paragraph()
    p_kp_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_kp_title.paragraph_format.space_after = Pt(12)
    run = p_kp_title.add_run("Kata Pengantar")
    set_run_font(run, 14, bold=True)

    add_body_paragraph(
        doc,
        "Puji syukur kami panjatkan kepada Tuhan Yang Maha Esa atas terselesaikannya "
        "makalah kelompok ini dalam rangka memenuhi tugas mata kuliah MNK202 Pelaporan "
        "Keuangan Korporat, Program Magister Akuntansi STIE YKPN Yogyakarta. Makalah ini "
        "menyajikan analisis komprehensif terhadap Kerangka Konseptual Financial Accounting "
        "Standards Board (FASB) — mulai dari SFAC 1 (1978) hingga SFAC 8 edisi September "
        "2024 — serta mengaplikasikan kerangka tersebut pada laporan tahunan PT Indofood "
        "Sukses Makmur Tbk (INDF) tahun 2024 sebagai studi kasus empiris. Kami berharap "
        "dokumen ini dapat memberikan kontribusi pemahaman yang mendalam mengenai fondasi "
        "konseptual yang mendasari standar pelaporan keuangan modern."
    )
    add_body_paragraph(
        doc,
        "Penyusunan makalah ini tidak terlepas dari bimbingan dosen pengampu dan kontribusi "
        "seluruh anggota kelompok. Kami menyadari bahwa dokumen ini masih jauh dari sempurna "
        "dan terbuka untuk masukan konstruktif demi peningkatan kualitas pemahaman bersama."
    )

    p_ttd = doc.add_paragraph()
    p_ttd.paragraph_format.space_before = Pt(20)
    p_ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_ttd = p_ttd.add_run("Yogyakarta, April 2026\n\nKelompok 3 — MNK202")
    set_run_font(run_ttd, 12)

    add_page_break(doc)
