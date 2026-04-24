from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Font helper ────────────────────────────────────────────────────────────────

def set_run_font(run, size_pt, bold=False, italic=False,
                 color_rgb=(0, 0, 0)):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)
    # Suppress theme-font override so TNR is actually used
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._r.get_or_add_rPr().append(rFonts)


# ── Paragraph borders ──────────────────────────────────────────────────────────

def _add_para_border(paragraph, sides, sz="6", color="000000"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "1")
        el.set(qn("w:color"), color)
        pBdr.append(el)
    pPr.append(pBdr)


# ── Typography hierarchy ───────────────────────────────────────────────────────

def add_doc_title(doc, text):
    """16pt Bold, centered — for cover page."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 16, bold=True)
    return p


def add_section_heading(doc, text):
    """14pt Bold + bottom border, for numbered sections I, II, …"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, 14, bold=True)
    _add_para_border(p, ["bottom"], sz="6", color="000000")
    return p


def add_sub_heading(doc, text):
    """12pt Bold, for sub-sections 1.1, 1.2, …"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, 12, bold=True)
    return p


def add_body_paragraph(doc, text):
    """12pt Regular, justify, 1.5 line spacing."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)   # 12pt × 1.5
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, 12)
    return p


def add_caption(doc, text):
    """10pt Italic, centered — for tables and figures."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, 10, italic=True)
    return p


def add_footnote_text(doc, text):
    """10pt Regular with top border — for section footnotes."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run(text)
    set_run_font(run, 10)
    _add_para_border(p, ["top"], sz="4", color="999999")
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    from docx.enum.text import WD_BREAK
    p.add_run().add_break(WD_BREAK.PAGE)
    return p


# ── Table helper ───────────────────────────────────────────────────────────────

def _set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_table_with_headers(doc, headers, rows_data):
    """
    Build a Word table with black header row (white bold text) and
    alternating gray rows. All fonts are 11pt TNR.

    headers   : list[str]
    rows_data : list[list[str]]
    Returns   : docx.table.Table
    """
    n_cols = len(headers)
    n_rows = len(rows_data)
    table = doc.add_table(rows=1 + n_rows, cols=n_cols)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, hdr_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(hdr_text)
        set_run_font(run, 11, bold=True, color_rgb=(255, 255, 255))
        _set_cell_shading(cell, "000000")

    # Data rows
    for row_idx, row_data in enumerate(rows_data):
        fill = "F5F5F5" if row_idx % 2 == 1 else "FFFFFF"
        cells = table.rows[row_idx + 1].cells
        for j, cell_text in enumerate(row_data):
            cell = cells[j]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(cell_text)
            set_run_font(run, 11)
            if fill != "FFFFFF":
                _set_cell_shading(cell, fill)

    return table


# ── Document setup ─────────────────────────────────────────────────────────────

def setup_document(doc):
    """Apply A4 page size, 2.5 cm margins."""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)


def add_page_numbers(doc):
    """Insert centered page number field in the footer."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for tag, text in [("begin", None), (None, "PAGE"), ("end", None)]:
        if tag:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
            run._r.append(el)
        else:
            instr = OxmlElement("w:instrText")
            instr.text = text
            run._r.append(instr)
