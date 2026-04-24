from docx.shared import Inches
from scripts.word_doc.styles import (
    add_section_heading, add_sub_heading, add_body_paragraph,
    add_caption, add_table_with_headers, add_footnote_text,
)
from scripts.word_doc.visuals import generate_bagan_fasb_to_psak


def build(doc) -> None:
    """Bagian VI — Sintesis Teori: SFAC 8 Keseluruhan"""

    add_section_heading(doc, "VI. SINTESIS TEORI: SFAC 8 SEBAGAI KERANGKA KONSEPTUAL KOMPREHENSIF")

    add_sub_heading(doc, "VI.1 Peta Delapan Chapters SFAC 8 (September 2024)")
    add_body_paragraph(
        doc,
        "SFAC 8 edisi September 2024 merupakan konsolidasi menyeluruh dari kerangka konseptual "
        "FASB dalam satu dokumen terpadu dengan delapan chapter. Setiap chapter menangani "
        "lapisan konseptual yang berbeda — dari tujuan paling mendasar (Chapter 1) hingga "
        "peran catatan keuangan (Chapter 8) — membentuk arsitektur hierarkis yang koheren."
    )
    add_caption(doc, "Tabel 5.1 — Delapan Chapters SFAC 8 (September 2024): Ringkasan")
    add_table_with_headers(
        doc,
        headers=["Ch.", "Judul", "Tanggal", "Isi Pokok"],
        rows_data=[
            ["1", "The Objective of General Purpose Financial Reporting",
             "2010", "Tujuan GPFR; primary users; stewardship"],
            ["2", "The Reporting Entity", "2018",
             "Definisi entitas pelaporan; batas konsolidasi"],
            ["3", "Qualitative Characteristics of Useful Financial Information",
             "2010", "QC Fundamental (Relevansi + FR) + Enhancing"],
            ["4", "Elements of Financial Statements", "Juli 2024",
             "Definisi 10 elemen; aset = 'present right'; BC4.7"],
            ["5", "Recognition and Derecognition", "Juli 2024",
             "3 kriteria pengakuan; RD3; Disclosure ≠ Recognition"],
            ["6", "Measurement", "Juli 2024",
             "Entry Price vs. Exit Price; M30–M34 decision guidance"],
            ["7", "Presentation and Disclosure", "Juli 2024",
             "PR12: Notes ≠ Recognition; BC7.21: basis OCI tidak jelas"],
            ["8", "Notes to Financial Statements", "Agustus 2023",
             "Fungsi catatan; 4 limitasi; D32 adverse consequences"],
        ],
    )

    add_sub_heading(doc, "VI.2 Jalur Pengaruh: FASB → IASB → PSAK → Praktik INDF")
    add_body_paragraph(
        doc,
        "Pemahaman SFAC 8 FASB secara langsung relevan bagi entitas Indonesia karena "
        "terdapat jalur pengaruh yang linear: (1) SFAC 8 FASB menjadi basis dialog konvergensi "
        "dalam Norwalk Agreement 2002; (2) konvergensi tersebut menghasilkan IASB Conceptual "
        "Framework for Financial Reporting 2018 yang selaras dengan SFAC 8; (3) DSAK-IAI "
        "mengadopsi IASB CF 2018 menjadi KKPK (Kerangka Konseptual Pelaporan Keuangan) IAI; "
        "(4) seluruh PSAK berbasis IFRS yang diterbitkan DSAK-IAI berlandaskan pada KKPK "
        "tersebut; dan (5) PT Indofood Sukses Makmur Tbk menyusun laporan keuangan "
        "konsolidasinya berdasarkan PSAK — merupakan manifestasi akhir dari rantai pengaruh ini."
    )
    add_caption(doc, "Bagan 5.1 — Jalur Pengaruh: FASB CF → IASB → PSAK → Praktik PT INDF")
    p_flow = doc.add_paragraph()
    p_flow.alignment = 1
    p_flow.add_run().add_picture(generate_bagan_fasb_to_psak(), width=Inches(5.8))

    add_footnote_text(
        doc,
        "⁶ KKPK = Kerangka Konseptual Pelaporan Keuangan, DSAK-IAI 2022 (adopsi IASB CF 2018). "
        "SFAC 8 September 2024 tersedia di fasb.org. Sumber: Wolk et al. (2017) Ch.7."
    )
