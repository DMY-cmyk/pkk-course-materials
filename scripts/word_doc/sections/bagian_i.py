from docx.shared import Inches
from scripts.word_doc.styles import (
    add_section_heading, add_sub_heading, add_body_paragraph,
    add_caption, add_table_with_headers, add_footnote_text,
)
from scripts.word_doc.visuals import generate_bagan_timeline, generate_bagan_primary_users


def build(doc) -> None:
    """Bagian I — Konteks dan Evolusi Kerangka Konseptual FASB"""

    add_section_heading(doc, "I. KONTEKS DAN EVOLUSI KERANGKA KONSEPTUAL FASB")

    add_sub_heading(doc, "I.1 Sejarah Perkembangan: SFAC 1 (1978) hingga SFAC 8 (2024)")
    add_body_paragraph(
        doc,
        "Kerangka Konseptual (Conceptual Framework/CF) FASB merupakan fondasi hierarkis "
        "yang mendasari seluruh standar pelaporan keuangan yang diterbitkan oleh Financial "
        "Accounting Standards Board (FASB). Berbeda dengan standar individual yang mengatur "
        "transaksi spesifik, kerangka konseptual menetapkan prinsip-prinsip dasar yang menjadi "
        "rujukan ketika standar yang ada tidak mencakup suatu transaksi tertentu — sebuah fungsi "
        "yang oleh Wolk, Dodd & Rozycki (2017) disebut sebagai 'panduan pengambilan keputusan "
        "ketika aturan spesifik tidak tersedia.'"
    )
    add_body_paragraph(
        doc,
        "Perjalanan CF FASB dimulai pada tahun 1978 dengan diterbitkannya SFAC 1 (Statement of "
        "Financial Accounting Concepts No. 1), yang untuk pertama kalinya menetapkan tujuan "
        "pelaporan keuangan secara eksplisit. Selama lebih dari empat dekade, FASB telah "
        "menerbitkan delapan SFAC, dengan edisi komprehensif terakhir SFAC 8 diterbitkan pada "
        "September 2024 — mengintegrasikan seluruh pembaruan konseptual termasuk definisi aset "
        "baru (Chapter 4), tiga kriteria pengakuan (Chapter 5), dua sistem pengukuran "
        "(Chapter 6), dan panduan penyajian (Chapter 7)."
    )

    add_caption(doc, "Bagan 1.1 — Evolusi Kerangka Konseptual FASB (1978–2024)")
    p_img = doc.add_paragraph()
    p_img.alignment = 1
    p_img.add_run().add_picture(generate_bagan_timeline(), width=Inches(5.8))

    add_caption(doc, "Tabel 1.1 — Ringkasan Delapan Statement of Financial Accounting Concepts (SFAC)")
    add_table_with_headers(
        doc,
        headers=["No.", "Tahun", "Judul / Topik", "Status"],
        rows_data=[
            ["SFAC 1", "1978", "Tujuan Pelaporan Keuangan Bisnis",
             "Digantikan SFAC 8 (2010)"],
            ["SFAC 2", "1980", "Karakteristik Kualitatif Informasi Akuntansi",
             "Digantikan SFAC 8 (2010)"],
            ["SFAC 3", "1980", "Elemen Laporan Keuangan Bisnis",
             "Digantikan SFAC 6 (1985)"],
            ["SFAC 4", "1980", "Tujuan LK Entitas Non-Bisnis",
             "Aktif"],
            ["SFAC 5", "1984", "Pengakuan dan Pengukuran dalam LK Bisnis",
             "Aktif — dikritisi akademisi"],
            ["SFAC 6", "1985", "Elemen Laporan Keuangan (Revisi)",
             "Aktif"],
            ["SFAC 7", "2000", "Penggunaan Cash Flow dan Nilai Sekarang",
             "Aktif"],
            ["SFAC 8", "2010–2024", "Kerangka Konseptual (8 Chapters, Sept 2024)",
             "Aktif — edisi terkini"],
        ],
    )

    add_sub_heading(doc, "I.2 SFAC 1: Tujuan Pelaporan Keuangan dan Primary Users")
    add_body_paragraph(
        doc,
        "SFAC 1 (1978) menetapkan bahwa tujuan utama pelaporan keuangan adalah menyediakan "
        "informasi yang berguna bagi pengguna utama (primary users): investor saat ini dan "
        "potensial, serta kreditor. Informasi tersebut harus membantu mereka dalam mengevaluasi "
        "jumlah, waktu, dan ketidakpastian arus kas masa depan entitas — sebuah perspektif yang "
        "secara eksplisit menempatkan investor dan kreditur sebagai fokus utama, bukan manajemen "
        "atau regulator."
    )
    add_body_paragraph(
        doc,
        "SFAC 8 (2010) mempertegas konsep ini dengan mendefinisikan primary users sebagai "
        "mereka yang tidak dapat meminta laporan khusus dari entitas (cannot require entities "
        "to provide information directly). Definisi ini penting karena menetapkan batasan "
        "yang jelas mengenai siapa yang menjadi audiens utama laporan keuangan bertujuan "
        "umum (general purpose financial statements/GPFS)."
    )

    add_caption(doc, "Bagan 1.2 — Primary Users dan Zona Pengguna Laporan Keuangan (SFAC 8)")
    p_pu = doc.add_paragraph()
    p_pu.alignment = 1
    p_pu.add_run().add_picture(generate_bagan_primary_users(), width=Inches(5.0))

    add_caption(doc, "Tabel 1.2 — Perbandingan Tujuan Pelaporan Keuangan: SFAC 1 (1978) vs. SFAC 8 Ch.1 (2010)")
    add_table_with_headers(
        doc,
        headers=["Dimensi", "SFAC 1 (1978)", "SFAC 8 Ch.1 (2010)"],
        rows_data=[
            ["Tujuan Utama",
             "Informasi berguna untuk keputusan investasi dan kredit",
             "Informasi berguna bagi primary users untuk keputusan penyediaan sumber daya"],
            ["Definisi Primary Users",
             "Investor, kreditor, dan pengguna potensial",
             "Existing/potential investors, lenders, other creditors yang tidak dapat meminta info langsung"],
            ["Fokus Informasi",
             "Arus kas masa depan entitas",
             "Arus kas masa depan + posisi keuangan + kinerja (earning power)"],
            ["Perspektif",
             "Stewardship sebagai tujuan sekunder",
             "Stewardship (akuntabilitas manajemen) sebagai tujuan setara dengan keputusan"],
        ],
    )

    add_sub_heading(doc, "I.3 Norwalk Agreement (2002) dan Konvergensi FASB–IASB")
    add_body_paragraph(
        doc,
        "Tonggak penting dalam sejarah konvergensi standar akuntansi global adalah "
        "Norwalk Agreement pada 18 September 2002. Dalam perjanjian ini, FASB dan "
        "International Accounting Standards Board (IASB) secara resmi berkomitmen untuk "
        "mengembangkan standar akuntansi yang kompatibel secara internasional dan "
        "mengoordinasikan agenda kerja mereka. Salah satu produk langsung dari komitmen "
        "ini adalah proyek bersama pengembangan kerangka konseptual, yang menghasilkan "
        "SFAC 8 pada tahun 2010 — sebuah dokumen yang mencerminkan konsensus FASB dan IASB "
        "setelah hampir satu dekade deliberasi bersama."
    )

    add_sub_heading(doc, "I.4 Relevansi bagi Indonesia: IASB CF 2018 → PSAK")
    add_body_paragraph(
        doc,
        "Di Indonesia, Dewan Standar Akuntansi Keuangan — Ikatan Akuntan Indonesia (DSAK-IAI) "
        "mengadopsi Kerangka Konseptual Pelaporan Keuangan (KKPK) berdasarkan IASB Conceptual "
        "Framework for Financial Reporting 2018. Dengan demikian, terdapat jalur pengaruh yang "
        "jelas: FASB CF (SFAC 8) → Konvergensi FASB–IASB (2002–2010) → IASB CF 2018 → "
        "PSAK/KKPK IAI → Praktik pelaporan keuangan entitas Indonesia. PT Indofood Sukses "
        "Makmur Tbk, sebagai entitas yang menyusun laporan keuangan konsolidasi berdasarkan "
        "PSAK, merupakan manifestasi langsung dari rantai pengaruh konseptual ini."
    )

    add_footnote_text(
        doc,
        "¹ SFAC = Statement of Financial Accounting Concepts. PSAK = Pernyataan Standar "
        "Akuntansi Keuangan. KKPK = Kerangka Konseptual Pelaporan Keuangan (DSAK-IAI). "
        "Sumber utama: Wolk, Dodd & Rozycki (2017), Ch.7; SFAC 8 (September 2024)."
    )
