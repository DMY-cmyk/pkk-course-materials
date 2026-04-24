from docx.shared import Inches
from scripts.word_doc.styles import (
    add_section_heading, add_sub_heading, add_body_paragraph,
    add_caption, add_table_with_headers, add_footnote_text,
)
from scripts.word_doc.visuals import generate_bagan_ci_waterfall


def build(doc) -> None:
    """Bagian III — Elemen-Elemen Laporan Keuangan"""

    add_section_heading(doc, "III. ELEMEN-ELEMEN LAPORAN KEUANGAN")

    add_sub_heading(doc, "III.1 SFAC 6 (1985): Sepuluh Elemen Laporan Keuangan")
    add_body_paragraph(
        doc,
        "SFAC 6 (1985) mendefinisikan sepuluh elemen laporan keuangan yang membentuk "
        "kerangka dasar penyusunan neraca dan laporan laba rugi. Elemen-elemen ini bukan "
        "sekadar kategori akuntansi — masing-masing memiliki definisi konseptual yang presisi "
        "yang membedakannya dari elemen lain dan menentukan kondisi pengakuan dalam laporan "
        "keuangan."
    )
    add_caption(doc, "Tabel 3.1 — Sepuluh Elemen Laporan Keuangan (SFAC 6, 1985)")
    add_table_with_headers(
        doc,
        headers=["Elemen", "Definisi Konseptual (ringkas)", "Posisi dalam LK"],
        rows_data=[
            ["Assets", "Manfaat ekonomi masa depan yang mungkin diperoleh atau dikendalikan entitas",
             "Neraca (Aset)"],
            ["Liabilities", "Pengorbanan manfaat ekonomi masa depan yang probable dari kewajiban sekarang",
             "Neraca (Liabilitas)"],
            ["Equity", "Kepentingan residual aset setelah dikurangi liabilitas",
             "Neraca (Ekuitas)"],
            ["Investments by Owners", "Kenaikan ekuitas dari transfer aset oleh pemilik",
             "Laporan Ekuitas"],
            ["Distributions to Owners", "Penurunan ekuitas dari transfer ke pemilik",
             "Laporan Ekuitas"],
            ["Comprehensive Income (CI)", "Perubahan ekuitas dari transaksi non-pemilik",
             "Laporan Laba Rugi Komprehensif"],
            ["Revenues", "Arus masuk/peningkatan aset dari aktivitas utama entitas",
             "Laporan Laba Rugi"],
            ["Expenses", "Arus keluar/penggunaan aset dari aktivitas utama entitas",
             "Laporan Laba Rugi"],
            ["Gains", "Kenaikan ekuitas dari transaksi peripheral (bukan aktivitas utama)",
             "Laporan Laba Rugi"],
            ["Losses", "Penurunan ekuitas dari transaksi peripheral",
             "Laporan Laba Rugi"],
        ],
    )

    add_sub_heading(doc, "III.2 SFAC 8 Chapter 4 (Juli 2024): Redefinisi Fundamental Aset")
    add_body_paragraph(
        doc,
        "SFAC 8 Chapter 4 (Juli 2024) melakukan perubahan yang secara konseptual lebih radikal "
        "dari sekadar memperbarui bahasa: definisi aset diganti dari 'probable future economic "
        "benefits' (SFAC 6) menjadi 'a present right of the entity to an economic benefit.' "
        "Perubahan ini memiliki implikasi praktis yang substansial: kata 'probable' dan "
        "konsep 'kontrol' (control) dihapus dari definisi formal."
    )
    add_body_paragraph(
        doc,
        "Dalam Basis for Conclusions paragraf BC4.7, FASB menjelaskan bahwa definisi baru "
        "memberikan primasi konseptual pada hak hukum (present right) sebagai dasar pengakuan "
        "aset, bukan pada kemungkinan realisasi manfaat di masa depan. Implikasinya: suatu hak "
        "yang ada saat ini sudah cukup untuk memenuhi definisi aset, meskipun probabilitas "
        "realisasi manfaat ekonomi masih rendah. Penentuan apakah aset tersebut diakui dalam "
        "laporan keuangan merupakan fungsi dari kriteria pengakuan (Recognition — Chapter 5), "
        "bukan definisi elemen."
    )
    add_caption(doc, "Tabel 3.2 — Perbandingan Definisi Aset: SFAC 6 (1985) vs. SFAC 8 Ch.4 (Juli 2024)")
    add_table_with_headers(
        doc,
        headers=["Aspek", "SFAC 6 (1985)", "SFAC 8 Ch.4 (Juli 2024)"],
        rows_data=[
            ["Definisi Aset",
             "Probable future economic benefits obtained or controlled by an entity as a result of past transactions",
             "A present right of the entity to an economic benefit"],
            ["Kata Kunci", "'Probable' + 'controlled' + 'future benefits'",
             "'Present right' saja"],
            ["Dasar Pengakuan", "Kemungkinan (probability) manfaat masa depan",
             "Keberadaan hak hukum sekarang (present right)"],
            ["Konsep Kontrol", "Eksplisit sebagai bagian definisi",
             "Dihapus dari definisi; diatur di Chapter 5 (Recognition)"],
            ["Relevansi BC", "—",
             "BC4.7: 'present right' memberikan primasi konseptual pada hak hukum"],
        ],
    )

    add_sub_heading(doc, "III.3 Comprehensive Income (CI) dan OCI: Hubungan Konseptual")
    add_body_paragraph(
        doc,
        "Comprehensive Income (CI) merupakan elemen yang paling luas mencakup seluruh perubahan "
        "ekuitas dari transaksi non-pemilik. CI terdiri dari: (1) Net Income — hasil aktivitas "
        "utama dan peripheral yang dilaporkan dalam Laporan Laba Rugi, dan (2) Other Comprehensive "
        "Income (OCI) — perubahan nilai yang belum direalisasi (unrealized gains/losses) yang "
        "dilaporkan terpisah. Contoh OCI: perubahan nilai investasi tersedia untuk dijual, "
        "translasi mata uang asing, dan re-measurement imbalan kerja."
    )
    add_body_paragraph(
        doc,
        "SFAC 8 Chapter 7, paragraf BC7.21, secara eksplisit menyatakan bahwa 'tidak terdapat "
        "basis konseptual yang jelas untuk pengklasifikasian item ke dalam OCI vs. Net Income.' "
        "Ini merupakan pengakuan FASB bahwa pemisahan OCI lebih bersifat pragmatis-politis "
        "ketimbang konseptual — entitas dan standar setter menggunakannya untuk menghindari "
        "volatilitas Laporan Laba Rugi dari perubahan nilai wajar jangka pendek."
    )

    add_sub_heading(doc, "III.4 Aplikasi di INDF 2024: Goodwill sebagai Aset Intangible")
    add_body_paragraph(
        doc,
        "PT Indofood Sukses Makmur Tbk mencatat Goodwill sebesar Rp52,2 triliun per "
        "31 Desember 2024, mewakili sekitar 26 persen dari total aset konsolidasi. "
        "Di bawah SFAC 6, Goodwill memenuhi definisi aset karena berasal dari transaksi masa "
        "lalu (akuisisi entitas anak) dan mengandung manfaat ekonomi masa depan yang diharapkan. "
        "Namun di bawah definisi SFAC 8 Ch.4, Goodwill dapat juga dipandang sebagai "
        "'present right' atas manfaat sinergistik yang diperoleh dari pengendalian entitas "
        "yang diakuisisi."
    )

    add_caption(doc, "Bagan 3.1 — Waterfall Comprehensive Income PT INDF 2024 (Rp Triliun, Ilustratif)")
    p_wf = doc.add_paragraph()
    p_wf.alignment = 1
    p_wf.add_run().add_picture(generate_bagan_ci_waterfall(), width=Inches(5.5))

    add_footnote_text(
        doc,
        "³ Catatan waterfall: nilai beban pokok dan beban operasi bersifat estimasi ilustratif "
        "— angka aktual harus diverifikasi dari sumber: sources/indf-2024-ar.pdf. "
        "Goodwill Rp52,2T dan Net Sales Rp115,79T telah dikonfirmasi dari AR 2024. "
        "BC4.7 mengacu Basis for Conclusions SFAC 8 Ch.4 (Juli 2024)."
    )
