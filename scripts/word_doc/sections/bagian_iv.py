from docx.shared import Inches
from scripts.word_doc.styles import (
    add_section_heading, add_sub_heading, add_body_paragraph,
    add_caption, add_table_with_headers, add_footnote_text,
)
from scripts.word_doc.visuals import generate_bagan_entry_exit_tree, generate_bagan_fs_hierarchy


def build(doc) -> None:
    """Bagian IV — Pengakuan dan Pengukuran"""

    add_section_heading(doc, "IV. PENGAKUAN DAN PENGUKURAN")

    add_sub_heading(doc, "IV.1 SFAC 5 (1984): Empat Kriteria Pengakuan dan Kritiknya")
    add_body_paragraph(
        doc,
        "SFAC 5 (1984) menetapkan empat kriteria untuk pengakuan suatu item dalam laporan "
        "keuangan: (1) item memenuhi definisi elemen laporan keuangan; (2) memiliki atribut "
        "yang dapat diukur dengan cukup andal; (3) informasinya relevan bagi pengguna; dan "
        "(4) informasinya dapat diandalkan (representasi tepat, dapat diverifikasi, netral). "
        "Selain itu, SFAC 5 mengidentifikasi lima atribut pengukuran yang dapat digunakan: "
        "Historical Cost, Current Replacement Cost, Current Market Value, Net Realizable Value, "
        "dan Present Value of Future Cash Flows."
    )
    add_body_paragraph(
        doc,
        "Namun para akademisi mengkritik SFAC 5 sebagai 'Achilles' heel of the FASB's conceptual "
        "framework' (Solomons, dikutip dalam Wolk, Dodd & Rozycki 2017, Ch.7). Kritik utama: "
        "SFAC 5 bersifat terlalu deskriptif dan permisif — ia mendeskripsikan praktik yang ada "
        "daripada menetapkan standar normatif yang jelas. Akibatnya, standar ini gagal "
        "memberikan panduan definitif kapan suatu item harus diakui vs. diungkapkan saja, "
        "menjadikannya sumber ambiguitas yang berkelanjutan dalam penetapan standar."
    )
    add_caption(doc, "Tabel 4.1 — Kriteria Pengakuan: SFAC 5 (1984) vs. SFAC 8 Chapter 5 (Juli 2024)")
    add_table_with_headers(
        doc,
        headers=["Kriteria", "SFAC 5 (1984)", "SFAC 8 Ch.5 (Juli 2024)"],
        rows_data=[
            ["Jumlah Kriteria", "4 kriteria", "3 kriteria (lebih sederhana)"],
            ["Kriteria 1", "Memenuhi definisi elemen LK",
             "Memenuhi definisi elemen LK (sama)"],
            ["Kriteria 2", "Dapat diukur dengan atribut relevan yang andal",
             "Informasi relevan dan tersedia (terkait benefit vs. cost)"],
            ["Kriteria 3", "Relevansi — mempengaruhi keputusan pengguna",
             "Faithful representation — merepresentasikan realitas ekonomi"],
            ["Kriteria 4 (SFAC 5 saja)", "Reliability — dapat diukur andal",
             "Tidak ada — telah disederhanakan"],
            ["Prinsip Utama", "Recognition = memenuhi 4 kriteria",
             "RD3: Recognition = 'words + numbers' dalam LK formal"],
            ["Disclosure vs. Recognition", "Tidak dibedakan secara tegas",
             "PR12: Notes ≠ Recognition (tidak dapat menggantikan pengakuan formal)"],
        ],
    )

    add_sub_heading(doc, "IV.2 SFAC 8 Chapter 5 (Juli 2024): RD3 dan Prinsip Recognition")
    add_body_paragraph(
        doc,
        "SFAC 8 Chapter 5 (Juli 2024) menyederhanakan kerangka pengakuan menjadi tiga kriteria "
        "dan memperkenalkan Recognized Definition 3 (RD3): sebuah item diakui ketika informasi "
        "tentang item tersebut menghasilkan manfaat yang melebihi biaya penyajiannya, baik "
        "dalam Laporan Posisi Keuangan maupun Laporan Kinerja Keuangan. Implikasi RD3 adalah "
        "bahwa pengakuan (recognition) selalu melibatkan representasi ganda: dalam kata-kata "
        "(label dalam laporan) dan dalam angka (nilai terukur). Ini secara eksplisit membedakan "
        "recognition dari mere disclosure."
    )

    add_sub_heading(doc, "IV.3 SFAC 8 Chapter 6 (Juli 2024): Dua Sistem Pengukuran")
    add_body_paragraph(
        doc,
        "SFAC 8 Chapter 6 (Juli 2024) mengorganisasi atribut pengukuran ke dalam dua sistem "
        "yang berbeda secara konseptual: Entry Price (harga masuk) dan Exit Price (harga keluar). "
        "Entry Price mencerminkan perspektif akuisisi — berapa yang dibayarkan untuk memperoleh "
        "aset atau menerima liabilitas — mencakup Historical Cost, Current Replacement Cost, "
        "dan Value in Use (untuk aset). Exit Price mencerminkan perspektif pasar — berapa yang "
        "akan diterima jika aset dijual atau dibayarkan untuk menyelesaikan liabilitas — "
        "mencakup Fair Value (IFRS 13) dan Net Realizable Value."
    )
    add_body_paragraph(
        doc,
        "Paragraf M30–M34 SFAC 8 Ch.6 memberikan panduan pemilihan basis pengukuran: ketika "
        "suatu aset atau liabilitas memiliki harga pasar yang unik (unique market price) — "
        "artinya harga tersebut spesifik untuk situasi entitas — maka Entry Price lebih tepat "
        "digunakan. Sebaliknya, ketika harga tidak unik (non-unique), artinya aset/liabilitas "
        "dapat diperdagangkan di pasar yang aktif dan transparan, Exit Price (Fair Value) "
        "menjadi basis yang lebih representatif."
    )
    add_caption(doc, "Tabel 4.2 — Lima Atribut Pengukuran SFAC 5 dan Pemetaan ke SFAC 8 Ch.6")
    add_table_with_headers(
        doc,
        headers=["Atribut (SFAC 5)", "Sistem (SFAC 8 Ch.6)", "Contoh Aplikasi di INDF"],
        rows_data=[
            ["Historical Cost", "Entry Price", "Aset tetap (pabrik Bogasari, mesin CPO)"],
            ["Current Replacement Cost", "Entry Price", "Persediaan bahan baku (terigu, CPO)"],
            ["Net Realizable Value (entry)", "Entry Price",
             "Persediaan barang jadi (mi instan siap jual)"],
            ["Fair Value / Market Value", "Exit Price",
             "Aset keuangan diperdagangkan, investasi"],
            ["Present Value of Future CF", "Exit Price",
             "Goodwill impairment test (discounted CF)"],
        ],
    )

    add_caption(doc, "Bagan 4.1 — Pohon Keputusan Basis Pengukuran (SFAC 8 Ch.6, M30–M34)")
    p_tree = doc.add_paragraph()
    p_tree.alignment = 1
    p_tree.add_run().add_picture(generate_bagan_entry_exit_tree(), width=Inches(5.5))

    add_sub_heading(doc, "IV.4 SFAC 8 Chapter 7 (Juli 2024): Penyajian dan Hierarki LK")
    add_body_paragraph(
        doc,
        "SFAC 8 Chapter 7 mengatur prinsip penyajian informasi dalam laporan keuangan. "
        "Prinsip kunci PR12 menegaskan: 'Notes to financial statements cannot substitute for "
        "recognition.' Catatan kaki (notes) hanya berfungsi sebagai suplemen penjelasan — "
        "bukan pengganti pengakuan formal yang memerlukan words + numbers dalam laporan utama. "
        "Implikasinya bagi praktik: entitas tidak dapat menghindari pengakuan suatu kewajiban "
        "dengan cara mengungkapkannya hanya dalam catatan kaki."
    )
    add_body_paragraph(
        doc,
        "Paragraf BC7.21 menyatakan bahwa 'tidak terdapat basis konseptual yang jelas untuk "
        "memisahkan item ke dalam Other Comprehensive Income (OCI) versus Net Income.' "
        "FASB mengakui bahwa pemisahan OCI lebih bersifat pragmatis dan historis, bukan "
        "merupakan prinsip konseptual yang koheren. Ini menjelaskan mengapa banyak analis "
        "keuangan menggunakan Total Comprehensive Income (bukan Net Income saja) sebagai "
        "ukuran kinerja yang lebih komprehensif."
    )
    add_caption(doc, "Bagan 4.2 — Hierarki Laporan Keuangan: SFAC 8 Ch.7")
    p_hier = doc.add_paragraph()
    p_hier.alignment = 1
    p_hier.add_run().add_picture(generate_bagan_fs_hierarchy(), width=Inches(4.5))

    add_footnote_text(
        doc,
        "⁴ RD3 = Recognized Definition 3, SFAC 8 Ch.5. M30–M34 = paragraf pengukuran "
        "SFAC 8 Ch.6. PR12 = paragraf penyajian SFAC 8 Ch.7. BC7.21 = Basis for Conclusions "
        "SFAC 8 Ch.7. Sumber: SFAC 8 (September 2024); Wolk et al. (2017) Ch.7."
    )
