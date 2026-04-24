from docx.shared import Inches
from scripts.word_doc.styles import (
    add_section_heading, add_sub_heading, add_body_paragraph,
    add_caption, add_table_with_headers, add_footnote_text,
)
from scripts.word_doc.visuals import (
    generate_bagan_indf_org, generate_bagan_kepemilikan_pie,
    generate_bagan_matrix_relevance_fr,
)


def build(doc) -> None:
    """Bagian VII — Studi Kasus: PT Indofood Sukses Makmur Tbk (INDF) 2024"""

    add_section_heading(
        doc,
        "VII. STUDI KASUS: PT INDOFOOD SUKSES MAKMUR TBK (INDF) 2024"
    )

    add_sub_heading(doc, "VII.1 Profil Perusahaan: Struktur Konglomerat INDF")
    add_body_paragraph(
        doc,
        "PT Indofood Sukses Makmur Tbk (INDF) merupakan salah satu perusahaan makanan terbesar "
        "di Asia Tenggara dengan struktur konglomerat yang terdiri dari empat divisi utama: "
        "Bogasari (penggilingan terigu), Agribisnis (kelapa sawit dan karet), Consumer Branded "
        "Products/CBP (terutama melalui PT Indofood CBP Sukses Makmur Tbk/ICBP yang memproduksi "
        "mi instan Indomie), dan Distribusi. Kerumitan struktur ini — dengan entitas induk, "
        "entitas anak yang tercatat di bursa (ICBP), dan kepentingan nonpengendali yang "
        "substansial — menjadikan INDF sebagai kasus studi yang kaya untuk aplikasi kerangka "
        "konseptual FASB."
    )
    add_body_paragraph(
        doc,
        "Per 31 Desember 2024, INDF mencatat pendapatan bersih (net sales) sebesar "
        "Rp115,79 triliun. Pemegang saham pengendali utama adalah First Pacific Limited, "
        "sebuah perusahaan investasi yang berkantor pusat di Hong Kong, dengan kepemilikan "
        "sebesar 50,07 persen. Laporan keuangan konsolidasi INDF 2024 diaudit oleh "
        "Kantor Akuntan Publik Purwantono, Sungkoro & Surja — anggota Ernst & Young Global."
    )
    add_caption(doc, "Bagan 6.1 — Struktur Konglomerat dan Angka Kunci PT INDF (2024)")
    p_org = doc.add_paragraph()
    p_org.alignment = 1
    p_org.add_run().add_picture(generate_bagan_indf_org(), width=Inches(5.8))

    add_caption(doc, "Bagan 6.2 — Struktur Kepemilikan PT INDF per 31 Desember 2024")
    p_pie = doc.add_paragraph()
    p_pie.alignment = 1
    p_pie.add_run().add_picture(generate_bagan_kepemilikan_pie(), width=Inches(4.0))

    add_sub_heading(doc, "VII.2 Identifikasi Primary Users: Konteks INDF")
    add_body_paragraph(
        doc,
        "Menerapkan konsep primary users (SFAC 8 Ch.1) pada INDF menghasilkan identifikasi "
        "yang spesifik. Primary users utama INDF adalah: (1) First Pacific Ltd. sebagai "
        "pemegang saham pengendali yang memonitor kinerja dan nilai investasinya; "
        "(2) investor publik dan institusional pemegang saham minoritas yang bergantung pada "
        "GPFS untuk keputusan beli/jual/tahan; (3) pemegang obligasi dan bank kreditur INDF "
        "yang menilai kemampuan pembayaran utang; dan (4) pemegang kepentingan nonpengendali "
        "(Non-Controlling Interest/NCI) sebesar Rp43,077 triliun yang merupakan pemegang "
        "saham minoritas di entitas anak INDF (terutama ICBP)."
    )
    add_caption(doc, "Tabel 6.0 — Identifikasi Primary Users PT INDF 2024 (SFAC 8 Ch.1)")
    add_table_with_headers(
        doc,
        headers=["Kategori Primary User", "Pihak Spesifik (INDF)", "Keputusan Utama",
                 "Informasi Kunci"],
        rows_data=[
            ["Pemegang saham pengendali", "First Pacific Ltd. (50,07%)",
             "Monitoring kinerja investasi", "Net Income, Dividen, EPS"],
            ["Investor publik/minoritas", "Pemegang saham di BEI",
             "Beli/jual/tahan saham INDF", "EPS, P/E, Net Sales"],
            ["Kreditur & pemegang obligasi", "Bank sindikasi, pemegang obligasi",
             "Penilaian kemampuan bayar utang", "Arus kas operasi, DER, Coverage"],
            ["Pemegang NCI", "Minoritas ICBP (Rp43,077 T)",
             "Nilai residual entitas anak", "Laba NCI, Ekuitas NCI"],
        ],
    )

    add_sub_heading(doc, "VII.3 Relevansi Informasi Keuangan INDF")
    add_body_paragraph(
        doc,
        "Informasi keuangan INDF mendemonstrasikan dua dimensi relevansi SFAC 8 secara "
        "bersamaan. Dari sisi Predictive Value: tren EPS 2020–2024 (Rp735 → Rp873 → Rp724 "
        "→ Rp928 → Rp984) memberikan basis bagi analis untuk memproyeksikan profitabilitas "
        "masa depan — tren kenaikan jangka panjang dengan volatilitas 2022 yang dapat "
        "dijelaskan oleh kenaikan harga komoditas (CPO, terigu). Dari sisi Confirmatory Value: "
        "Net Sales Rp115,79 triliun 2024 mengkonfirmasi proyeksi analis mengenai pemulihan "
        "volume penjualan pasca-pandemi dan normalisasi harga bahan baku."
    )

    add_sub_heading(doc, "VII.4 Faithful Representation dalam Laporan Konsolidasi INDF")
    add_body_paragraph(
        doc,
        "Tantangan Faithful Representation yang paling kompleks di INDF berkaitan dengan "
        "konsolidasi laporan keuangan empat divisi dengan model bisnis yang berbeda: "
        "penggilingan terigu (Bogasari, margin relatif stabil), perkebunan (Agribisnis, "
        "sangat terpengaruh harga komoditas), consumer goods (CBP/ICBP, margin konsisten "
        "tinggi), dan distribusi. Kerangka SFAC 8 mensyaratkan bahwa laporan konsolidasi "
        "merepresentasikan realitas ekonomi group secara keseluruhan, termasuk dampak "
        "kepentingan nonpengendali."
    )
    add_body_paragraph(
        doc,
        "NCI sebesar Rp43,077 triliun mencerminkan kepentingan pemegang saham minoritas "
        "entitas anak (terutama ICBP) yang tidak dikendalikan INDF. Penyajian NCI secara "
        "terpisah dalam laporan konsolidasi — sesuai PSAK 65 dan SFAC 8 Ch.7 — merupakan "
        "implementasi Faithful Representation: pembaca laporan dapat memahami bahwa bagian "
        "substansial ekuitas bukan milik pemegang saham INDF (parent)."
    )

    add_sub_heading(doc, "VII.5 Aset dan Goodwill: Implikasi Definisi SFAC 8 Ch.4")
    add_body_paragraph(
        doc,
        "Goodwill INDF sebesar Rp52,2 triliun (sekitar 26 persen dari total aset konsolidasi) "
        "merupakan salah satu Goodwill terbesar di antara emiten Indonesia. Goodwill ini "
        "muncul terutama dari akuisisi Bogasari dan berbagai entitas anak lainnya. "
        "Di bawah definisi SFAC 8 Ch.4, Goodwill merupakan 'present right' atas manfaat "
        "sinergistik yang diperoleh dari pengendalian entitas yang diakuisisi — sebuah "
        "interpretasi yang lebih kuat dari basis konseptual Goodwill dibandingkan SFAC 6."
    )
    add_caption(doc, "Tabel 6.1 — Pengujian Penurunan Nilai (Impairment Test) Goodwill INDF 2024")
    add_table_with_headers(
        doc,
        headers=["Unit Penghasil Kas (UPK)", "Nilai Tercatat (Rp T)", "Metode Uji", "Indikasi"],
        rows_data=[
            ["Bogasari (Tepung Terigu)", "[lihat AR 2024]", "Value in Use (DCF)",
             "[lihat AR 2024]"],
            ["Agribisnis (CPO & Karet)", "[lihat AR 2024]", "Value in Use (DCF)",
             "[lihat AR 2024]"],
            ["CBP / ICBP (Mi Instan)", "[lihat AR 2024]", "Value in Use (DCF)",
             "[lihat AR 2024]"],
            ["Total Goodwill INDF", "Rp 52,2 T", "Berbasis DCF", "Tidak ada penurunan nilai 2024"],
        ],
    )
    add_sub_heading(doc, "VII.6 Liabilitas dan Ekuitas: Klasifikasi NCI")
    add_body_paragraph(
        doc,
        "Salah satu isu klasifikasi yang relevan di INDF adalah apakah kepentingan "
        "nonpengendali (NCI) harus diklasifikasikan sebagai Ekuitas atau sebagai suatu bentuk "
        "Liabilitas. Kerangka konseptual SFAC 8 Ch.4 — sejalan dengan PSAK 65 — menempatkan "
        "NCI sebagai bagian dari Ekuitas konsolidasi (bukan liabilitas), karena NCI merupakan "
        "residual interest pemegang saham minoritas entitas anak, bukan kewajiban kontraktual "
        "entitas induk. NCI INDF sebesar Rp43,077 triliun dilaporkan sebagai bagian Ekuitas "
        "dalam Laporan Posisi Keuangan konsolidasi."
    )

    add_sub_heading(doc, "VII.7 Pengakuan di INDF: Tiga Kriteria SFAC 8 Ch.5")
    add_body_paragraph(
        doc,
        "Menerapkan tiga kriteria pengakuan SFAC 8 Ch.5 pada item-item laporan keuangan INDF "
        "menghasilkan penilaian sebagai berikut. Pendapatan (Revenue) memenuhi ketiga kriteria: "
        "definisi elemen terpenuhi (pendapatan dari penjualan barang), informasi relevan "
        "(prediktif arus kas masa depan), dan representasi tepat (pendapatan direalisasi dari "
        "transaksi aktual dengan pihak ketiga). Goodwill memenuhi kriteria definisi (present "
        "right atas sinergi) dan relevansi, namun representasi tepatnya lebih dipertanyakan "
        "karena nilai Goodwill sangat bergantung pada asumsi DCF yang subjektif."
    )

    add_sub_heading(doc, "VII.8 Pengukuran di INDF: Pemetaan Entry vs. Exit Price")
    add_caption(doc, "Tabel 6.2 — Peta Basis Pengukuran Item LK INDF 2024 (SFAC 8 Ch.6)")
    add_table_with_headers(
        doc,
        headers=["Item LK", "Basis Pengukuran", "Sistem", "Referensi SFAC 8 Ch.6"],
        rows_data=[
            ["Aset Tetap (mesin, pabrik)", "Historical Cost — net of depreciation",
             "Entry Price", "M30: harga unik"],
            ["Goodwill", "Cost less accumulated impairment",
             "Entry Price", "M30: tidak ada pasar aktif"],
            ["Persediaan (terigu, CPO)", "Cost atau NRV, yang lebih rendah",
             "Entry Price", "M31: replacement cost"],
            ["Investasi efek diperdagangkan", "Fair Value (IFRS 13)",
             "Exit Price", "M34: harga tidak unik"],
            ["Piutang usaha", "Amortized Cost",
             "Entry Price", "M30: nilai khusus entitas"],
            ["Liabilitas keuangan jangka panjang", "Amortized Cost (EIR)",
             "Entry Price", "M30: harga khusus pada originasi"],
            ["Aset Biologis (kelapa sawit)", "Fair Value less cost to sell",
             "Exit Price", "M34: pasar komoditas aktif"],
        ],
    )

    add_sub_heading(doc, "VII.9 Penyajian (Presentation): SFAC 8 Ch.7 di Laporan INDF")
    add_body_paragraph(
        doc,
        "Laporan keuangan konsolidasi INDF 2024 menyajikan komponen utama sesuai dengan "
        "hierarki penyajian SFAC 8 Ch.7: Laporan Posisi Keuangan, Laporan Laba Rugi dan "
        "Penghasilan Komprehensif Lain (yang memisahkan Net Income dari OCI), Laporan Arus Kas, "
        "Laporan Perubahan Ekuitas, dan Catatan atas Laporan Keuangan. Penerapan PR12 "
        "(Notes ≠ Recognition) terlihat dalam cara INDF mengakui kewajiban imbalan kerja "
        "(PSAK 24) langsung dalam laporan keuangan, bukan hanya dalam catatan — suatu praktik "
        "yang selaras dengan prinsip bahwa pengakuan formal tidak dapat digantikan oleh "
        "pengungkapan semata."
    )

    add_sub_heading(doc, "VII.10 Catatan atas LK: Empat Limitasi dan D32 di Konteks INDF")
    add_body_paragraph(
        doc,
        "SFAC 8 Chapter 8 (Agustus 2023) mengidentifikasi empat limitasi inherent dari catatan "
        "atas laporan keuangan: (1) catatan tidak dapat menggantikan pengakuan; (2) catatan "
        "dapat menjadi terlalu panjang sehingga mengurangi kegunaan (information overload); "
        "(3) beberapa informasi dalam catatan mungkin terlalu sensitif untuk diungkapkan secara "
        "penuh (D32: adverse consequences); dan (4) catatan kurang terstandarisasi "
        "dibandingkan laporan utama. Catatan atas LK INDF 2024 mencakup lebih dari 60 "
        "item pengungkapan — termasuk penjelasan Goodwill, analisis sensitivitas DCF, "
        "dan transaksi pihak berelasi — yang relevan dengan limitasi (2) di atas."
    )
    add_caption(doc, "Tabel 6.3 — Penilaian Disclosure INDF 2024 vs. Standar SFAC 8 Ch.8")
    add_table_with_headers(
        doc,
        headers=["Area Disclosure", "Standar SFAC 8 Ch.8", "Praktik INDF 2024", "Penilaian"],
        rows_data=[
            ["Goodwill — basis pengukuran", "Wajib diungkapkan metode & asumsi",
             "Diungkapkan di Catatan (DCF, WACC)", "Sesuai"],
            ["NCI — komposisi kepemilikan", "Wajib diungkapkan terpisah",
             "Diungkapkan di Neraca & Catatan", "Sesuai"],
            ["Risiko keuangan (FX, bunga)", "Wajib kuantitatif (IFRS 7)",
             "Analisis sensitivitas tersedia", "Sesuai"],
            ["Transaksi pihak berelasi", "Wajib diungkapkan nilai & syarat",
             "Tersedia di Catatan", "Sesuai"],
            ["Asumsi DCF impairment test", "Wajib — namun rentan D32",
             "Sebagian diungkapkan (range WACC)",
             "Parsial — potensi information gap"],
        ],
    )

    add_caption(doc, "Bagan 6.3 — Matrix Relevansi × Faithful Representation: Item LK INDF 2024")
    p_mat = doc.add_paragraph()
    p_mat.alignment = 1
    p_mat.add_run().add_picture(generate_bagan_matrix_relevance_fr(), width=Inches(5.0))

    add_footnote_text(
        doc,
        "⁷ Semua angka INDF (Net Sales Rp115,79T; Goodwill Rp52,2T; NCI Rp43,077T; "
        "EPS 2024 Rp984) bersumber dari Annual Report PT Indofood Sukses Makmur Tbk 2024 "
        "(sources/indf-2024-ar.pdf). Nilai per UPK Goodwill dalam Tabel 6.1 harus diverifikasi "
        "implementor dari halaman Catatan LK yang relevan. "
        "D32 = paragraf D32, SFAC 8 Ch.8 (Agustus 2023) — 'Adverse Consequences of Disclosure.'"
    )
