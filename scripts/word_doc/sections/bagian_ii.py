from docx.shared import Inches
from scripts.word_doc.styles import (
    add_section_heading, add_sub_heading, add_body_paragraph,
    add_caption, add_table_with_headers, add_footnote_text,
)
from scripts.word_doc.visuals import (
    generate_bagan_qc_sfac2, generate_bagan_qc_sfac8, generate_grafik_eps_indf,
)


def build(doc) -> None:
    """Bagian II — Karakteristik Kualitatif Informasi Keuangan"""

    add_section_heading(doc, "II. KARAKTERISTIK KUALITATIF INFORMASI KEUANGAN")

    add_sub_heading(doc, "II.1 SFAC 2 (1980): Hierarki Keandalan (Reliability)")
    add_body_paragraph(
        doc,
        "SFAC 2 (1980) mendefinisikan dua karakteristik kualitatif fundamental informasi "
        "keuangan: Relevansi (Relevance) dan Keandalan (Reliability). Keandalan diurai lebih "
        "lanjut menjadi tiga komponen: Representational Faithfulness (ketepatan representasi), "
        "Verifiability (dapat diverifikasi pihak independen), dan Neutrality (kebebasan dari "
        "bias). Posisi Verifiability sebagai komponen setara dari Keandalan — bukan subordinat — "
        "mencerminkan paradigma akuntansi era 1980-an yang sangat menekankan objektivitas "
        "dan kemampuan audit."
    )
    add_body_paragraph(
        doc,
        "Selain dua karakteristik fundamental tersebut, SFAC 2 menempatkan Komparabilitas "
        "(Comparability) sebagai karakteristik interaktif yang meningkatkan kegunaan "
        "informasi ketika digunakan bersama relevansi dan keandalan. Materialitas berfungsi "
        "sebagai ambang batas (threshold) yang menentukan apakah suatu informasi cukup "
        "signifikan untuk diungkapkan — bukan sebagai karakteristik kualitatif itu sendiri."
    )
    add_caption(doc, "Bagan 2.1 — Hierarki QC SFAC 2 (1980): Reliability sebagai Karakteristik Fundamental")
    p1 = doc.add_paragraph()
    p1.alignment = 1
    p1.add_run().add_picture(generate_bagan_qc_sfac2(), width=Inches(5.5))

    add_sub_heading(doc, "II.2 SFAC 8 (2010): Perubahan Fundamental ke Faithful Representation")
    add_body_paragraph(
        doc,
        "SFAC 8 (2010) melakukan perubahan konseptual yang paling signifikan dalam sejarah "
        "kerangka konseptual FASB: penggantian istilah Reliability dengan Faithful Representation "
        "(representasi tepat) sebagai karakteristik kualitatif fundamental kedua. Perubahan ini "
        "bukan sekadar pergantian nama — ia mencerminkan pergeseran paradigma mendasar dari "
        "pertanyaan 'dapatkah kita memverifikasi informasi ini?' menjadi 'apakah informasi ini "
        "merepresentasikan realitas ekonomi yang mendasarinya dengan tepat?' "
        "(Wolk, Dodd & Rozycki 2017, hal. 170–172)."
    )
    add_body_paragraph(
        doc,
        "Faithful Representation dalam SFAC 8 terdiri dari tiga sub-komponen: "
        "Completeness (kelengkapan informasi yang diperlukan untuk pemahaman), "
        "Neutrality (bebas dari bias dan pilihan kebijakan yang menguntungkan pihak tertentu), "
        "dan Free from Error (tidak terdapat kesalahan material dalam deskripsi fenomena). "
        "Sementara itu, empat karakteristik peningkat (Enhancing QCs) — Comparability, "
        "Verifiability, Timeliness, dan Understandability — berfungsi memaksimalkan kegunaan "
        "informasi yang sudah memenuhi dua QC fundamental."
    )
    add_caption(doc, "Bagan 2.2 — Hierarki QC SFAC 8 (2010): Faithful Representation sebagai Fundamental")
    p2 = doc.add_paragraph()
    p2.alignment = 1
    p2.add_run().add_picture(generate_bagan_qc_sfac8(), width=Inches(5.5))

    add_sub_heading(doc, "II.3 Analisis Komparatif: SFAC 2 vs. SFAC 8 — Perubahan Mendasar")
    add_body_paragraph(
        doc,
        "Penurunan status Verifiability dari komponen Keandalan (SFAC 2) menjadi Enhancing QC "
        "(SFAC 8) merupakan salah satu perubahan paling dramatis dalam evolusi kerangka "
        "konseptual. Implikasinya: informasi yang sangat mudah diverifikasi (seperti harga "
        "perolehan historis) tidak serta-merta superior dibandingkan informasi yang sulit "
        "diverifikasi tetapi lebih representatif (seperti nilai wajar aset yang diperdagangkan "
        "di pasar aktif). SFAC 8 secara eksplisit menempatkan representasi realitas ekonomi "
        "di atas kemudahan verifikasi."
    )
    add_body_paragraph(
        doc,
        "Penghapusan Conservatism (Konservatisme) dari SFAC 8 merupakan perubahan kedua yang "
        "paling diperdebatkan. FASB dalam Basis for Conclusions paragraf BC3.27–BC3.29 "
        "menjelaskan bahwa konservatisme bertentangan secara fundamental dengan Neutrality — "
        "komponen Faithful Representation. Sebuah kerangka konseptual tidak dapat sekaligus "
        "menuntut netralitas dan mengizinkan kecenderungan sistematis ke arah pesimisme "
        "(under-statement aset, over-statement liabilitas)."
    )
    add_caption(doc, "Tabel 2.1 — Perbandingan Hierarki QC: SFAC 2 (1980) vs. SFAC 8 (2010)")
    add_table_with_headers(
        doc,
        headers=["Dimensi", "SFAC 2 (1980)", "SFAC 8 (2010)", "Perubahan"],
        rows_data=[
            ["QC Fundamental 1", "Relevansi", "Relevansi", "Dipertahankan"],
            ["QC Fundamental 2", "Reliability (Keandalan)", "Faithful Representation",
             "Diganti nama + direkonstitusi"],
            ["Status Verifiability", "Komponen Keandalan (setara FR & Neutrality)",
             "Enhancing QC (subordinat)", "Diturunkan drastis"],
            ["Konservatisme", "Disebutkan sebagai 'konvensi yang berguna'",
             "Dihapus sepenuhnya (BC3.27–29)", "Konflik dengan Neutrality"],
            ["Materialitas", "Threshold bawah relevansi", "Threshold bawah relevansi",
             "Dipertahankan"],
        ],
    )

    add_sub_heading(doc, "II.4 Relevansi dalam Konteks INDF: Tren EPS sebagai Ilustrasi")
    add_body_paragraph(
        doc,
        "Konsep Relevansi dalam SFAC 8 mencakup dua dimensi: Predictive Value (nilai "
        "prediktif — membantu pengguna memproyeksikan hasil masa depan) dan Confirmatory Value "
        "(nilai konfirmasi — mengkonfirmasi atau mengoreksi ekspektasi sebelumnya). Tren "
        "Earnings Per Share (EPS) PT Indofood Sukses Makmur Tbk selama 2020–2024 "
        "mendemonstrasikan kedua dimensi ini: EPS berfungsi sebagai indikator prediktif "
        "profitabilitas per saham sekaligus konfirmasi realisasi kinerja terhadap proyeksi analis."
    )
    add_caption(doc, "Grafik 2.1 — Tren EPS PT INDF 2020–2024 (Rp per saham)")
    p3 = doc.add_paragraph()
    p3.alignment = 1
    p3.add_run().add_picture(generate_grafik_eps_indf(), width=Inches(5.0))

    add_footnote_text(
        doc,
        "² Sumber EPS: Annual Report PT Indofood Sukses Makmur Tbk 2024. "
        "BC3.27–BC3.29 mengacu pada Basis for Conclusions SFAC 8 (September 2024). "
        "Wolk et al. (2017) Ch.7 Exhibit 7.1 (SFAC 2) dan Exhibit 7.5 (SFAC 8)."
    )
