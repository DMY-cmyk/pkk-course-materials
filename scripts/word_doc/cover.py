import os
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from scripts.word_doc.styles import set_run_font, _add_para_border


NAVY = "1E3A5F"


def _blue_rule(doc, side):
    """Paragraph with a thick navy border on top or bottom."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _add_para_border(p, [side], sz="24", color=NAVY)
    return p


def _centered(doc, text, size_pt, bold=False, italic=False,
               space_before=6, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size_pt, bold=bold, italic=italic)
    return p


def build_cover(doc, logo_path: str) -> None:
    """
    Build the formal cover page (Opsi A — Formal Akademik Klasik).
    Mutates `doc` in place.
    """
    # Top blue rule
    _blue_rule(doc, "top")

    # Institution header
    _centered(doc,
              "PROGRAM MAGISTER AKUNTANSI — STIE YKPN YOGYAKARTA",
              size_pt=11, bold=True, space_before=18, space_after=12)

    # Logo
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(6)
        p_logo.paragraph_format.space_after = Pt(14)
        run = p_logo.add_run()
        run.add_picture(logo_path, height=Inches(1.4))
    else:
        _centered(doc, "[Logo STIE YKPN]", 10, italic=True)

    # Document title (two lines)
    _centered(doc,
              "KERANGKA KONSEPTUAL FASB:\nFONDASI STANDAR PELAPORAN KEUANGAN",
              size_pt=16, bold=True, space_before=10, space_after=8)

    # Subtitle
    _centered(doc,
              "PT Indofood Sukses Makmur Tbk (INDF) sebagai Studi Kasus",
              size_pt=12, italic=True, space_before=4, space_after=20)

    # Course info
    _centered(doc, "MNK202 Pelaporan Keuangan Korporat",
              size_pt=11, space_before=8, space_after=4)
    _centered(doc, "Kelompok 3  |  Tahun Akademik 2025/2026",
              size_pt=11, space_before=2, space_after=16)

    # Member list
    members = [
        ("1.", "Efri Nurmalinda",          "NIM: 1225 01049"),
        ("2.", "Dzaki Muhammad Yusfian",   "NIM: 1225 01079"),
        ("3.", "Nuradila",                 "NIM: 1225 01080"),
        ("4.", "Achmad Dimas Wibawa",      "NIM: 1225 01083"),
        ("5.", "Adinda Putri Dewi",        "NIM: 1225 01086"),
        ("6.", "Setiabudi Yudha Pratama",  "NIM: 1225 01098"),
    ]
    for no, name, nim in members:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{no:<4}{name:<35}{nim}")
        set_run_font(run, 12)

    # Spacer
    _centered(doc, "", size_pt=10, space_before=20, space_after=0)

    # Bottom blue rule
    _blue_rule(doc, "bottom")
