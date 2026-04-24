import os
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from scripts.word_doc.styles import set_run_font, add_page_break


def build_bibliography(doc, txt_path: str) -> None:
    """Parse Daftar Pustaka .txt and append formatted bibliography to doc."""
    add_page_break(doc)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(16)
    run = p_title.add_run("DAFTAR PUSTAKA")
    set_run_font(run, 14, bold=True)

    if not os.path.exists(txt_path):
        p_warn = doc.add_paragraph()
        run_w = p_warn.add_run(
            f"[Peringatan: File daftar pustaka tidak ditemukan di: {txt_path}]"
        )
        set_run_font(run_w, 11, italic=True)
        return

    with open(txt_path, encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        stripped = line.rstrip("\r\n")
        if not stripped.strip():
            doc.add_paragraph()
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)

        if stripped.endswith(":") or (stripped.isupper() and len(stripped) < 60):
            p.paragraph_format.space_before = Pt(10)
            run = p.add_run(stripped)
            set_run_font(run, 11, bold=True)
        else:
            p.paragraph_format.first_line_indent = Pt(-24)
            p.paragraph_format.left_indent = Pt(24)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(stripped)
            set_run_font(run, 11)
