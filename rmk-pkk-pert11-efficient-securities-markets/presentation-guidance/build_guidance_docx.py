"""
build_guidance_docx.py — renders panduan-presentasi-kelompok3-pert11.md into a
formatted Word document for sharing with Kelompok 3.

Source of truth: panduan-presentasi-kelompok3-pert11.md (git-tracked, editable).
Output: output/Panduan Presentasi Kelompok 3 - Pert. 11.docx (Calibri 12 / 1.5 / A4).

Markdown subset supported: # title, ## H1, ### H2, '- ' bullets, GitHub tables,
blank-line-separated paragraphs, **bold**, *italic*.
"""
import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor

FONT = "Calibri"
HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "panduan-presentasi-kelompok3-pert11.md")
OUT_DIR = os.path.join(HERE, "output")
OUT_NAME = "Panduan Presentasi Kelompok 3 - Pert. 11.docx"

INLINE_RE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*)')


def add_runs(paragraph, text):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        else:
            r = paragraph.add_run(tok[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def style_base(doc):
    s = doc.styles["Normal"]
    s.font.name = FONT
    s.font.size = Pt(12)
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    s.paragraph_format.space_after = Pt(6)
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)  # A4 portrait
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(2.5))


def heading(doc, text, size, color=None, space_before=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def parse_table(lines, i, doc):
    rows = []
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        rows.append(lines[i].strip())
        i += 1
    cells = [[c.strip() for c in r.strip("|").split("|")] for r in rows]
    # drop the |---| separator row
    body = [r for r in cells if not all(set(c) <= set("-: ") for c in r)]
    if not body:
        return i
    t = doc.add_table(rows=0, cols=len(body[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(body):
        tr = t.add_row().cells
        for ci, val in enumerate(row):
            if ci >= len(tr):
                continue
            cp = tr[ci].paragraphs[0]
            add_runs(cp, val)
            for run in cp.runs:
                run.font.name = FONT
                run.font.size = Pt(10.5)
                if ri == 0:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return i


def build():
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    style_base(doc)
    accent = RGBColor(0x1F, 0x3A, 0x5F)

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue
        if line.startswith("# "):
            p = heading(doc, line[2:].strip(), 16, accent, space_before=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            heading(doc, line[3:].strip(), 13.5, accent, space_before=14)
        elif line.startswith("### "):
            heading(doc, line[4:].strip(), 11.5, space_before=8)
        elif line.lstrip().startswith("|"):
            i = parse_table(lines, i, doc)
            continue
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            add_runs(p, stripped)
        i += 1

    os.makedirs(OUT_DIR, exist_ok=True)
    out = os.path.join(OUT_DIR, OUT_NAME)
    doc.save(out)
    print("Saved:", out)
    print("Size: ", f"{os.path.getsize(out):,} bytes")


if __name__ == "__main__":
    build()
