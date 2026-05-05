"""
build_docx.py
Builds 01079_Dzaki Muhammad Yusfian_Tugas Individual 3.docx
using python-docx 1.2.0.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_paragraph_format(para, font_name="Times New Roman", font_size=12,
                          bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          line_spacing_pt=18, space_before_pt=0, space_after_pt=6):
    """Apply standard body formatting to a paragraph."""
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)

    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        # Ensure East Asian font is also set
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        existing = rPr.find(qn('w:rFonts'))
        if existing is not None:
            rPr.remove(existing)
        rPr.insert(0, rFonts)


def add_paragraph(doc, text, font_name="Times New Roman", font_size=12,
                  bold=False, italic=False,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  line_spacing_pt=18, space_before_pt=0, space_after_pt=6):
    """Add a paragraph with given formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic

    pf = para.paragraph_format
    pf.alignment = alignment
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)

    # Force font via XML
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)

    return para


def add_section_heading(doc, text, number_prefix, font_size=13):
    """Add a section heading (Roman numeral level)."""
    full_text = f"{number_prefix}   {text}"
    para = add_paragraph(doc, full_text,
                         font_size=font_size, bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         line_spacing_pt=18,
                         space_before_pt=12, space_after_pt=6)
    return para


def add_subsection_heading(doc, text, letter_prefix):
    """Add a subsection heading (A. B. C. D. level)."""
    full_text = f"{letter_prefix}   {text}"
    para = add_paragraph(doc, full_text,
                         font_size=12, bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         line_spacing_pt=18,
                         space_before_pt=12, space_after_pt=6)
    return para


def add_body_para(doc, text):
    """Add a standard body paragraph."""
    return add_paragraph(doc, text,
                         font_size=12, bold=False,
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         line_spacing_pt=18,
                         space_before_pt=0, space_after_pt=6)


def add_blank(doc):
    """Add a blank paragraph (spacer) — no space_after."""
    para = add_paragraph(doc, "",
                         line_spacing_pt=18,
                         space_before_pt=0, space_after_pt=0)
    return para


def add_daftar_pustaka_entry(doc, text_parts):
    """
    Add a bibliography entry.  text_parts is a list of (text, italic) tuples.
    Hanging indent via first_line_indent = -0.5cm, left = 0.5cm.
    """
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.left_indent = Cm(0.75)
    pf.first_line_indent = Cm(-0.75)

    for text, italic in text_parts:
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.italic = italic
        run.bold = False

        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), "Times New Roman")
        rFonts.set(qn('w:hAnsi'), "Times New Roman")
        rFonts.set(qn('w:cs'), "Times New Roman")
        existing = rPr.find(qn('w:rFonts'))
        if existing is not None:
            rPr.remove(existing)
        rPr.insert(0, rFonts)

    return para


# ---------------------------------------------------------------------------
# Cover helpers
# ---------------------------------------------------------------------------

def add_cover_para(doc, text, font_size=12, bold=False):
    return add_paragraph(doc, text,
                         font_size=font_size, bold=bold,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing_pt=18,
                         space_before_pt=0, space_after_pt=6)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # -----------------------------------------------------------------------
    # Page setup: Letter 8.5 × 11 in, 3 cm margins on all sides
    # -----------------------------------------------------------------------
    section = doc.sections[0]
    section.page_width = Cm(21.59)   # 8.5 in
    section.page_height = Cm(27.94)  # 11 in
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    # Remove default header/footer
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    # -----------------------------------------------------------------------
    # Default document style
    # -----------------------------------------------------------------------
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # -----------------------------------------------------------------------
    # COVER PAGE
    # -----------------------------------------------------------------------
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)

    add_cover_para(doc, "TUGAS INDIVIDU 3", font_size=16, bold=True)
    add_cover_para(doc, "Pelaporan Keuangan Korporat", font_size=14, bold=False)
    add_cover_para(doc,
                   "Perbandingan Kualitas Informasi Akuntansi: SFAC No. 2 (1980) dan SFAC No. 8 (2010)",
                   font_size=14, bold=True)

    add_blank(doc)
    add_blank(doc)

    add_cover_para(doc, "Disusun oleh:", font_size=12, bold=False)
    add_cover_para(doc, "Dzaki Muhammad Yusfian", font_size=12, bold=True)
    add_cover_para(doc, "1225 01079", font_size=12, bold=False)

    add_blank(doc)
    add_blank(doc)

    add_cover_para(doc, "Dosen Pengampu: Prof. Djoko Susanto", font_size=12, bold=False)
    add_cover_para(doc, "Program Studi Magister Akuntansi", font_size=12, bold=False)
    add_cover_para(doc, "5 Mei 2026", font_size=12, bold=True)

    # -----------------------------------------------------------------------
    # Page break before body
    # -----------------------------------------------------------------------
    doc.add_page_break()

    # -----------------------------------------------------------------------
    # SECTION I — Pendahuluan
    # -----------------------------------------------------------------------
    add_section_heading(doc, "Pendahuluan", "I.")

    pendahuluan = [
        "Kerangka konseptual merupakan landasan logis yang mengikat seluruh standar akuntansi ke dalam satu sistem yang koheren. Tanpa kerangka ini, standar-standar individual akan bersifat ad hoc — sekadar respons terhadap tekanan praktis tanpa panduan normatif yang konsisten. Dalam konteks Dewan Standar Akuntansi Keuangan (FASB) Amerika Serikat, kerangka konseptual telah mengalami dua fase utama: era SFAC No. 2 (1980) yang meletakkan dasar kualitas informasi akuntansi, dan era SFAC No. 8 (2010) yang merevisi landasan tersebut secara mendasar sebagai hasil proyek konvergensi bersama dengan Dewan Standar Akuntansi Internasional (IASB).",
        "Transisi dari SFAC No. 2 ke SFAC No. 8 bukan sekadar pembaruan terminologis. Perubahan tersebut mencerminkan reorientasi konseptual: dari kerangka yang memandang kualitas informasi melalui lensa keandalan pengukuran, menuju kerangka yang menempatkan representasi tepat fenomena ekonomi sebagai cita-cita utama. Pergeseran ini membawa implikasi mendalam terhadap bagaimana akuntan, regulator, dan analis memahami dan mengevaluasi laporan keuangan.",
        "Esai ini menelaah kualitas informasi akuntansi menurut SFAC No. 2 (Bagian II), kemudian menurut SFAC No. 8 (Bagian III), sebelum melakukan perbandingan dan analisis kritis terhadap pergeseran konseptual yang terjadi (Bagian IV), dan menutup dengan sintesis evaluatif (Bagian V).",
    ]
    for p in pendahuluan:
        add_body_para(doc, p)

    # -----------------------------------------------------------------------
    # SECTION II — SFAC No. 2
    # -----------------------------------------------------------------------
    add_section_heading(doc,
                        "Kualitas Informasi Akuntansi menurut SFAC No. 2 (1980)",
                        "II.")

    # A.
    add_subsection_heading(doc,
                           "Latar Belakang dan Tujuan SFAC No. 2",
                           "A.")
    add_body_para(doc,
        "SFAC No. 2, Qualitative Characteristics of Accounting Information, diterbitkan FASB pada Mei 1980 sebagai bagian dari proyek kerangka konseptual yang dimulai pascapembentukan FASB pada 1973. Kerangka konseptual ini merespons Laporan Trueblood (1973) yang menggeser orientasi akuntansi dari pertanggungjawaban stewardship menuju decision-usefulness — yakni, informasi keuangan harus berguna bagi para pengambil keputusan ekonomi (FASB, 1980; Wolk et al., 2019). SFAC No. 2 secara khusus mendefinisikan karakteristik-karakteristik yang membuat informasi akuntansi bernilai bagi pengguna."
    )

    # B.
    add_subsection_heading(doc,
                           "Karakteristik Kualitatif Primer: Relevansi dan Keandalan",
                           "B.")
    add_body_para(doc,
        "SFAC No. 2 menetapkan dua karakteristik kualitatif primer yang harus dipenuhi agar informasi berguna: relevansi dan keandalan."
    )
    add_body_para(doc,
        "Relevansi mensyaratkan bahwa informasi mampu membuat perbedaan dalam pengambilan keputusan. Informasi relevan memiliki tiga komponen: nilai prediktif (membantu pengguna memprakirakan hasil di masa depan), nilai umpan balik (membantu pengguna mengkonfirmasi atau merevisi ekspektasi sebelumnya), dan ketepatan waktu (tersedia sebelum kehilangan kapasitasnya untuk memengaruhi keputusan). Perlu dicatat bahwa ketepatan waktu dalam SFAC No. 2 diposisikan sebagai komponen relevansi — sebuah pilihan klasifikasi yang kemudian dipersoalkan dalam SFAC No. 8."
    )
    add_body_para(doc,
        "Keandalan (reliability) mensyaratkan bahwa informasi secara tepat mewakili apa yang diklaim untuk diwakilinya, bebas dari kesalahan dan bias yang material. Keandalan terdiri atas: ketepatan representasional (korespondensi antara angka akuntansi dan fenomena ekonomi yang diwakilinya), verifikabilitas (pengukur independen yang kompeten akan menghasilkan kesimpulan yang sama), dan netralitas (tidak ada bias dalam pemilihan atau penyajian informasi). Ketiga komponen ini diperlakukan sebagai satu kesatuan dalam keandalan."
    )

    # C.
    add_subsection_heading(doc,
                           "Karakteristik Kualitatif Sekunder dan Batasan",
                           "C.")
    add_body_para(doc,
        "Di bawah karakteristik primer, SFAC No. 2 mengidentifikasi komparabilitas (termasuk konsistensi) sebagai karakteristik sekunder — meningkatkan kegunaan informasi tetapi tidak menentukan apakah informasi tersebut berguna secara fundamental."
    )
    add_body_para(doc,
        "Dua batasan penting juga diakui. Pertama, materialitas berperan sebagai ambang batas: informasi hanya perlu diungkapkan jika kelalaian atau kesalahannya dapat memengaruhi keputusan pengguna yang rasional. Kedua, pertimbangan manfaat versus biaya membatasi tuntutan kualitas — manfaat informasi harus melebihi biaya penyediaannya."
    )
    add_body_para(doc,
        "Selain itu, SFAC No. 2 mengakui konservatisme sebagai konvensi yang tepat dalam menghadapi ketidakpastian: ketika dua estimasi sama-sama dapat dipertahankan, akuntan sebaiknya memilih yang lebih konservatif (§95). Pengakuan ini bersifat hibrida: konservatisme dalam SFAC No. 2 tidak secara formal diintegrasikan ke dalam hierarki kualitatif primer maupun sekunder, melainkan ditempatkan sebagai konvensi pragmatis yang berdampingan dengan kerangka tersebut — diterima tanpa diberi status konseptual yang setara. Posisi hibrida inilah yang membuat penghapusannya dalam SFAC No. 8 menjadi tindakan yang melampaui sekadar penataan ulang hierarki: ia adalah pernyataan konseptual bahwa konservatisme tidak dapat direkonsiliasi dengan netralitas sebagai nilai inti."
    )

    # -----------------------------------------------------------------------
    # SECTION III — SFAC No. 8
    # -----------------------------------------------------------------------
    add_section_heading(doc,
                        "Kualitas Informasi Akuntansi menurut SFAC No. 8 (2010)",
                        "III.")

    # A.
    add_subsection_heading(doc,
                           "Konteks Konvergensi FASB-IASB dan Penggantian SFAC No. 2",
                           "A.")
    add_body_para(doc,
        "SFAC No. 8, Conceptual Framework for Financial Reporting, diterbitkan FASB pada September 2010 sebagai hasil proyek kerangka konseptual bersama antara FASB dan IASB — sebuah inisiatif yang berakar pada Norwalk Agreement (2002) yang mewajibkan kedua dewan untuk menyelaraskan standar mereka secara bertahap. SFAC No. 8 sekaligus menggantikan SFAC No. 1 (tujuan pelaporan keuangan) dan SFAC No. 2 (karakteristik kualitatif) dalam satu dokumen terintegrasi yang secara substansial identik dengan Conceptual Framework for Financial Reporting IASB (2018). Penggantian ini bukan hanya pembaruan administratif: ia merefleksikan perubahan paradigmatik dalam cara keduanya memandang kualitas informasi akuntansi (FASB, 2010; IASB, 2018)."
    )

    # B.
    add_subsection_heading(doc,
                           "Karakteristik Fundamental: Relevansi dan Representasi Tepat",
                           "B.")
    add_body_para(doc,
        "SFAC No. 8 mempertahankan relevansi sebagai karakteristik fundamental, tetapi dengan modifikasi penting. Nilai prediktif dipertahankan. Namun nilai umpan balik diganti dengan istilah nilai konfirmatoris (confirmatory value) — perubahan yang lebih dari sekadar terminologis. Istilah ‘nilai umpan balik’ (feedback value) berpotensi menyiratkan bahwa fungsinya bersifat murni retrospektif — sekadar melaporkan apa yang telah terjadi. ‘Nilai konfirmatoris’ menegaskan bahwa fungsi tersebut secara inheren prospektif: informasi mengkonfirmasi atau merevisi ekspektasi yang berorientasi masa depan, sehingga ia tidak berdiri sendiri dari nilai prediktif melainkan beroperasi bersamanya dalam satu siklus pengambilan keputusan yang koheren. Yang lebih signifikan, materialitas — yang dalam SFAC No. 2 berdiri sebagai ambang batas terpisah — kini diredefinisi sebagai aspek dari relevansi: informasi yang tidak material bukan sekadar di bawah ambang batas, melainkan secara definitif tidak relevan bagi pengambil keputusan (FASB, 2010)."
    )
    add_body_para(doc,
        "Perubahan paling fundamental adalah penggantian keandalan (reliability) dengan representasi tepat (faithful representation). SFAC No. 8 mendefinisikan representasi tepat melalui tiga komponen: kelengkapan (semua informasi yang diperlukan untuk memahami fenomena yang disajikan harus dimasukkan), netralitas (bebas dari bias dalam pemilihan atau penyajian), dan bebas dari kesalahan (tidak ada kesalahan dalam mendeskripsikan fenomena — meski estimasi terbaik pun dapat mengandung ketidakpastian inherent). Kelengkapan merupakan tambahan eksplisit yang tidak bernama dalam SFAC No. 2."
    )

    # C.
    add_subsection_heading(doc,
                           "Karakteristik Penambah dan Batasan Biaya",
                           "C.")
    add_body_para(doc,
        "SFAC No. 8 memperkenalkan kategori karakteristik kualitatif penambah (enhancing qualitative characteristics): komparabilitas, verifikabilitas, ketepatan waktu, dan keterpahaman (understandability). Kategori ini bukan sekadar “sekunder” — ia secara konseptual berbeda: karakteristik penambah meningkatkan kegunaan informasi yang sudah memiliki relevansi dan representasi tepat, tetapi tidak dapat menebus ketiadaan karakteristik fundamental. Implikasi praktisnya signifikan: verifikabilitas tidak lagi dapat digunakan sebagai proksi kecukupan kualitas — seorang auditor tidak dapat menyimpulkan bahwa informasi yang kurang terverifikasi namun merepresentasikan fenomena ekonomi secara tepat harus dikecualikan semata-mata karena ia lemah dalam karakteristik penambah."
    )
    add_body_para(doc,
        "Verifikabilitas dan ketepatan waktu — yang sebelumnya tertanam dalam keandalan dan relevansi SFAC No. 2 — kini berdiri mandiri sebagai karakteristik penambah. Keterpahaman (understandability), yang dalam SFAC No. 2 diperlakukan sebagai kualitas yang bergantung pada pengguna, kini menjadi karakteristik eksplisit dengan asumsi bahwa pengguna memiliki pengetahuan keuangan yang memadai."
    )
    add_body_para(doc,
        "Paling penting, SFAC No. 8 secara eksplisit mengecualikan konservatisme dari kerangka kualitas: “including prudence or conservatism in the description of faithful representation would be incompatible with neutrality” (SFAC No. 8, BC3.27). Pengecualian ini — sebagaimana akan dianalisis pada Bagian IV — merupakan salah satu keputusan paling kontroversial dalam sejarah standar penetapan FASB."
    )

    # -----------------------------------------------------------------------
    # SECTION IV — Perbandingan dan Analisis Kritis
    # -----------------------------------------------------------------------
    add_section_heading(doc,
                        "Perbandingan dan Analisis Kritis",
                        "IV.")

    # A.
    add_subsection_heading(doc,
                           "Pergeseran Terminologis: Keandalan → Representasi Tepat",
                           "A.")
    add_body_para(doc,
        "Penggantian istilah “keandalan” dengan “representasi tepat” merupakan perubahan yang tampaknya semantis namun secara teoritis bersifat substantif. Kritik terhadap “keandalan” telah lama mengendap dalam literatur: Schipper dan Vincent (2003) mendokumentasikan bagaimana akuntan dan regulator kerap mengidentikkan keandalan dengan ketepatan pengukuran atau verifikabilitas — bukan dengan korespondensi antara angka akuntansi dan fenomena ekonomi yang diklaim untuk diwakilinya. Akibatnya, aset yang diukur dengan biaya historis cenderung dianggap “lebih andal” ketimbang estimasi nilai wajar, semata-mata karena biaya historis lebih mudah diverifikasi — padahal verifikabilitas dan keandalan representasional adalah dua konstruk yang berbeda."
    )
    add_body_para(doc,
        "“Representasi tepat” (SFAC No. 8) memutus asosiasi keliru tersebut: tujuannya adalah bahwa angka akuntansi secara tepat merepresentasikan fenomena ekonomi yang diklaim diwakilinya — bukan bahwa angka tersebut mudah diverifikasi atau presisi secara pengukuran. Pergeseran ini membuka ruang teoretis bagi pengukuran nilai wajar: estimasi nilai wajar Level 3 (berdasarkan model) dapat mencapai representasi tepat meski tidak mudah diverifikasi oleh pihak independen. Dengan demikian, SFAC No. 8 secara implisit mendukung orientasi fair value yang menjadi tren dominan standar akuntansi internasional pasca-2005."
    )

    # B.
    add_subsection_heading(doc,
                           "Restrukturisasi Hierarki Karakteristik",
                           "B.")
    add_body_para(doc,
        "Pergeseran dari hierarki primer/sekunder (SFAC No. 2) ke fundamental/penambah (SFAC No. 8) membawa sejumlah perpindahan signifikan yang secara menyeluruh dianalisis oleh Wolk et al. (2019). Pertama, ketepatan waktu berpindah dari komponen relevansi (SFAC No. 2) menjadi karakteristik penambah (SFAC No. 8). Alasan yang diajukan: informasi dapat relevan secara substansial namun tetap terlambat tersedia — keduanya adalah dimensi berbeda yang tidak saling mendefinisikan. Kedua, verifikabilitas berpindah dari komponen keandalan (SFAC No. 2) menjadi karakteristik penambah (SFAC No. 8). Logikanya serupa: representasi tepat dapat tercapai bahkan jika pihak independen tidak dapat mereplikasi pengukurannya secara tepat."
    )
    add_body_para(doc,
        "Perpindahan yang paling konsekuensial adalah materialitas: dari ambang batas terpisah (SFAC No. 2) menjadi aspek dari relevansi (SFAC No. 8). Dalam SFAC No. 2, materialitas berdiri sebagai pertanyaan yang terpisah dari relevansi — informasi mungkin saja relevan namun tidak material. SFAC No. 8 mengintegrasikannya: informasi yang tidak material secara definitif tidak relevan bagi pengambil keputusan, sehingga tidak memerlukan pengakuan atau pengungkapan. Integrasi ini memperketat konsistensi logis kerangka, meski dalam praktiknya garis batas materialitas tetap merupakan pertimbangan kualitatif yang bergantung pada konteks."
    )

    # C.
    add_subsection_heading(doc,
                           "Eliminasi Konservatisme dan Implikasinya",
                           "C.")
    add_body_para(doc,
        "Salah satu keputusan paling kontroversial SFAC No. 8 adalah pengecualian eksplisit konservatisme dari kerangka kualitas. Argumen FASB bersandar pada netralitas: jika representasi tepat mensyaratkan informasi bebas dari bias, maka konservatisme — yang secara sistematis mengutamakan nilai yang lebih rendah di bawah kondisi ketidakpastian — menciptakan bias asimetris yang tidak dapat dibenarkan secara konseptual (SFAC No. 8, BC3.27)."
    )
    add_body_para(doc,
        "Namun argumen ini menghadapi serangan balik yang kuat dari perspektif teori kontrak. Watts (2003) berargumen bahwa konservatisme bukan sekadar konvensi usang, melainkan respons yang rasional terhadap asimetri verifikasi dalam kontrak. Perjanjian utang (debt covenants), kompensasi berbasis kinerja, dan risiko litigasi asimetris semuanya menciptakan permintaan ekonomi atas akuntansi konservatif yang tidak bergantung pada preferensi badan standar. Konservatisme, dalam pandangan Watts, adalah fitur, bukan bug. Dengan menghapusnya dari kerangka konseptual, SFAC No. 8 menghasilkan ketegangan yang belum terselesaikan: standar-standar seperti pengujian penurunan nilai (impairment testing) dan pengakuan kerugian ekspektasian (expected credit loss) tetap mengandung asimetri perlakuan yang mencerminkan konservatisme de facto — bahkan setelah konservatisme dihapus de jure dari kerangka."
    )

    # D.
    add_subsection_heading(doc,
                           "Dinamika Politik Konvergensi FASB-IASB",
                           "D.")
    add_body_para(doc,
        "SFAC No. 8 merupakan produk intelektual tertinggi proyek konvergensi FASB-IASB — dokumen yang secara substansial identik dengan kerangka konseptual IASB. Namun ia juga menandai batas konvergensi tersebut. Perpecahan struktural antara kedua tradisi regulasi tetap tidak terselesaikan: GAAP Amerika Serikat berakar pada pendekatan rules-based yang memberikan panduan terperinci dan garis batas yang jelas — sebuah preferensi yang secara pragmatis melindungi auditor dan preparer dari risiko litigasi. IFRS, sebaliknya, mengandalkan pendekatan principles-based yang mengutamakan substansi ekonomi di atas bentuk hukum dan menuntut lebih banyak pertimbangan profesional."
    )
    add_body_para(doc,
        "Namun akar kegagalan konvergensi bersifat lebih struktural dari sekadar perbedaan teknis atau tekanan kelompok kepentingan. FASB adalah badan penetapan standar nasional yang dibentuk oleh undang-undang dan bertanggung jawab kepada SEC — sebuah lembaga pemerintah Amerika Serikat. IASB, sebaliknya, adalah entitas swasta yang bertanggung jawab kepada konstituen global tanpa mekanisme penegakan yang mengikat. Perbedaan mandat konstitusional ini berarti bahwa bahkan ketika kedua badan menyepakati prinsip konseptual yang sama, penerjemahannya ke dalam aturan yang dapat ditegakkan niscaya berbeda: konteks institusional yang memberi makna pada aturan tersebut tidak dapat disamakan hanya dengan harmonisasi dokumen kerangka konseptual."
    )
    add_body_para(doc,
        "Tekanan politik yang mengikutinya adalah manifestasi dari perbedaan struktural ini, bukan penyebab independennya. Kongres Amerika Serikat, SEC, dan kelompok industri secara konsisten menolak adopsi penuh IFRS di Amerika, khawatir kehilangan kendali regulasi atas standar akuntansi nasional. Di sisi lain, IASB menghadapi tekanan dari Uni Eropa dan negara-negara Asia yang berupaya mempertahankan kekhasan tradisi akuntansi lokal mereka. Akibatnya, meskipun SFAC No. 8 dan kerangka IASB berbagi arsitektur yang hampir identik, implementasi standar-standar turunannya — seperti pengakuan pendapatan, sewa, dan instrumen keuangan — tetap mengandung divergensi yang signifikan. Konvergensi pada tingkat kerangka konseptual tidak serta-merta menghasilkan konvergensi pada tingkat standar operasional, justru karena standar operasional harus bekerja dalam sistem hukum dan regulasi yang berbeda secara mendasar."
    )

    # -----------------------------------------------------------------------
    # SECTION V — Kesimpulan
    # -----------------------------------------------------------------------
    add_section_heading(doc, "Kesimpulan", "V.")

    add_body_para(doc,
        "Perbandingan antara SFAC No. 2 dan SFAC No. 8 mengungkapkan bahwa transisi kerangka kualitas informasi akuntansi FASB bukan sekadar pembaruan nomenklatur, melainkan pergeseran orientasi yang mencerminkan dua transformasi besar: pertama, menuju kerangka yang lebih konsisten secara internal — di mana setiap karakteristik diposisikan berdasarkan perannya yang sesungguhnya (fundamental vs. penambah), bukan berdasarkan hierarki kepentingan yang kadang tumpang tindih; dan kedua, menuju kerangka yang secara eksplisit mendukung pengukuran berbasis nilai ekonomi alih-alih berbasis biaya historis."
    )
    add_body_para(doc,
        "Namun kemajuan ini tidak bebas tegangan. Penghapusan konservatisme dari kerangka konseptual menciptakan ketidaksesuaian yang belum terpecahkan: standar-standar operasional yang lahir dari kerangka yang sama tetap mengandung asimetri perlakuan yang mencerminkan konservatisme de facto. Ini menunjukkan bahwa konservatisme bukan semata-mata pilihan normatif, melainkan respons terhadap struktur insentif yang lebih dalam — sebagaimana diteorikan Watts (2003). Kerangka konseptual yang lebih baik bukan yang menghapus ketegangan ini, melainkan yang mengakuinya dan memberi panduan untuk menavigasinya. Kasus pengakuan kerugian ekspektasian (expected credit loss) dan pengujian penurunan nilai — di mana asimetri perlakuan konservatif dipertahankan dalam standar yang lahir dari kerangka yang menolak konservatisme — adalah bukti bahwa ketegangan tersebut bukan anomali historis, melainkan pertanyaan yang terus menanti jawaban konseptual yang lebih jujur."
    )

    # -----------------------------------------------------------------------
    # DAFTAR PUSTAKA
    # -----------------------------------------------------------------------
    add_section_heading(doc, "Daftar Pustaka", "")

    # Entry 1
    add_daftar_pustaka_entry(doc, [
        ("FASB. (1980). ", False),
        ("Statement of Financial Accounting Concepts No. 2: Qualitative Characteristics of Accounting Information", True),
        (". Financial Accounting Standards Board.", False),
    ])

    # Entry 2
    add_daftar_pustaka_entry(doc, [
        ("FASB. (2010). ", False),
        ("Statement of Financial Accounting Concepts No. 8: Conceptual Framework for Financial Reporting", True),
        (". Financial Accounting Standards Board.", False),
    ])

    # Entry 3
    add_daftar_pustaka_entry(doc, [
        ("IASB. (2018). ", False),
        ("Conceptual Framework for Financial Reporting", True),
        (". IFRS Foundation.", False),
    ])

    # Entry 4
    add_daftar_pustaka_entry(doc, [
        ("Schipper, K., & Vincent, L. (2003). Earnings quality. ", False),
        ("Accounting Horizons", True),
        (", ", False),
        ("17", True),
        ("(Supplement), 97–110.", False),
    ])

    # Entry 5
    add_daftar_pustaka_entry(doc, [
        ("Watts, R. L. (2003). Conservatism in accounting part I: Explanations and implications. ", False),
        ("Accounting Horizons", True),
        (", ", False),
        ("17", True),
        ("(3), 207–221.", False),
    ])

    # Entry 6
    add_daftar_pustaka_entry(doc, [
        ("Wolk, H. I., Dodd, J. L., & Rozycki, J. J. (2019). ", False),
        ("Accounting Theory: Conceptual Issues in a Political and Economic Environment", True),
        (" (9th ed.). SAGE Publications.", False),
    ])

    # -----------------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------------
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir,
                            "01079_Dzaki Muhammad Yusfian_Tugas Individual 3.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    print(f"Size:  {os.path.getsize(out_path):,} bytes")


if __name__ == "__main__":
    build()
