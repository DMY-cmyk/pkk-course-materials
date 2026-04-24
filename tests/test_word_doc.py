import pytest


def test_all_modules_import():
    from scripts.word_doc import styles, visuals, cover, front_matter, bibliography
    from scripts.word_doc.sections import (
        bagian_i, bagian_ii, bagian_iii, bagian_iv,
        bagian_v, bagian_vi, bagian_vii, bagian_viii,
    )


from docx import Document


def test_set_run_font_sets_times_new_roman():
    from scripts.word_doc.styles import set_run_font
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("test")
    set_run_font(run, size_pt=12)
    assert run.font.name == "Times New Roman"
    assert run.font.size.pt == 12
    assert run.font.bold is None or run.font.bold is False


def test_add_section_heading_is_14pt_bold():
    from scripts.word_doc.styles import add_section_heading
    doc = Document()
    p = add_section_heading(doc, "I. TEST HEADING")
    run = p.runs[0]
    assert run.font.bold is True
    assert run.font.size.pt == 14


def test_add_body_paragraph_is_12pt():
    from scripts.word_doc.styles import add_body_paragraph
    doc = Document()
    p = add_body_paragraph(doc, "Some text here.")
    run = p.runs[0]
    assert run.font.size.pt == 12
    assert run.font.bold is None or run.font.bold is False


def test_add_table_with_headers_shape():
    from scripts.word_doc.styles import add_table_with_headers
    doc = Document()
    headers = ["Kolom A", "Kolom B", "Kolom C"]
    rows = [["a1", "b1", "c1"], ["a2", "b2", "c2"]]
    table = add_table_with_headers(doc, headers, rows)
    assert len(table.rows) == 3   # 1 header + 2 data
    assert len(table.columns) == 3


PNG_MAGIC = b"\x89PNG\r\n\x1a\n"


def _is_valid_png(buf):
    return buf.read(8) == PNG_MAGIC


def test_grafik_eps_is_png():
    from scripts.word_doc.visuals import generate_grafik_eps_indf
    buf = generate_grafik_eps_indf()
    assert _is_valid_png(buf)


def test_bagan_timeline_is_png():
    from scripts.word_doc.visuals import generate_bagan_timeline
    buf = generate_bagan_timeline()
    assert _is_valid_png(buf)


def test_all_visuals_return_bytesio():
    from scripts.word_doc import visuals
    funcs = [
        visuals.generate_bagan_timeline,
        visuals.generate_bagan_primary_users,
        visuals.generate_bagan_qc_sfac2,
        visuals.generate_bagan_qc_sfac8,
        visuals.generate_grafik_eps_indf,
        visuals.generate_bagan_ci_waterfall,
        visuals.generate_bagan_entry_exit_tree,
        visuals.generate_bagan_fs_hierarchy,
        visuals.generate_bagan_fasb_to_psak,
        visuals.generate_bagan_indf_org,
        visuals.generate_bagan_kepemilikan_pie,
        visuals.generate_bagan_matrix_relevance_fr,
    ]
    for fn in funcs:
        buf = fn()
        assert hasattr(buf, "read"), f"{fn.__name__} must return BytesIO"
        data = buf.read()
        assert data[:8] == PNG_MAGIC, f"{fn.__name__} must return PNG"


import os
BASE_DIR = os.path.join(
    os.path.dirname(__file__), ".."
)


def test_cover_adds_all_six_members():
    from docx import Document
    from scripts.word_doc.cover import build_cover
    doc = Document()
    logo_path = os.path.join(BASE_DIR, "STIE Logo.png")
    build_cover(doc, logo_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for name in [
        "Efri Nurmalinda", "Dzaki Muhammad Yusfian",
        "Nuradila", "Achmad Dimas Wibawa",
        "Adinda Putri Dewi", "Setiabudi Yudha Pratama",
    ]:
        assert name in full_text, f"Missing member: {name}"


def test_cover_contains_title_keywords():
    from docx import Document
    from scripts.word_doc.cover import build_cover
    doc = Document()
    logo_path = os.path.join(BASE_DIR, "STIE Logo.png")
    build_cover(doc, logo_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "KERANGKA KONSEPTUAL FASB" in full_text
    assert "MNK202" in full_text


def test_front_matter_contains_kata_pengantar():
    from docx import Document
    from scripts.word_doc.front_matter import build_front_matter
    doc = Document()
    build_front_matter(doc)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Kata Pengantar" in full_text
    assert "DAFTAR ISI" in full_text


def test_bagian_i_contains_key_content():
    from docx import Document
    from scripts.word_doc.sections.bagian_i import build
    doc = Document()
    build(doc)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "SFAC 1" in text
    assert "Norwalk Agreement" in text
    assert "PSAK" in text
    # Tabel 1.1 and Tabel 1.2
    assert len(doc.tables) >= 2


def test_bagian_ii_contains_key_content():
    from docx import Document
    from scripts.word_doc.sections.bagian_ii import build
    doc = Document()
    build(doc)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Faithful Representation" in text
    assert "Reliability" in text
    assert "BC3.27" in text
    assert len(doc.tables) >= 1
    assert len(doc.inline_shapes) >= 3   # 2 trees + 1 bar chart


def test_bagian_iii_contains_key_content():
    from docx import Document
    from scripts.word_doc.sections.bagian_iii import build
    doc = Document()
    build(doc)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "present right" in text
    assert "BC4.7" in text
    assert "Comprehensive Income" in text
    assert len(doc.tables) >= 2


def test_bagian_iv_contains_key_content():
    from docx import Document
    from scripts.word_doc.sections.bagian_iv import build
    doc = Document()
    build(doc)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "RD3" in text
    assert "M30" in text
    assert "Entry Price" in text
    assert "PR12" in text
    assert len(doc.tables) >= 2


def test_bagian_v_contains_key_content():
    from docx import Document
    from scripts.word_doc.sections.bagian_v import build
    doc = Document()
    build(doc)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Achilles" in text
    assert "pro-cyclicality" in text or "pro-siklikalitas" in text


def test_bagian_vi_contains_key_content():
    from docx import Document
    from scripts.word_doc.sections.bagian_vi import build
    doc = Document()
    build(doc)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Chapter 8" in text or "Catatan" in text
    assert len(doc.tables) >= 1


def test_bagian_vii_contains_key_content():
    from docx import Document
    from scripts.word_doc.sections.bagian_vii import build
    doc = Document()
    build(doc)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "115,79" in text or "115.79" in text   # Net Sales
    assert "52,2" in text or "52.2" in text        # Goodwill
    assert "First Pacific" in text
    assert "43,077" in text or "43.077" in text    # NCI
    assert len(doc.tables) >= 4
    assert len(doc.inline_shapes) >= 3             # org, pie, matrix


def test_bagian_viii_contains_conclusion():
    from docx import Document
    from scripts.word_doc.sections.bagian_viii import build
    doc = Document()
    build(doc)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Kesimpulan" in text or "KESIMPULAN" in text
    assert "DSAK" in text or "OJK" in text


def test_bibliography_parses_txt_file():
    import os
    from docx import Document
    from scripts.word_doc.bibliography import build_bibliography
    txt_path = os.path.join(
        os.path.dirname(__file__), "..",
        "Daftar Pustaka — Presentasi Kelompok 3.txt"
    )
    doc = Document()
    build_bibliography(doc, txt_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Daftar Pustaka" in text or "DAFTAR PUSTAKA" in text
    # At least some reference entries should be present
    assert len([p for p in doc.paragraphs if p.text.strip()]) > 5


import tempfile


def test_generate_full_document_produces_docx():
    import subprocess, sys, os
    output_dir = tempfile.mkdtemp()
    output_path = os.path.join(output_dir, "test_output.docx")
    result = subprocess.run(
        [sys.executable, "scripts/generate_word_doc.py", "--output", output_path],
        capture_output=True, text=True, timeout=120
    )
    assert result.returncode == 0, f"Script failed:\n{result.stderr}"
    assert os.path.exists(output_path), "Output .docx not created"
    from docx import Document
    doc = Document(output_path)
    assert len(doc.paragraphs) > 50, "Document too short"
