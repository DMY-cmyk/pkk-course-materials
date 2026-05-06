"""
build_docx.py
Builds 01079_Dzaki Muhammad Yusfian_Tugas Individual 4.docx
using python-docx 1.2.0.
Adapted from ti3/output/build_docx.py — same helper functions, new content.
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers (identical to ti3/output/build_docx.py)
# ---------------------------------------------------------------------------

def add_paragraph(doc, text, font_name="Times New Roman", font_size=12,
                  bold=False, italic=False,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  line_spacing_pt=18, space_before_pt=0, space_after_pt=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)
    return para


def add_section_heading(doc, text, number_prefix, font_size=13):
    full_text = f"{number_prefix}   {text}" if number_prefix else text
    return add_paragraph(doc, full_text, font_size=font_size, bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         line_spacing_pt=18, space_before_pt=12, space_after_pt=6)


def add_subsection_heading(doc, text, letter_prefix):
    full_text = f"{letter_prefix}   {text}"
    return add_paragraph(doc, full_text, font_size=12, bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         line_spacing_pt=18, space_before_pt=12, space_after_pt=6)


def add_body_para(doc, text):
    return add_paragraph(doc, text, font_size=12, bold=False,
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         line_spacing_pt=18, space_before_pt=0, space_after_pt=6)


def add_blank(doc):
    return add_paragraph(doc, "", line_spacing_pt=18,
                         space_before_pt=0, space_after_pt=0)


def add_cover_para(doc, text, font_size=12, bold=False):
    return add_paragraph(doc, text, font_size=font_size, bold=bold,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing_pt=18, space_before_pt=0, space_after_pt=6)


def add_daftar_pustaka_entry(doc, text_parts):
    """text_parts: list of (text, italic) tuples. Hanging indent."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.left_indent = Cm(0.75)
    pf.first_line_indent = Cm(-0.75)
    for text, italic in text_parts:
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.italic = italic
        run.bold = False
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), "Times New Roman")
        rFonts.set(qn('w:hAnsi'), "Times New Roman")
        rFonts.set(qn('w:cs'), "Times New Roman")
        existing = rPr.find(qn('w:rFonts'))
        if existing is not None:
            rPr.remove(existing)
        rPr.insert(0, rFonts)
    return para


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # Page setup: Letter 8.5x11 in, 3cm all margins
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # ------------------------------------------------------------------
    # COVER
    # ------------------------------------------------------------------
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)
    add_cover_para(doc, "TUGAS INDIVIDU 4", font_size=16, bold=True)
    add_cover_para(doc, "Pelaporan Keuangan Korporat", font_size=14, bold=False)
    add_cover_para(doc,
        "Kegunaan Informasi Laporan Keuangan bagi Investor dan Kreditur",
        font_size=14, bold=True)
    add_blank(doc)
    add_blank(doc)
    add_cover_para(doc, "Disusun oleh:", font_size=12, bold=False)
    add_cover_para(doc, "Dzaki Muhammad Yusfian", font_size=12, bold=True)
    add_cover_para(doc, "1225 01079", font_size=12, bold=False)
    add_blank(doc)
    add_blank(doc)
    add_cover_para(doc, "Dosen Pengampu: Prof. Djoko Susanto", font_size=12, bold=False)
    add_cover_para(doc, "Program Studi Magister Akuntansi", font_size=12, bold=False)
    add_cover_para(doc, "Mei 2026", font_size=12, bold=True)

    doc.add_page_break()

    # ------------------------------------------------------------------
    # SECTION I — Pendahuluan
    # ------------------------------------------------------------------
    add_section_heading(doc, "Pendahuluan", "I.")

    add_body_para(doc,
        'Kerangka Konseptual IASB (2018) menetapkan tujuan pelaporan keuangan dalam satu kalimat yang deceptively sederhana: "The objective of general purpose financial reporting is to provide financial information about the reporting entity that is useful to existing and potential investors, lenders and other creditors in making decisions relating to providing resources to the entity" (§1.2). Kata kunci dalam pernyataan ini bukan "informasi" atau "keuangan", melainkan "berguna" — sebuah kriteria yang mengandaikan pemakai tertentu, keputusan tertentu, dan konteks tertentu.')
    add_body_para(doc,
        'Pertanyaan "bagaimana informasi laporan keuangan berguna bagi investor dan kreditur?" tidak dapat dijawab dengan mendeskripsikan isi laporan keuangan. Pertanyaan itu menuntut analisis yang dimulai dari fungsi utilitas masing-masing pemakai: investor ekuitas sebagai residual claimant yang berfokus pada upside, dan kreditur sebagai fixed claimant yang berfokus pada downside protection. Dua kelompok ini membaca laporan yang sama, tetapi mengekstrak informasi yang berbeda, melalui lensa yang berbeda, untuk keputusan yang berbeda.')
    add_body_para(doc,
        'Esai ini mengorganisasi analisis sebagai berikut: Bagian II membangun kerangka decision usefulness dan bukti empiris bahwa laporan keuangan benar-benar digunakan pasar. Bagian III menganalisis kegunaan informasi bagi investor — meliputi karakteristik kualitatif yang relevan, bukti empiris (Ball & Brown 1968; model Ohlson 1995), dan aplikasi pada PT Bank Central Asia Tbk (BBCA). Bagian IV melakukan analisis paralel bagi kreditur. Bagian V menutup dengan sintesis.')

    # ------------------------------------------------------------------
    # SECTION II — Kerangka Decision Usefulness
    # ------------------------------------------------------------------
    add_section_heading(doc, "Kerangka Decision Usefulness", "II.")

    add_body_para(doc,
        'Wolk, Dodd dan Rozycki (2017) mendefinisikan decision usefulness sebagai tujuan tunggal pelaporan keuangan bertujuan umum: menyediakan informasi yang memungkinkan penyedia modal — investor, pemberi pinjaman, dan kreditur lain — membuat keputusan alokasi sumber daya secara lebih rasional. Proposisi ini bukan sekadar pernyataan normatif tentang apa yang "seharusnya" dilakukan laporan keuangan; ia adalah standar evaluasi: setiap PSAK, setiap pengungkapan, setiap kebijakan akuntansi harus diuji dengan pertanyaan "apakah ini membuat laporan lebih berguna bagi pemakai primer?"')
    add_body_para(doc,
        'Kegunaan tidak identik dengan akurasi. Informasi yang akurat secara pengukuran tetapi terlambat tersedia — setelah jendela keputusan tertutup — gagal memenuhi kriteria usefulness meskipun secara teknis benar. Inilah mengapa Kerangka Konseptual IASB (2018) menempatkan ketepatan waktu (timeliness) sebagai karakteristik enhancing, bukan sekadar fitur tambahan. Dua karakteristik fundamental — relevansi (predictive value + confirmatory value + materiality) dan representasi tepat (complete, neutral, free from error) — keduanya harus terpenuhi; empat karakteristik enhancing (comparability, verifiability, timeliness, understandability) meningkatkan kegunaan yang sudah memenuhi fondasi fundamental. Satu constraint membatasi seluruh hierarki: manfaat informasi harus melebihi biaya penyediaannya.')
    add_body_para(doc,
        'Bahwa laporan keuangan benar-benar digunakan pasar — bukan sekadar diterbitkan dan diabaikan — dibuktikan secara empiris oleh Ball dan Brown (1968) dalam studi perintis mereka. Menggunakan sampel 261 perusahaan NYSE (1957–1965), mereka mengklasifikasikan pengumuman laba sebagai Good News (laba di atas ekspektasi) atau Bad News (di bawah ekspektasi) dan mengukur Cumulative Average Abnormal Return (CAAR) dalam jendela ±12 bulan. Temuan: perusahaan GN menghasilkan CAAR positif, perusahaan BN menghasilkan CAAR negatif. Meski 80–90% informasi sudah diantisipasi pasar sebelum pengumuman resmi — melalui laporan kuartalan, guidance manajemen, dan analisis analis — reaksi residual yang signifikan pada bulan pengumuman membuktikan bahwa laporan keuangan memiliki information content yang independen dan tidak dapat digantikan sepenuhnya oleh sumber lain. Perlu dicatat bahwa laba tidak terduga hanya menjelaskan sekitar 5% dari variasi total harga sekuritas (Scott, 2015) — bukan karena laporan keuangan kurang berguna, melainkan karena harga saham merespons banyak faktor makroekonomi, sentimen, dan informasi non-akuntansi secara simultan. Temuan ini justru memotivasi analisis yang lebih mendalam: investor yang mampu mengekstrak sinyal dari komponen laporan keuangan di luar headline earnings berpotensi memperoleh keunggulan informasional yang bermakna.')

    # ------------------------------------------------------------------
    # SECTION III — Investor
    # ------------------------------------------------------------------
    add_section_heading(doc, "Kegunaan Informasi bagi Investor", "III.")

    add_subsection_heading(doc, "Karakteristik yang Relevan bagi Investor", "A.")
    add_body_para(doc,
        'Investor ekuitas adalah residual claimant: mereka menerima apa yang tersisa setelah seluruh kewajiban perusahaan kepada pihak lain terpenuhi. Konsekuensinya, fungsi utilitas investor berorientasi ke atas (upside-oriented) — mereka berminat pada potensi pertumbuhan laba, arus kas bebas, kebijakan dividen, dan kualitas tata kelola yang menentukan seberapa besar "sisa" yang akan mereka terima di masa depan.')
    add_body_para(doc,
        'Dari hierarki karakteristik kualitatif IASB, dua karakteristik paling kritis bagi investor adalah predictive value dan timeliness. Predictive value — sub-komponen relevansi — memungkinkan investor menggunakan angka akuntansi historis untuk memperkirakan laba, arus kas, dan nilai intrinsik masa depan. Tanpa predictive value, laporan keuangan hanya mendokumentasikan masa lalu tanpa memberikan landasan untuk keputusan yang berorientasi masa depan. Timeliness memastikan bahwa informasi tersedia sebelum jendela keputusan investasi tertutup; laporan yang akurat tetapi terlambat tidak mencegah investor dari keputusan yang salah arah.')
    add_body_para(doc,
        'Confirmatory value melengkapi predictive value: ketika laba aktual BBCA 2025 sebesar Rp 57,5 triliun dibandingkan dengan target manajemen ROE 21%–23%, investor dapat memperbarui model valuasi mereka secara Bayesian — mengkonfirmasi atau merevisi ekspektasi tentang kemampuan menghasilkan laba BBCA di masa depan. Komparabilitas (data 5 tahun, 2021–2025) menambah dimensi ketiga: analisis time-series yang mengidentifikasi tren, bukan sekadar titik data tunggal.')

    add_subsection_heading(doc, "Bukti Empiris: Ball & Brown, ERC, dan Model Ohlson", "B.")
    add_body_para(doc,
        'Bukti empiris paling langsung tentang kegunaan laporan keuangan bagi investor datang dari penelitian pasar modal. Ball dan Brown (1968) membuktikan bahwa pengumuman laba menghasilkan abnormal return yang signifikan — implikasinya bagi investor bukan sekadar bahwa "laporan berguna", melainkan bahwa angka laba merupakan sinyal yang diproses oleh pasar dan diterjemahkan ke dalam harga aset. Sensitivitas harga terhadap laba kejutan ini diukur oleh Earnings Response Coefficient (ERC): CARⱼj = α + β × UEⱼj + εⱼj, di mana CAR adalah cumulative abnormal return, UE adalah unexpected earnings (realisasi dikurangi ekspektasi), dan β adalah ERC. ERC tinggi ditemukan pada perusahaan dengan: (a) laba yang persisten, (b) beta pasar yang rendah, (c) leverage yang rendah, dan (d) peluang pertumbuhan yang tinggi (Collins & Kothari, 1989).')
    add_body_para(doc,
        'Melengkapi temuan Ball dan Brown, penelitian Lintner (1956) tentang kebijakan dividen mengungkap mekanisme transmisi penting bagi investor. Lintner menemukan bahwa manajer menetapkan target payout ratio jangka panjang dan melakukan penyesuaian bertahap menuju target tersebut — dengan konsekuensi kritis: manajer hanya menaikkan dividen ketika mereka meyakini bahwa peningkatan laba bersifat permanen, bukan transitory. Kebijakan dividen dengan demikian berfungsi sebagai sinyal kualitas laba kepada pasar: kenaikan dividen merupakan pernyataan implisit manajemen tentang keyakinan mereka atas keberlanjutan laba tinggi. Rantai ini — laba akuntansi → sinyal dividen → ekspektasi permanensi → harga saham — menegaskan mengapa investor membutuhkan informasi yang dapat membedakan komponen laba persisten dari komponen transitory. Dalam konteks model valuasi, laba yang persisten menghasilkan earnings response coefficient (ERC) yang lebih tinggi dan parameter ω mendekati 1 dalam model Ohlson. Bernard dan Thomas (1989) menemukan bahwa pasar sering gagal memproses perbedaan ini secara instan: harga saham cenderung terus bergerak ke arah yang sama dengan kejutan laba selama beberapa bulan pasca pengumuman — post-earnings-announcement drift — sebuah anomali yang menciptakan peluang bagi investor yang melakukan analisis laporan keuangan secara lebih mendalam melampaui angka laba utama.')
    add_body_para(doc,
        'Model Ohlson (1995) menyatukan pendekatan balance sheet dan income statement dalam satu kerangka koheren. Nilai intrinsik ekuitas diekspresikan sebagai: Vₜ = BVₜ + Σ Eₜ[xᵃₜ₊τ] / (1+r)ᵗ, di mana xᵃₜ = earningsₜ − r × BVₜ₋₁ adalah laba abnormal — selisih antara laba aktual dan required return atas nilai buku. Syarat clean surplus mengharuskan semua perubahan ekuitas (selain transaksi dengan pemilik) melewati laporan laba rugi: BVₜ = BVₜ₋₁ + earningsₜ − dividendsₜ. Parameter ω ∈ [0,1) mengukur persistensi laba abnormal: nilai ω mendekati 1 berarti keunggulan kompetitif perusahaan bertahan lama; mendekati 0 berarti laba abnormal bersifat transitory. Model Ohlson menjadikan laporan keuangan — khususnya nilai buku dan laba — sebagai input langsung dalam valuasi ekuitas.')

    add_subsection_heading(doc, "Aplikasi: BBCA sebagai Kasus Pasar Modal Indonesia", "C.")
    add_body_para(doc,
        'PT Bank Central Asia Tbk menyediakan laboratorium ideal untuk menguji decision usefulness bagi investor. Per 31 Desember 2025, BBCA membukukan EPS Rp 467, ROE 23,3%, P/BV 3,8x, dan P/E 17,3x (BBCA, 2025, hal. 16) — metrik yang menjadi input langsung dalam model valuasi investor institusional dan ritel.')
    add_body_para(doc,
        'Aplikasi model Ohlson pada BBCA menghasilkan kalkulasi yang instruktif. Dengan nilai buku per saham (BVPS) Rp 2.288 dan asumsi biaya ekuitas 12%, laba abnormal per saham tahun pertama adalah: xᵃ₁ = (23,3% − 12%) × Rp 2.288 = 11,3% × Rp 2.288 ≈ Rp 258 per saham (BBCA, 2025, hal. 16). Dengan cakrawala 5 tahun dan asumsi ROE tetap 23,3% serta retention ratio 40%, nilai intrinsik yang dihasilkan model adalah sekitar Rp 3.600–3.800 per saham. Harga pasar aktual Rp 8.075 per saham mencerminkan gap yang signifikan — sinyal bahwa pasar memperkirakan laba abnormal BBCA akan persisten jauh melampaui 5 tahun. Ekspektasi persistensi tinggi ini konsisten secara fundamental: CASA ratio 83,7% (BBCA, 2025, hal. 200) menciptakan biaya pendanaan yang sulit ditiru kompetitor digital — sebuah economic moat yang mendukung ω mendekati 1.')
    add_body_para(doc,
        'Dari sisi standar, transisi dari PSAK 55 (incurred loss model) ke PSAK 71 (expected credit loss) meningkatkan predictive value laba bank bagi investor. Model ECL 3-stage BBCA — dengan CKPN kredit Rp 31,6 triliun dan beban CKPN 2025 Rp 4,0 triliun (BBCA, 2025, hal. 201, 210) — memberikan informasi forward-looking tentang risiko kredit yang tidak tersedia di bawah PSAK 55. Ini memungkinkan investor untuk memproyeksikan cost of credit masa depan — komponen penting dalam pemodelan laba bank — dengan akurasi yang lebih tinggi.')

    # ------------------------------------------------------------------
    # SECTION IV — Kreditur
    # ------------------------------------------------------------------
    add_section_heading(doc, "Kegunaan Informasi bagi Kreditur", "IV.")

    add_subsection_heading(doc, "Karakteristik yang Relevan bagi Kreditur", "A.")
    add_body_para(doc,
        'Kreditur — baik pemegang obligasi, bank pemberi pinjaman antarbank, maupun pemasok dengan piutang dagang — adalah fixed claimant: mereka menerima pembayaran tetap berupa bunga dan pokok, tanpa hak atas kelebihan nilai di atas klaim tersebut. Konsekuensinya, fungsi utilitas kreditur berorientasi ke bawah (downside-oriented): pertanyaan mereka bukan "seberapa tinggi laba perusahaan bisa tumbuh?" melainkan "apakah perusahaan ini mampu tidak gagal bayar, bahkan dalam skenario makroekonomi yang buruk?"')
    add_body_para(doc,
        'Dari hierarki karakteristik kualitatif, kreditur mengutamakan verifikabilitas dan kelengkapan (completeness). Verifikabilitas — kemampuan pengamat independen mencapai kesimpulan serupa — memberikan kredibilitas pada estimasi manajemen yang menjadi dasar penilaian risiko kredit. Kelengkapan memastikan tidak ada risiko material yang tersembunyi dalam laporan. Representasi tepat (faithful representation) secara keseluruhan lebih kritis bagi kreditur daripada bagi investor, karena kreditur tidak memiliki mekanisme untuk menikmati upside jika laporan "terlalu optimis" — mereka hanya menanggung kerugian jika laporan menyembunyikan risiko.')
    add_body_para(doc,
        'Hal ini menciptakan ketegangan langsung dengan preferensi investor: pengukuran nilai wajar (fair value) meningkatkan relevansi dan predictive value bagi investor, tetapi menciptakan volatilitas laba yang dapat memicu pelanggaran debt covenant bagi kreditur — meskipun kualitas kredit perusahaan secara fundamental tidak berubah. Wolk et al. (2017) mengidentifikasi ketegangan ini sebagai salah satu alasan mengapa pelaporan keuangan tidak dapat secara simultan optimal bagi semua pemakai.')

    add_subsection_heading(doc, "Instrumen Analisis Kreditur", "B.")
    add_body_para(doc,
        'Kreditur menggunakan seperangkat instrumen analisis yang berbeda dari investor. Tiga yang paling umum adalah: (1) Altman Z-score, yang mengagregatkan rasio-rasio keuangan ke dalam skor probabilitas kebangkrutan; (2) stress testing, yang mensimulasikan kinerja perusahaan di bawah skenario makroekonomi yang buruk (resesi, kenaikan suku bunga, depresiasi rupiah); dan (3) covenant monitoring, yang memantau apakah perusahaan mendekati atau melampaui ambang batas perjanjian utang — seperti rasio debt-to-equity atau interest coverage ratio. Ketiga instrumen ini membutuhkan informasi yang lengkap, verifiable, dan konsisten antar periode.')
    add_body_para(doc,
        'Relevansi instrumen-instrumen ini didukung oleh bukti empiris yang kuat. Penelitian tentang kebangkrutan mengkonfirmasi kemampuan prediktif model Altman Z-score terhadap financial distress — profitabilitas rendah, arus kas lemah, dan likuiditas terbatas merupakan indikator awal yang terdeteksi melalui data akuntansi jauh sebelum default terjadi. Penelitian tentang peringkat kredit menunjukkan bahwa lembaga seperti Moody\’s Corporation dan Standard & Poor\’s mengandalkan rasio akuntansi — profitabilitas, arus kas operasi, leverage, dan likuiditas — sebagai input primer dalam penentuan peringkat obligasi dan biaya pinjaman perusahaan (Scott, 2015; Wolk et al., 2017). Kualitas informasi akuntansi yang dapat diverifikasi secara langsung memengaruhi biaya modal perusahaan: perusahaan dengan laporan yang lebih transparan dan andal cenderung memperoleh peringkat lebih tinggi, sehingga membayar bunga lebih rendah — konsekuensi ekonomi nyata dari pemenuhan standar verifikabilitas dan kelengkapan. Konsisten dengan temuan ini, obligasi subordinasi BBCA memperoleh peringkat idAA dari Pefindo (BBCA, 2025, hal. 17) — peringkat yang mencerminkan kualitas informasi akuntansi BBCA yang dapat diverifikasi melalui audit Big-4 dan pengawasan prudensial OJK.')

    add_subsection_heading(doc, "Aplikasi: Rasio Prudensial BBCA", "C.")
    add_body_para(doc,
        'Laporan keuangan BBCA menyediakan informasi yang sangat relevan bagi kreditur melalui pengungkapan rasio prudensial yang komprehensif. Liquidity Coverage Ratio (LCR) BBCA 2025 sebesar 310,8% — jauh di atas ambang minimum OJK 100% — mengindikasikan bahwa bank mampu memenuhi kewajiban jangka pendeknya bahkan dalam skenario tekanan likuiditas selama 30 hari (BBCA, 2025, hal. 15). Net Stable Funding Ratio (NSFR) 159,9% menunjukkan bahwa pendanaan jangka panjang BBCA secara substansial melebihi kebutuhan aset yang tidak likuid (hal. 15). Capital Adequacy Ratio (CAR) sebesar 29,8% — hampir tiga kali lipat ambang minimum 10,5% — memberikan cushion modal yang luas bagi kreditur terhadap risiko kerugian tak terduga (hal. 15). Non-Performing Loan (NPL) 1,7% dengan NPL Coverage 183,8% menunjukkan bahwa cadangan kerugian BBCA melampaui seluruh kredit bermasalah yang ada (hal. 205).')
    add_body_para(doc,
        'Kreditur BBCA mencakup pemberi pinjaman antarbank (pinjaman diterima Rp 6,0 triliun) dan pemegang obligasi subordinasi Rp 65 miliar berperingkat idAA oleh Pefindo — jatuh tempo Juli 2030 (BBCA, 2025, hal. 17). Bagi kelompok ini, kenaikan beban CKPN dari Rp 2,0 triliun (2024) menjadi Rp 4,0 triliun (2025) — kenaikan 100% — merupakan sinyal yang perlu diinterpretasikan dengan hati-hati (BBCA, 2025, hal. 210). Di satu sisi, kenaikan ini mencerminkan implementasi model ECL yang prudent — BBCA membangun buffer terhadap risiko makroekonomi meskipun NPL justru turun dari 1,8% ke 1,7%. Di sisi lain, volatilitas beban CKPN yang tinggi dapat mempersulit kreditur dalam memprediksi kemampuan BBCA mempertahankan interest coverage ratio di atas ambang covenant.')
    add_body_para(doc,
        'Keberadaan auditor KAP Rintis, Jumadi, Rianto & Rekan (anggota jaringan PwC) mengurangi — meski tidak menghilangkan — asimetri informasi ini: auditor Big-4 meningkatkan verifikabilitas estimasi internal manajemen, termasuk asumsi model ECL 3-skenario BBCA. Regulator OJK juga memiliki akses ke data granular yang tidak tersedia di laporan publik, sehingga pengawasan prudensial mereka berfungsi sebagai lapisan verifikasi tambahan yang tidak dimiliki kreditur privat.')

    # ------------------------------------------------------------------
    # SECTION V — Kesimpulan
    # ------------------------------------------------------------------
    add_section_heading(doc, "Kesimpulan", "V.")

    add_body_para(doc,
        'Analisis di atas mengungkapkan bahwa laporan keuangan tidak berguna secara abstrak — ia berguna bagi pemakai tertentu, untuk keputusan tertentu, melalui mekanisme tertentu. Investor ekuitas dan kreditur membaca dokumen yang sama, tetapi "kegunaan" yang mereka ekstrak berbeda secara fundamental: investor mengutamakan predictive value dan timeliness untuk memodelkan laba dan nilai intrinsik masa depan; kreditur mengutamakan verifikabilitas dan kelengkapan untuk menilai kemampuan perusahaan tidak gagal bayar bahkan dalam skenario buruk.')
    add_body_para(doc,
        'Ketegangan antara kedua kebutuhan ini bukan anomali yang perlu dieliminasi, melainkan konsekuensi yang tidak terhindarkan dari satu dokumen yang melayani dua fungsi utilitas yang berbeda. Setiap revisi standar — dari PSAK 55 ke PSAK 71 — merupakan repositioning pada trade-off relevansi versus keandalan, bukan resolusinya. Pemilihan ECL model menaikkan predictive value bagi investor tetapi menciptakan volatilitas laba yang perlu dikelola secara hati-hati oleh kreditur.')
    add_body_para(doc,
        'BBCA membuktikan bahwa laporan keuangan yang dirancang dengan standar tertinggi — pengungkapan ECL 3-stage, 16+ rasio prudensial, data 5 tahun, dan audit Big-4 — dapat melayani investor institusional dan kreditur secara bersamaan. Bagi investor, laporan BBCA mendukung model valuasi yang kompleks. Bagi kreditur, ia menyediakan sinyal kehati-hatian yang komprehensif. Dengan demikian, decision usefulness bukan tujuan yang bersaing dengan tujuan lain — ia adalah lensa evaluasi yang menyatukan semua pertimbangan akuntansi dalam satu pertanyaan: berguna bagi siapa, dan untuk keputusan apa?')

    # ------------------------------------------------------------------
    # DAFTAR PUSTAKA
    # ------------------------------------------------------------------
    add_section_heading(doc, "Daftar Pustaka", "")

    add_daftar_pustaka_entry(doc, [
        ("Ball, R., & Brown, P. (1968). An empirical evaluation of accounting income numbers. ", False),
        ("Journal of Accounting Research", True),
        (", 6(2), 159–178.", False),
    ])
    add_daftar_pustaka_entry(doc, [
        ("Bernard, V. L., & Thomas, J. K. (1989). Post-earnings-announcement drift: Delayed price response or risk premium? ", False),
        ("Journal of Accounting Research", True),
        (", 27(Supplement), 1–36.", False),
    ])
    add_daftar_pustaka_entry(doc, [
        ("IASB. (2018). ", False),
        ("Conceptual framework for financial reporting", True),
        (". IFRS Foundation.", False),
    ])
    add_daftar_pustaka_entry(doc, [
        ("Ikatan Akuntan Indonesia. (2020). ", False),
        ("PSAK 71: Instrumen keuangan", True),
        (". Dewan Standar Akuntansi Keuangan IAI.", False),
    ])
    add_daftar_pustaka_entry(doc, [
        ("Lintner, J. (1956). Distribution of incomes of corporations among dividends, retained earnings, and taxes. ", False),
        ("American Economic Review", True),
        (", 46(2), 97–113.", False),
    ])
    add_daftar_pustaka_entry(doc, [
        ("Ohlson, J. A. (1995). Earnings, book values, and dividends in equity valuation. ", False),
        ("Contemporary Accounting Research", True),
        (", 11(2), 661–687.", False),
    ])
    add_daftar_pustaka_entry(doc, [
        ("PT Bank Central Asia Tbk. (2025). ", False),
        ("Laporan tahunan 2025: Unity for a better future", True),
        (". BCA.", False),
    ])
    add_daftar_pustaka_entry(doc, [
        ("Scott, W. R. (2015). ", False),
        ("Financial accounting theory", True),
        (" (7th ed.). Pearson Education Canada.", False),
    ])
    add_daftar_pustaka_entry(doc, [
        ("Wolk, H. I., Dodd, J. L., & Rozycki, J. J. (2017). ", False),
        ("Accounting theory: Conceptual issues in a political and economic environment", True),
        (" (9th ed.). SAGE Publications.", False),
    ])

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir,
                            "01079_Dzaki Muhammad Yusfian_Tugas Individual 4.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    print(f"Size:  {os.path.getsize(out_path):,} bytes")


if __name__ == "__main__":
    build()
