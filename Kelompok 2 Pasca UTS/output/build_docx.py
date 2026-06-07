"""
build_docx.py
Builds "Kelompok 2_RMK Chapter 13 Statement of Cash Flows.docx"
from ../content/rmk-ch13.md using python-docx.
Styling helpers adapted from ti5/output/build_docx.py.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"

MEMBERS = [
    ("Satriyo Nugroho", "122501039"),
    ("Mario Da Costa", "122501048"),
    ("Amelda Putri Zhany Wiguna", "122501067"),
    ("Ahmad Ramadhan", "122501078"),
    ("Nida Nur Cahyati", "122501084"),
    ("Priska Putri Parungky", "122501094"),
]


# ---------------------------------------------------------------------------
# Markdown parsing layer (unit-tested)
# ---------------------------------------------------------------------------

def parse_inline_runs(text):
    """Split markdown text into (text, bold, italic) run tuples.

    Supports **bold** and *italic*, non-nested.
    """
    runs = []
    pattern = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], False, False))
        token = m.group(0)
        if token.startswith('**'):
            runs.append((token[2:-2], True, False))
        else:
            runs.append((token[1:-1], False, True))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], False, False))
    return [r for r in runs if r[0]]


def parse_blocks(md_text):
    """Parse markdown into ('heading'|'para'|'bullet'|'ref', text) blocks.

    '# ' title lines are skipped (cover is hardcoded). Consecutive non-blank
    lines join into one paragraph. After a heading containing 'Daftar
    Pustaka', paragraphs become 'ref' blocks (hanging indent).
    """
    blocks = []
    in_refs = False
    buf = []

    def flush():
        nonlocal buf
        if buf:
            blocks.append(("ref" if in_refs else "para", " ".join(buf)))
            buf = []

    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            flush()
        elif line.startswith("## "):
            flush()
            text = line[3:].strip()
            blocks.append(("heading", text))
            in_refs = "daftar pustaka" in text.lower()
        elif line.startswith("# "):
            flush()
        elif line.startswith("- "):
            flush()
            blocks.append(("bullet", line[2:].strip()))
        else:
            buf.append(line.strip())
    flush()
    return blocks


# ---------------------------------------------------------------------------
# Docx styling layer (adapted from ti5/output/build_docx.py)
# ---------------------------------------------------------------------------

def _style_run(run, font_size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rFonts.set(qn('w:cs'), FONT)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _add_para(doc, runs, font_size=12, bold=False,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
              space_before_pt=0, space_after_pt=6,
              left_indent_cm=None, hanging=False):
    """runs: list of (text, bold, italic) tuples."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    if left_indent_cm is not None:
        pf.left_indent = Cm(left_indent_cm)
    if hanging:
        pf.left_indent = Cm(0.75)
        pf.first_line_indent = Cm(-0.75)
    for text, b, i in runs:
        run = para.add_run(text)
        _style_run(run, font_size=font_size, bold=bold or b, italic=i)
    return para


def add_cover_line(doc, text, font_size=12, bold=False):
    return _add_para(doc, [(text, False, False)], font_size=font_size,
                     bold=bold, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def add_blank(doc):
    return _add_para(doc, [("", False, False)], space_after_pt=0)


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def build():
    here = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(here, "..", "content", "rmk-ch13.md")
    with open(md_path, encoding="utf-8") as f:
        md_text = f.read()

    doc = Document()

    # Page setup: A4, 3 cm margins (ti5 pattern)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)

    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(12)

    # Cover
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)
    add_cover_line(doc, "RINGKASAN MATERI KULIAH (RMK)", font_size=16, bold=True)
    add_cover_line(doc, "Pelaporan Keuangan Korporat", font_size=14)
    add_cover_line(doc, "Statement of Cash Flows", font_size=14, bold=True)
    add_cover_line(doc,
        "Wolk, Dodd & Rozycki — Accounting Theory: Conceptual Issues in a "
        "Political and Economic Environment (Chapter 13)", font_size=12)
    add_blank(doc)
    add_blank(doc)
    add_cover_line(doc, "Disusun oleh Kelompok 2:", font_size=12)
    for name, nim in MEMBERS:
        add_cover_line(doc, f"{name} — {nim}", font_size=12, bold=True)
    add_blank(doc)
    add_blank(doc)
    add_cover_line(doc, "Program Studi Magister Akuntansi", font_size=12)
    add_cover_line(doc, "Juni 2026", font_size=12, bold=True)
    doc.add_page_break()

    # Content from markdown
    for kind, text in parse_blocks(md_text):
        runs = parse_inline_runs(text)
        if kind == "heading":
            _add_para(doc, runs, font_size=13, bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT,
                      space_before_pt=12, space_after_pt=6)
        elif kind == "bullet":
            _add_para(doc, [("• ", False, False)] + runs,
                      left_indent_cm=0.75)
        elif kind == "ref":
            _add_para(doc, runs, hanging=True)
        else:
            _add_para(doc, runs)

    out_path = os.path.join(
        here, "Kelompok 2_RMK Chapter 13 Statement of Cash Flows.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    print(f"Size:  {os.path.getsize(out_path):,} bytes")


if __name__ == "__main__":
    build()
