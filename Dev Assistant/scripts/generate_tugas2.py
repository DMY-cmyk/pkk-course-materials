from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "01079_Dzaki Muhammad Yusfian_Tugas Individual 2.docx"
)


def set_margins(doc, top=3, bottom=3, left=3, right=3):
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def add_para(doc, text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=6, line_spacing=18,
             first_indent=False, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(line_spacing)
    if first_indent:
        pf.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_heading(doc, text, level=1):
    add_para(doc, text, bold=True, size=12 if level > 1 else 13,
             space_before=12, space_after=6, line_spacing=18,
             align=WD_ALIGN_PARAGRAPH.LEFT)


def add_ref(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = Pt(18)
    pf.left_indent = Cm(1.25)
    pf.first_line_indent = Cm(-1.25)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def build_document():
    doc = Document()
    set_margins(doc)

    # ── COVER ─────────────────────────────────────────────────────────
    add_para(doc, 'MNK202 PELAPORAN KEUANGAN KORPORAT', bold=True, center=True,
             size=12, space_after=4)
    add_para(doc, 'TUGAS INDIVIDUAL 2', bold=True, center=True,
             size=14, space_after=4)
    add_para(doc,
             'Regulasi Standar Akuntansi: Antara Kepentingan Publik dan '
             'Mekanisme Pasar \u2014 Sebuah Tinjauan Kritis',
             bold=True, center=True, size=12, space_before=4, space_after=16)
    add_para(doc, 'Disusun oleh:', center=True, size=12, space_after=2)
    add_para(doc, 'Dzaki Muhammad Yusfian', bold=True, center=True, size=12, space_after=2)
    add_para(doc, 'NIM: 1125 01079', center=True, size=12, space_after=2)
    add_para(doc, 'Program Studi Magister Akuntansi', center=True, size=12, space_after=2)
    add_para(doc, 'STIE YKPN Yogyakarta', center=True, size=12, space_after=2)
    add_para(doc, 'Dosen: Prof. Djoko Susanto', center=True, size=12, space_after=2)
    add_para(doc, 'April 2026', center=True, size=12, space_after=24)
    doc.add_page_break()

    # ── I. PENDAHULUAN ────────────────────────────────────────────────
    add_heading(doc, 'I. Pendahuluan')
    add_para(doc,
        'Dalam khazanah akuntansi, angka-angka keuangan bukan sekadar representasi '
        'teknis atas transaksi ekonomi. Wolk, Dodd, dan Rozycki (2017) mengingatkan '
        'bahwa angka akuntansi memiliki \u201crealitas sosial\u201d (social reality): '
        'pilihan metode inventori memengaruhi beban pajak, angka laba memengaruhi '
        'bonus manajemen, dan rasio neraca memengaruhi kemampuan perusahaan mengakses '
        'kredit. Dengan konsekuensi nyata sebesar ini, pertanyaan apakah standar '
        'akuntansi sebaiknya diregulasi atau diserahkan kepada mekanisme pasar bebas '
        'menjadi perdebatan yang jauh melampaui ranah teknis.',
        first_indent=True)
    add_para(doc,
        'Esai ini menelaah kedua kubu argumentasi secara berimbang. Kubu pro-regulasi '
        'bersandar pada teori kegagalan pasar, asimetri informasi, dan masalah '
        'keagenan. Kubu anti-regulasi \u2014 yang diwakili oleh Teori Akuntansi '
        'Positif dan hipotesis pasar efisien \u2014 berargumen bahwa pasar dapat '
        'mendisiplinkan dirinya sendiri. Setelah mensintesis kedua kubu dalam konteks '
        'pasar modal Indonesia, penulis berpendapat bahwa regulasi standar akuntansi '
        'lebih superior daripada laissez-faire, dengan catatan regulasi harus '
        'dirancang dengan proses yang transparan dan bebas dari regulatory capture.',
        first_indent=True)

    # ── II. PRO-REGULASI ──────────────────────────────────────────────
    add_heading(doc, 'II. Argumen Pro-Regulasi')

    add_heading(doc, 'A. Kegagalan Pasar dan Sifat Barang Publik Informasi', level=2)
    add_para(doc,
        'Informasi akuntansi memiliki karakteristik barang publik (public good): '
        'sekali diproduksi dan dipublikasikan, penggunaannya oleh satu pihak tidak '
        'mengurangi ketersediaannya bagi pihak lain (non-rival), dan sulit untuk '
        'mengecualikan pihak tertentu dari mengaksesnya (non-excludable). Dalam '
        'kondisi ini, perusahaan yang tidak diwajibkan mengungkap informasi akan '
        'cenderung underinvest dalam pengungkapan karena biaya produksi informasi '
        'ditanggung sendiri sementara manfaatnya dinikmati oleh semua pihak tanpa '
        'harus membayar (free rider problem). Stigler (1971) menunjukkan bahwa '
        'kegagalan pasar semacam ini adalah justifikasi fundamental bagi intervensi '
        'regulasi. Tanpa standar wajib, pasar akan menghasilkan pengungkapan yang '
        'secara sistematis lebih sedikit dari yang optimal secara sosial.',
        first_indent=True)

    add_heading(doc, 'B. Asimetri Informasi dan Seleksi Berlawanan (Adverse Selection)', level=2)
    add_para(doc,
        'Akerlof (1970) dalam makalah seminalnya \u201cThe Market for \u2018Lemons\u2019\u201d '
        'menunjukkan bagaimana asimetri informasi dapat menyebabkan kegagalan pasar yang '
        'sistemik. Dalam pasar di mana penjual (manajemen perusahaan) memiliki informasi '
        'yang jauh lebih lengkap daripada pembeli (investor), investor rasional akan '
        'mendiskon semua sekuritas \u2014 baik yang berkualitas tinggi maupun rendah '
        '\u2014 ke tingkat rata-rata pasar. Akibatnya, perusahaan berkualitas tinggi '
        'menerima harga yang tidak adil dan cenderung keluar dari pasar, meninggalkan '
        'hanya perusahaan berkualitas rendah. Regulasi mandatory disclosure melalui '
        'standar akuntansi yang seragam memotong asimetri informasi ini dengan '
        'menetapkan lantai minimum pengungkapan yang dapat dipercaya (Scott, 2015).',
        first_indent=True)

    add_heading(doc, 'C. Konflik Keagenan dan Moral Hazard', level=2)
    add_para(doc,
        'Jensen dan Meckling (1976) memformalkan hubungan keagenan antara pemegang '
        'saham (principal) dan manajemen (agent). Manajer yang tidak diawasi memiliki '
        'insentif untuk berperilaku oportunistik \u2014 mengelola laba, menyembunyikan '
        'kerugian, atau mengambil risiko berlebihan \u2014 karena mereka menanggung '
        'sebagian kecil biaya keputusan buruk namun menikmati sebagian besar keuntungan '
        'dari keputusan baik. Regulasi standar akuntansi memaksa transparansi yang '
        'membatasi perilaku oportunistik ini. Bukti historis paling dramatis adalah '
        'skandal Enron (2001) dan WorldCom (2002), yang terjadi sebagian karena '
        'lemahnya regulasi atas special purpose entities (SPE) \u2014 sesuatu yang '
        'Wolk et al. (2017, hlm. 5) tunjukkan sebagai kegagalan FASB yang kemudian '
        'dipaksa dibenahi oleh tekanan publik pascaskandal.',
        first_indent=True)

    # ── III. ANTI-REGULASI ────────────────────────────────────────────
    add_heading(doc, 'III. Argumen Anti-Regulasi')

    add_heading(doc, 'A. Teori Akuntansi Positif dan Kontrak Privat', level=2)
    add_para(doc,
        'Watts dan Zimmerman (1978) meletakkan dasar Teori Akuntansi Positif '
        '(Positive Accounting Theory/PAT) dengan berargumen bahwa perusahaan dan '
        'pemegang kepentingannya dapat menyelesaikan masalah keagenan secara mandiri '
        'melalui kontrak privat. Debt covenants dalam perjanjian kredit, kompensasi '
        'manajemen berbasis kinerja, dan mekanisme bonding lainnya adalah contoh '
        'bagaimana pasar menciptakan insentif yang menyelaraskan kepentingan manajer '
        'dengan pemegang saham tanpa campur tangan regulator. Dari perspektif ini, '
        'regulasi akuntansi bukan hanya tidak perlu \u2014 ia dapat merusak '
        'keseimbangan kontraktual yang telah terbentuk secara organik dan '
        'menambahkan beban compliance cost yang tidak proporsional.',
        first_indent=True)

    add_heading(doc, 'B. Pasar Efisien dan Disiplin Pasar', level=2)
    add_para(doc,
        'Fama (1970) berargumen bahwa pasar modal yang efisien secara cepat '
        'menginkorporasi semua informasi publik yang tersedia ke dalam harga '
        'sekuritas. Dalam kondisi ini, investor canggih akan mengidentifikasi dan '
        'mendiskon perusahaan dengan pengungkapan yang buruk atau tidak dapat '
        'dipercaya \u2014 menciptakan insentif berbasis pasar bagi perusahaan untuk '
        'secara sukarela mengungkap informasi berkualitas tinggi demi mendapatkan '
        'biaya modal (cost of capital) yang lebih rendah. Jika mekanisme ini '
        'berfungsi, regulasi pengungkapan menjadi redundan: pasar sudah '
        '\u201cmemaksa\u201d pengungkapan optimal tanpa perlu intervensi. Scott '
        '(2015) mencatat bahwa implikasi efisiensi pasar bagi akuntansi adalah '
        'bahwa konten informasi lebih penting dari format atau lokasi pengungkapan.',
        first_indent=True)

    add_heading(doc, 'C. Regulatory Capture dan Distorsi Politik', level=2)
    add_para(doc,
        'Argumen terkuat dari kubu anti-regulasi mungkin bukan bahwa regulasi tidak '
        'diperlukan, melainkan bahwa regulasi dalam praktiknya cenderung gagal '
        'memenuhi tujuannya. Stigler (1971) menunjukkan bahwa regulator secara '
        'sistematis cenderung \u201cditangkap\u201d (captured) oleh industri yang '
        'mereka regulasi \u2014 menghasilkan aturan yang melayani kepentingan '
        'preparer daripada pengguna informasi keuangan. Ironinya, contoh yang '
        'disajikan Wolk et al. (2017, hlm. 5) \u2014 kegagalan FASB menyelesaikan '
        'masalah SPE karena tekanan politik dari Big Five \u2014 justru '
        'mengilustrasikan fenomena regulatory capture ini secara sempurna. Namun '
        'penulis melihat ironi ini sebagai argumen untuk regulasi yang lebih kuat '
        'dan independen, bukan untuk deregulasi.',
        first_indent=True)

    # ── IV. SINTESIS INDONESIA ────────────────────────────────────────
    add_heading(doc, 'IV. Sintesis: Regulasi dalam Konteks Indonesia')
    add_para(doc,
        'Perdebatan antara regulasi dan pasar bebas dalam akuntansi tidak dapat '
        'dilepaskan dari konteks institusional di mana standar tersebut diterapkan. '
        'Dalam konteks Indonesia, penulis berargumen bahwa asumsi-asumsi yang '
        'menopang argumen anti-regulasi tidak terpenuhi secara memadai.',
        first_indent=True)
    add_para(doc,
        'Pertama, pasar modal Indonesia masih dalam tahap berkembang dengan tingkat '
        'efisiensi yang lebih rendah dibandingkan pasar maju seperti Amerika Serikat '
        'atau Inggris Raya. Mekanisme kontrak privat yang menjadi andalan Teori '
        'Akuntansi Positif (Watts & Zimmerman, 1978) membutuhkan infrastruktur hukum '
        'dan pasar yang kuat \u2014 kondisi yang belum sepenuhnya tersedia di Indonesia.',
        first_indent=True)
    add_para(doc,
        'Kedua, konvergensi PSAK dengan International Financial Reporting Standards '
        '(IFRS) yang dimulai secara bertahap sejak 2012 telah memberikan bukti empiris '
        'keberhasilan regulasi: meningkatnya komparabilitas laporan keuangan '
        'antarperusahaan, meningkatnya kepercayaan investor asing, dan berkurangnya '
        'biaya modal bagi emiten Indonesia yang membutuhkan pendanaan internasional '
        '(DSAK IAI, 2024).',
        first_indent=True)
    add_para(doc,
        'Ketiga, regulasi mandatory disclosure oleh Otoritas Jasa Keuangan (OJK) '
        'memainkan peran krusial dalam melindungi jutaan investor ritel yang '
        'mendominasi basis investor BEI. Tanpa standar minimum yang dapat diandalkan, '
        'investor ritel \u2014 yang tidak memiliki sumber daya untuk melakukan '
        'analisis mendalam \u2014 akan sangat rentan terhadap manipulasi informasi.',
        first_indent=True)
    add_para(doc,
        'Keempat, Indonesia telah memilih jalur regulasi principles-based berbasis '
        'IFRS \u2014 bukan rules-based yang kaku \u2014 sambil tetap mempertahankan '
        'panduan implementasi melalui OJK. Model hybrid ini mengakomodasi fleksibilitas '
        'yang dibutuhkan oleh keragaman industri Indonesia sekaligus mempertahankan '
        'akuntabilitas yang diperlukan untuk melindungi kepentingan publik.',
        first_indent=True)

    # ── V. KESIMPULAN ─────────────────────────────────────────────────
    add_heading(doc, 'V. Kesimpulan')
    add_para(doc,
        'Perdebatan regulasi versus pasar bebas dalam akuntansi adalah perdebatan '
        'tentang seberapa sempurna pasar bekerja dalam menghasilkan pengungkapan yang '
        'optimal. Argumen Watts & Zimmerman (1978) dan Fama (1970) valid secara '
        'teoritis \u2014 namun mengandaikan kondisi pasar ideal: efisiensi tinggi, '
        'kontrak yang dapat ditegakkan, dan investor yang sepenuhnya rasional. Di '
        'dunia nyata, dan khususnya di konteks pasar modal Indonesia yang sedang '
        'berkembang, kondisi ideal ini jauh dari terpenuhi.',
        first_indent=True)
    add_para(doc,
        'Kegagalan pasar, asimetri informasi, dan konflik keagenan adalah realitas '
        'struktural \u2014 bukan anomali sementara yang akan hilang dengan sendirinya. '
        'Oleh karena itu, penulis berpendapat bahwa standar akuntansi lebih baik '
        'diregulasi daripada tidak diregulasi. Namun regulasi yang baik adalah '
        'regulasi yang dirancang dengan proses yang transparan, inklusif, dan bebas '
        'dari regulatory capture \u2014 sesuatu yang terus menjadi tantangan bagi '
        'FASB, IASB, maupun DSAK IAI. Indonesia, dengan model regulasi '
        'principles-based yang diawasi OJK, telah memilih jalur yang secara '
        'konseptual tepat \u2014 meskipun implementasinya masih memerlukan penguatan '
        'yang berkelanjutan.',
        first_indent=True)

    # ── DAFTAR PUSTAKA ────────────────────────────────────────────────
    add_heading(doc, 'Daftar Pustaka')
    refs = [
        'Akerlof, G. A. (1970). The market for \u201clemons\u201d: Quality uncertainty '
        'and the market mechanism. Quarterly Journal of Economics, 84(3), 488\u2013500.',

        'DSAK IAI. (2024). Standar Akuntansi Keuangan. Jakarta: Ikatan Akuntan Indonesia.',

        'Fama, E. F. (1970). Efficient capital markets: A review of theory and empirical '
        'work. Journal of Finance, 25(2), 383\u2013417.',

        'IASB. (2018). Conceptual Framework for Financial Reporting. London: IFRS Foundation.',

        'Jensen, M. C., & Meckling, W. H. (1976). Theory of the firm: Managerial behavior, '
        'agency costs and ownership structure. Journal of Financial Economics, 3(4), 305\u2013360.',

        'OJK. (2024). Peraturan OJK tentang Keterbukaan Informasi Emiten dan Perusahaan '
        'Publik. Jakarta: Otoritas Jasa Keuangan.',

        'Scott, W. R. (2015). Financial Accounting Theory (7th ed.). Toronto: Pearson.',

        'Stigler, G. J. (1971). The theory of economic regulation. Bell Journal of Economics '
        'and Management Science, 2(1), 3\u201321.',

        'Watts, R. L., & Zimmerman, J. L. (1978). Towards a positive theory of the '
        'determination of accounting standards. The Accounting Review, 53(1), 112\u2013134.',

        'Wolk, H. I., Dodd, J. L., & Rozycki, J. J. (2017). Accounting Theory: Conceptual '
        'Issues in a Political and Economic Environment (9th ed.). Thousand Oaks: SAGE.',
    ]
    for ref in refs:
        add_ref(doc, ref)

    doc.save(OUTPUT_PATH)
    print(f'Saved: {OUTPUT_PATH}')
    return OUTPUT_PATH


if __name__ == '__main__':
    path = build_document()
    doc = Document(path)
    all_text = ' '.join(p.text for p in doc.paragraphs)
    word_count = len(all_text.split())
    print(f'Word count (approx): {word_count}')
    headings = [p.text for p in doc.paragraphs
                if p.runs and p.runs[0].bold
                and p.text.startswith(('I.', 'II.', 'III.', 'IV.', 'V.'))]
    print(f'Main sections: {headings}')
    assert word_count >= 1200, f'Too short: {word_count}'
    assert len(headings) >= 5, f'Missing sections: {headings}'
    print('Verification passed.')
