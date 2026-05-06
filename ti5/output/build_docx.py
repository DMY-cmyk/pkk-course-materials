"""
build_docx.py
Builds 01079_Dzaki Muhammad Yusfian_Tugas Individual 5.docx
using python-docx 1.2.0.
Adapted from ti4/output/build_docx.py — identical helper functions, new content.
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers (identical to ti4/output/build_docx.py)
# ---------------------------------------------------------------------------

def add_paragraph(doc, text, font_name="Times New Roman", font_size=12,
                  bold=False, italic=False,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  line_spacing_pt=18, space_before_pt=0, space_after_pt=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)
    return para


def add_section_heading(doc, text, number_prefix, font_size=13):
    full_text = f"{number_prefix}   {text}" if number_prefix else text
    return add_paragraph(doc, full_text, font_size=font_size, bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         line_spacing_pt=18, space_before_pt=12, space_after_pt=6)


def add_body_para(doc, text):
    return add_paragraph(doc, text, font_size=12, bold=False,
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         line_spacing_pt=18, space_before_pt=0, space_after_pt=6)


def add_blank(doc):
    return add_paragraph(doc, "", line_spacing_pt=18,
                         space_before_pt=0, space_after_pt=0)


def add_cover_para(doc, text, font_size=12, bold=False):
    return add_paragraph(doc, text, font_size=font_size, bold=bold,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing_pt=18, space_before_pt=0, space_after_pt=6)


def add_daftar_pustaka_entry(doc, text_parts):
    """text_parts: list of (text, italic) tuples. Hanging indent."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.left_indent = Cm(0.75)
    pf.first_line_indent = Cm(-0.75)
    for text, italic in text_parts:
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.italic = italic
        run.bold = False
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), "Times New Roman")
        rFonts.set(qn('w:hAnsi'), "Times New Roman")
        rFonts.set(qn('w:cs'), "Times New Roman")
        existing = rPr.find(qn('w:rFonts'))
        if existing is not None:
            rPr.remove(existing)
        rPr.insert(0, rFonts)
    return para


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # Page setup: A4, 3cm all margins
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # ------------------------------------------------------------------
    # COVER
    # ------------------------------------------------------------------
    add_blank(doc)
    add_blank(doc)
    add_blank(doc)
    add_cover_para(doc, "TUGAS INDIVIDU 5", font_size=16, bold=True)
    add_cover_para(doc, "Pelaporan Keuangan Korporat", font_size=14, bold=False)
    add_cover_para(doc,
        "Mengapa Disclosure Segment dalam SFAS No. 131 Merupakan Perbaikan "
        "dari Disclosure Segment dalam SFAS 14",
        font_size=14, bold=True)
    add_blank(doc)
    add_blank(doc)
    add_cover_para(doc, "Disusun oleh:", font_size=12, bold=False)
    add_cover_para(doc, "Dzaki Muhammad Yusfian", font_size=12, bold=True)
    add_cover_para(doc, "1225 01079", font_size=12, bold=False)
    add_blank(doc)
    add_blank(doc)
    add_cover_para(doc, "Dosen Pengampu: Prof. Djoko Susanto", font_size=12, bold=False)
    add_cover_para(doc, "Program Studi Magister Akuntansi", font_size=12, bold=False)
    add_cover_para(doc, "Mei 2026", font_size=12, bold=True)

    doc.add_page_break()

    # ------------------------------------------------------------------
    # SECTION I — Pendahuluan
    # ------------------------------------------------------------------
    add_section_heading(doc, "Pendahuluan", "I.")

    add_body_para(doc,
        'Laporan keuangan segmental bertujuan memberikan informasi tentang kinerja berbagai komponen bisnis suatu perusahaan — informasi yang tidak tersedia jika hanya membaca laporan keuangan konsolidasi secara keseluruhan. Dalam kerangka Wolk, Dodd & Rozycki (2017), pengungkapan segmental adalah bentuk informative disclosure: ia ditujukan bukan sekadar untuk melindungi investor dari informasi tersembunyi, melainkan untuk memperkaya analisis risk-return dengan memberikan gambaran granular tentang dari mana laba dan risiko perusahaan sesungguhnya berasal. Ketika pengungkapan segmental berhasil, pengguna dapat mengevaluasi setiap komponen bisnis seolah mereka sedang menganalisis perusahaan yang lebih kecil dan lebih fokus.')

    add_body_para(doc,
        'SFAS No. 14 (1976) dan SFAS No. 131 (1997) mewakili dua filosofi yang sangat berbeda tentang bagaimana tujuan ini seharusnya dicapai. Penggantian SFAS 14 oleh SFAS 131 setelah lebih dari dua dekade bukan revisi kosmetik: ia mencerminkan pengakuan bahwa SFAS 14 telah secara sistematis gagal dalam tujuan informative disclosure yang menjadi alasan eksistensinya. SFAS 14 memberi manajemen kebebasan luas untuk mendefinisikan “industri” — dan kebebasan itu dieksploitasi untuk meminimalkan pengungkapan alih-alih memaksimalkan kegunaan informasi bagi pengguna.')

    add_body_para(doc,
        'Esai ini berargumen bahwa SFAS 131 merupakan perbaikan substantif atas SFAS 14 karena tiga alasan yang saling terkait. Pertama, SFAS 131 menggantikan kebebasan definitorial tanpa kendali dengan jangkar organisasional yang dapat diobservasi — konsep Chief Operating Decision Maker (CODM). Kedua, standar ini menyelaraskan laporan eksternal dengan laporan internal, mengurangi asimetri informasi secara struktural. Ketiga, perbaikan ini bukan hanya teoritis: bukti empiris dari studi komparatif lintas periode mengonfirmasi bahwa pengungkapan segmental di bawah SFAS 131 lebih granular, lebih lengkap, dan lebih prediktif secara ekonomi dibandingkan pendahulunya (Wolk et al., 2017; Scott, 2015).')

    # ------------------------------------------------------------------
    # SECTION II — SFAS 14 dan Tiga Kelemahannya
    # ------------------------------------------------------------------
    add_section_heading(doc, "SFAS No. 14 dan Tiga Kelemahannya", "II.")

    add_body_para(doc,
        'SFAS No. 14, yang diterbitkan pada tahun 1976, mewajibkan perusahaan publik untuk mengungkapkan informasi keuangan berdasarkan segmen industri dan area geografis. Namun, desain standar ini mengandung kelemahan mendasar yang baru sepenuhnya disadari setelah beberapa dekade implementasi.')

    add_body_para(doc,
        'Kelemahan pertama adalah flexibility yang menyamar sebagai uniformity. SFAS 14 memperbolehkan manajemen mendefinisikan “industri” menggunakan berbagai basis — produk, proses produksi, atau saluran pemasaran — tanpa menetapkan relevant circumstances yang spesifik sebagai dasar pemilihan. Dalam kerangka Wolk (2017, hal. 247), ini bukan finite uniformity — di mana metode berbeda diizinkan karena kondisi ekonomi yang berbeda secara substansial dapat diidentifikasi — melainkan lebih mendekati flexibility: kebebasan memilih tanpa kriteria ekonomi yang dapat dipertanggungjawabkan. Hasilnya adalah komparabilitas nominal tanpa komparabilitas substantif: dua perusahaan dengan bisnis yang identik secara ekonomi dapat menghasilkan laporan segmen yang sangat berbeda berdasarkan pilihan definitorial manajemen, bukan perbedaan ekonomi yang nyata.')

    add_body_para(doc,
        'Kelemahan kedua adalah single-segment gaming — eksploitasi kebebasan definitorial untuk meminimalkan pengungkapan. Karena SFAS 14 tidak menetapkan batasan yang mengikat terhadap agregasi, banyak perusahaan mengklaim beroperasi dalam satu segmen meskipun menjalankan bisnis yang beragam secara ekonomi. Street, Nichols & Gray (2000) mengonfirmasi bahwa proporsi yang tidak proporsional dari perusahaan sampel mereka melaporkan sebagai entitas satu-segmen di bawah SFAS 14 — sebuah pola yang tidak konsisten dengan realitas ekonomi perusahaan-perusahaan yang terdiversifikasi. Ini adalah manipulasi-by-design: ketika standar tidak memberikan kendala yang bermakna, manajemen rasional memilih definisi yang meminimalkan beban pengungkapan dan potensi kerugian kompetitif dari paparan informasi granular kepada pesaing.')

    add_body_para(doc,
        'Kelemahan ketiga — dan yang paling fundamental — adalah disconnect antara pengungkapan eksternal dan struktur manajemen internal. SFAS 14 tidak mensyaratkan bahwa segmen yang diungkap kepada publik harus sesuai dengan cara manajemen sesungguhnya mengorganisasi dan memantau bisnis. Konsekuensinya: pengguna laporan keuangan menerima peta segmen yang mungkin tidak digunakan oleh kepemimpinan perusahaan itu sendiri dalam pengambilan keputusan. Ini menghancurkan tujuan inti informative disclosure: ketika unit yang diungkap tidak mencerminkan cara manajemen mengevaluasi kinerja dan mengalokasikan sumber daya, pengguna eksternal tidak dapat menggunakan informasi segmen untuk menilai bagaimana perusahaan sesungguhnya dijalankan. Asimetri informasi yang dihasilkan bersifat struktural, bukan insidental — ia tertanam dalam desain standar itu sendiri.')

    # ------------------------------------------------------------------
    # SECTION III — SFAS 131: Pendekatan Manajemen sebagai Koreksi Struktural
    # ------------------------------------------------------------------
    add_section_heading(doc,
        "SFAS No. 131: Pendekatan Manajemen sebagai Koreksi Struktural", "III.")

    add_body_para(doc,
        'SFAS No. 131, yang diterbitkan FASB pada tahun 1997 setelah rekomendasi komite khusus AICPA, mengoreksi ketiga kelemahan SFAS 14 melalui satu konsep sentral: pendekatan manajemen (management approach). Alih-alih meminta perusahaan mendefinisikan segmen berdasarkan taksonomi industri yang dipilih manajemen, SFAS 131 mendefinisikan segmen operasi berdasarkan cara manajemen sendiri mengorganisasi perusahaan untuk tujuan pengambilan keputusan operasional dan evaluasi kinerja (FASB, 1997).')

    add_body_para(doc,
        'Inovasi pertama adalah konsep Chief Operating Decision Maker (CODM) sebagai jangkar definitorial. Segmen dilaporkan sebagai komponen-komponen yang tentangnya CODM membuat keputusan alokasi sumber daya dan mengevaluasi kinerja. Konsep ini menggantikan kebebasan definitorial SFAS 14 dengan kriteria yang dapat diobservasi dan dapat diaudit: segmen yang dilaporkan harus konsisten dengan struktur pelaporan internal yang sesungguhnya digunakan oleh kepemimpinan perusahaan. Jika CODM menerima laporan kinerja bulanan untuk tiga unit bisnis, maka tiga unit itulah yang menjadi segmen operasi yang dilaporkan. Tidak ada ruang untuk mendefinisikan ulang segmen agar lebih “ramping” secara strategis.')

    add_body_para(doc,
        'Inovasi kedua adalah keselarasan struktural antara laporan eksternal dan laporan internal. Dengan mensyaratkan bahwa segmen yang diungkap sesuai dengan segmen yang digunakan untuk manajemen internal, SFAS 131 mengurangi asimetri informasi secara struktural — bukan melalui pengungkapan tambahan yang diwajibkan secara artifisial, melainkan dengan menjadikan laporan publik sebagai cermin dari laporan yang sudah ada secara internal. Pengguna eksternal menerima dekomposisi bisnis yang sama yang digunakan manajemen untuk mengalokasikan modal dan mengevaluasi manajer unit — secara teoritis, dekomposisi yang paling relevan secara keputusan yang tersedia (Wolk et al., 2017, hal. 257). Ini adalah informative disclosure dalam pengertian yang paling murni: bukan informasi yang diproduksi khusus untuk pelaporan publik, melainkan informasi yang sudah ada dan digunakan secara internal, sekarang dibuat dapat diakses oleh pengguna eksternal.')

    add_body_para(doc,
        'Inovasi ketiga adalah ambang batas kuantitatif sebagai finite uniformity. SFAS 131 menetapkan bahwa segmen dilaporkan secara terpisah jika memenuhi setidaknya satu dari tiga kriteria: memiliki 10% atau lebih dari total pendapatan konsolidasi, 10% dari total laba/rugi semua segmen operasi, atau 10% dari total aset semua segmen. Selain itu, segmen-segmen yang dilaporkan secara terpisah harus secara kolektif mencakup setidaknya 75% dari total pendapatan konsolidasi (FASB, 1997). Ambang batas ini adalah bright-line criteria dalam pengertian Wolk: mereka membatasi agregasi tanpa mendiktekan struktur segmen. Dalam kerangka uniformity, SFAS 131 menerapkan finite uniformity pada pengungkapan: relevant circumstances (struktur organisasional yang tercermin dalam praktik CODM) menentukan batas segmen, sementara ambang kuantitatif memastikan granularitas minimum. SFAS 14, sebaliknya, tidak memberikan disiplin setara.')

    # ------------------------------------------------------------------
    # SECTION IV — Bukti Empiris dan Analisis Kritis
    # ------------------------------------------------------------------
    add_section_heading(doc, "Bukti Empiris dan Analisis Kritis", "IV.")

    add_body_para(doc,
        'Perbaikan konseptual yang ditawarkan SFAS 131 tidak akan bernilai jika tidak terwujud dalam praktik pelaporan. Bukti empiris yang tersedia memberikan konfirmasi kuat — dengan satu catatan analitis penting.')

    add_body_para(doc,
        'Bukti pertama dan paling langsung datang dari studi Street, Nichols & Gray (2000), yang membandingkan pelaporan segmental tahun 1997 (di bawah SFAS 14) dengan tahun 1998 (tahun pertama SFAS 131 berlaku) untuk populasi perusahaan yang sama. Temuan mereka konsisten pada tiga dimensi: (1) terjadi penurunan signifikan dalam jumlah perusahaan yang mengklaim beroperasi sebagai entitas satu-segmen, (2) rata-rata jumlah segmen yang dilaporkan meningkat, dan (3) jumlah item yang diungkap per segmen meningkat. Ini adalah bukti kausal paling bersih yang tersedia: variabel yang berubah hanyalah standar yang berlaku — perusahaan, industri, dan kondisi ekonomi tetap sama. Peningkatan dalam granularitas dan kelengkapan pengungkapan dapat secara wajar diatribusikan kepada perubahan standar.')

    add_body_para(doc,
        'Bukti kedua berkaitan dengan validitas prediktif informasi segmen. Penelitian terkini menemukan hubungan yang lebih kuat antara pendapatan segmen periode berjalan dan pendapatan periode berikutnya untuk perusahaan yang melaporkan di bawah SFAS 131 dibandingkan SFAS 14 (Wolk et al., 2017). Ini adalah uji decision usefulness dalam pengertian Scott (2015): informasi berguna jika membantu pengguna membentuk prediksi yang lebih baik. Validitas prediktif yang lebih tinggi berarti informasi segmen SFAS 131 secara nyata meningkatkan kualitas prakiraan — standar mencapai tujuan informative disclosure-nya.')

    add_body_para(doc,
        'Meskipun demikian, studi Herrmann & Thomas (2000) menambahkan catatan analitis yang perlu diperhatikan. Mereka menemukan bahwa 68 dari 100 perusahaan sampel mengubah definisi segmen mereka setelah mengadopsi SFAS 131. Temuan ini bersifat double-edged: sebagian besar perubahan definisi mencerminkan realignment yang tepat dengan struktur internal — standar yang berhasil memaksa perusahaan untuk melaporkan segmen sebagaimana mereka sungguh-sungguh diorganisasi secara internal. Namun, sebagian perubahan mungkin mencerminkan pemanfaatan aggregation loophole: karena segmen ditentukan berdasarkan cara CODM mengorganisasi bisnis, perusahaan yang mengorganisasi secara luas secara internal berhak melaporkan hanya satu atau dua segmen yang sangat agregat, bahkan jika unit-unit tersebut sangat beragam secara ekonomi. Ini adalah keterbatasan inheren dari pendekatan manajemen: kualitas pengungkapan dibatasi oleh pilihan struktur organisasional, bukan hanya pilihan akuntansi.')

    add_body_para(doc,
        'Namun catatan kritis ini tidak membalikkan verdict. Temuan validitas prediktif beroperasi pada tingkat populasi dan mengonfirmasi bahwa informasi SFAS 131 lebih berguna secara rata-rata — bahkan memperhitungkan perusahaan yang beragregasi secara luas. Gabungan bukti dari Street, Nichols & Gray (2000) dan penelitian validitas prediktif mendukung kesimpulan yang ditegaskan dalam literatur: SFAS No. 131 jelas merupakan peningkatan dibandingkan SFAS No. 14, meskipun bukan peningkatan tanpa batas.')

    # ------------------------------------------------------------------
    # SECTION V — Aplikasi: ICBP dan PSAK 5
    # ------------------------------------------------------------------
    add_section_heading(doc,
        "Aplikasi: ICBP dan PSAK 5 sebagai Cermin SFAS 131", "V.")

    add_body_para(doc,
        'Pendekatan manajemen yang diperkenalkan SFAS 131 tidak hanya diadopsi di Amerika Serikat. Pada tahun 2006, IASB menerbitkan IFRS 8 (Operating Segments) yang mengadopsi pendekatan manajemen dari SFAS 131 secara substansial, dan Indonesia mengikuti melalui PSAK 5 (Segmen Operasi) yang diterbitkan oleh DSAK IAI. Lineage ini — SFAS 131 (1997) ke IFRS 8 (2006) ke PSAK 5 — menjadikan SFAS 131 relevan secara langsung bagi pelaporan keuangan korporat Indonesia: prinsip CODM, ambang 10%/10%/10%, dan persyaratan 75% pendapatan konsolidasi semua berlaku dalam konteks Indonesia (IAI, 2020).')

    add_body_para(doc,
        'PT Indofood CBP Sukses Makmur Tbk (ICBP) menyajikan ilustrasi yang kaya tentang bagaimana pendekatan manajemen bekerja dalam praktik. Dalam laporan tahunan 2024, ICBP mengungkap enam segmen operasi: Mie Instan, Produk Susu, Penyedap Makanan, Makanan Ringan, Minuman, dan Pinehill International (operasi di Timur Tengah, Afrika, dan Eropa). Keenam segmen ini mencerminkan cara CODM ICBP mengorganisasi dan memantau kinerja bisnis secara internal. Nilai pendekatan manajemen terlihat jelas dalam perbandingan kontrafaktual: di bawah rezim SFAS 14, ICBP dapat saja melaporkan sebagai satu segmen “makanan dan minuman” — secara teknis tidak salah, tetapi secara informatif tidak berguna. Di bawah PSAK 5, pengguna laporan keuangan dapat mengamati bahwa segmen Mie Instan (bisnis domestik inti ICBP dengan margin tinggi dan risiko FX rendah) beroperasi dalam kondisi yang sangat berbeda dari segmen Pinehill International (bisnis internasional yang mencakup Nigeria, dengan eksposur signifikan terhadap depresiasi naira dan risiko geopolitik). Perbedaan profil risk-return ini — yang tidak terlihat dalam laporan konsolidasi — adalah informasi yang secara langsung berguna bagi investor dalam membangun model valuasi yang akurat (PT Indofood CBP Sukses Makmur Tbk, 2025).')

    add_body_para(doc,
        'Namun ICBP juga mengillustrasikan keterbatasan yang diidentifikasi Herrmann & Thomas (2000). Segmen Pinehill International mengagregatkan operasi di lebih dari dua belas negara — Nigeria, Arab Saudi, Turki, Maroko, dan lainnya — ke dalam satu unit pelaporan. Agregasi ini diizinkan karena CODM ICBP memang mengorganisasi dan memantau operasi internasional sebagai satu unit. Akibatnya, investor tidak dapat melihat profitabilitas Nigeria versus Arab Saudi secara terpisah, meskipun kedua negara menghadapi kondisi makroekonomi yang sangat berbeda pada 2024: Nigeria mengalami depresiasi naira yang dramatis, sementara Arab Saudi menikmati stabilitas fiskal berbasis harga minyak yang tinggi. Ini adalah aggregation loophole dalam bentuknya yang paling konkret: kualitas pengungkapan dibatasi oleh cara manajemen mengorganisasi bisnis, bukan oleh standar akuntansi. Keunggulan SFAS 131 dan PSAK 5 atas SFAS 14 adalah nyata — tetapi ia bukan tidak terbatas.')

    # ------------------------------------------------------------------
    # SECTION VI — Kesimpulan
    # ------------------------------------------------------------------
    add_section_heading(doc, "Kesimpulan", "VI.")

    add_body_para(doc,
        'Analisis di atas menunjukkan bahwa pertanyaan “mengapa SFAS 131 merupakan perbaikan atas SFAS 14” memiliki jawaban yang berlapis tetapi konsisten.')

    add_body_para(doc,
        'Pada tataran desain standar, SFAS 14 gagal bukan karena kurangnya niat baik, melainkan karena kebebasan definitorial yang diberikannya kepada manajemen tidak dibatasi oleh relevant circumstances yang dapat diidentifikasi. Ini menjadikannya lebih dekat ke flexibility — dalam pengertian Wolk yang negatif — daripada finite uniformity yang sesungguhnya. SFAS 131 mengoreksi kegagalan ini secara struktural dengan menjadikan struktur CODM sebagai jangkar definitorial yang dapat diobservasi dan diaudit. Laporan eksternal menjadi cermin laporan internal — bukan dokumen yang diproduksi khusus untuk memenuhi kewajiban pelaporan.')

    add_body_para(doc,
        'Pada tataran bukti empiris, perbaikan ini bukan sekadar klaim teoritis. Street, Nichols & Gray (2000) mendokumentasikan peningkatan langsung dan terukur dalam granularitas pengungkapan: lebih sedikit perusahaan mengklaim satu segmen, lebih banyak segmen dilaporkan, lebih banyak item diungkap per segmen. Penelitian validitas prediktif mengonfirmasi bahwa informasi yang dihasilkan SFAS 131 lebih berguna secara keputusan. Catatan kritis Herrmann & Thomas (2000) tentang aggregation loophole memperhalus — bukan membatalkan — verdict ini: SFAS 131 adalah perbaikan substantif yang kualitasnya dibatasi oleh pilihan struktur organisasional, bukan pilihan akuntansi.')

    add_body_para(doc,
        'Pada tataran relevansi kebijakan, lineage dari SFAS 131 ke IFRS 8 ke PSAK 5 menunjukkan bahwa komunitas standar internasional telah mengonfirmasi penilaian yang sama. ICBP — dengan enam segmen yang mencerminkan struktur internal CODM-nya — adalah bukti hidup bahwa pendekatan manajemen memberikan informasi yang lebih berguna daripada taksonomi industri yang sewenang-wenang. Dengan demikian, SFAS 131 bukan sekadar perbaikan teknis; ia adalah reposisi fundamental pengungkapan segmental dari kewajiban kepatuhan (protective disclosure) menuju alat dukungan keputusan (informative disclosure) — pergeseran yang Wolk identifikasi sebagai arah kebijakan standar yang menentukan (Wolk et al., 2017, hal. 260).')

    # ------------------------------------------------------------------
    # DAFTAR PUSTAKA
    # ------------------------------------------------------------------
    add_section_heading(doc, "Daftar Pustaka", "")

    add_daftar_pustaka_entry(doc, [
        ("FASB. (1976). ", False),
        ("Statement of Financial Accounting Standards No. 14: Financial reporting "
         "for segments of a business enterprise", True),
        (". Financial Accounting Standards Board.", False),
    ])
    add_daftar_pustaka_entry(doc, [
        ("FASB. (1997). ", False),
        ("Statement of Financial Accounting Standards No. 131: Disclosures about "
         "segments of an enterprise and related information", True),
        (". Financial Accounting Standards Board.", False),
    ])
    add_daftar_pustaka_entry(doc, [
        ("Herrmann, D., & Thomas, W. B. (2000). An analysis of segment disclosures "
         "under SFAS No. 131 and SFAS No. 14. ", False),
        ("Accounting Horizons", True),
        (", 14(3), 287–302.", False),
    ])
    add_daftar_pustaka_entry(doc, [
        ("Ikatan Akuntan Indonesia. (2020). ", False),
        ("PSAK 5: Segmen operasi", True),
        (". Dewan Standar Akuntansi Keuangan IAI.", False),
    ])
    add_daftar_pustaka_entry(doc, [
        ("PT Indofood CBP Sukses Makmur Tbk. (2025). ", False),
        ("Laporan tahunan 2024", True),
        (". ICBP.", False),
    ])
    add_daftar_pustaka_entry(doc, [
        ("Scott, W. R. (2015). ", False),
        ("Financial accounting theory", True),
        (" (7th ed.). Pearson Education Canada.", False),
    ])
    add_daftar_pustaka_entry(doc, [
        ("Street, D. L., Nichols, N. B., & Gray, S. J. (2000). Segment disclosures "
         "under SFAS No. 131: Has business segment reporting improved? ", False),
        ("Accounting Horizons", True),
        (", 14(3), 259–285.", False),
    ])
    add_daftar_pustaka_entry(doc, [
        ("Wolk, H. I., Dodd, J. L., & Rozycki, J. J. (2017). ", False),
        ("Accounting theory: Conceptual issues in a political and economic environment", True),
        (" (9th ed.). SAGE Publications.", False),
    ])

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir,
                            "01079_Dzaki Muhammad Yusfian_Tugas Individual 5.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    print(f"Size:  {os.path.getsize(out_path):,} bytes")


if __name__ == "__main__":
    build()
