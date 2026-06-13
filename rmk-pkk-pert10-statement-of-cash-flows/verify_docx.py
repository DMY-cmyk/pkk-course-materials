"""
Verification script for RMK Pert. 10 docx.
Prints evidence for every checked rule; does not assert (so it never masks a failure).
Run from PROJ root.
"""
import re, sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING

DOCX = Path("output/01079_Dzaki Muhammad Yusfian_RMK Pert. 10.docx")

doc = Document(str(DOCX))

# ── 1. Normal style font ──────────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
print(f"[FONT] Normal.font.name = {font.name!r}  size = {font.size}")
if font.size:
    print(f"       -> {font.size.pt} pt")

# ── 2. Page size (section 0) ─────────────────────────────────────────────────
sec = doc.sections[0]
w_cm = sec.page_width.cm
h_cm = sec.page_height.cm
print(f"[PAGE] width={w_cm:.2f} cm  height={h_cm:.2f} cm  (A4 target 21.0 x 29.7)")

# ── 3. Footer PAGE field ─────────────────────────────────────────────────────
footer_text = []
footer_has_field = False
for para in sec.footer.paragraphs:
    footer_text.append(para.text)
    for run in para.runs:
        for fld in run._r.findall(f".//{qn('w:fldChar')}"):
            footer_has_field = True
# Also look for w:fldSimple or instrText
for elem in sec.footer._element.iter():
    if elem.tag in (qn("w:instrText"), qn("w:fldSimple")):
        instr = elem.text or ""
        if "PAGE" in instr:
            footer_has_field = True
            break
print(f"[FOOTER] text={footer_text!r}  page_field_found={footer_has_field}")

# ── 4. Line spacing on first few body paragraphs ─────────────────────────────
checked = 0
for para in doc.paragraphs:
    if para.style.name in ("Normal", "Body Text") and para.text.strip():
        pf = para.paragraph_format
        rule = pf.line_spacing_rule
        ls   = pf.line_spacing
        print(f"[SPACING] style={para.style.name!r}  rule={rule}  ls={ls}  text={para.text[:60]!r}")
        checked += 1
        if checked >= 5:
            break

# ── 5. Inline images and tables ───────────────────────────────────────────────
inline_imgs = 0
for para in doc.paragraphs:
    for run in para.runs:
        inline_imgs += len(run._r.findall(f".//{qn('a:blip')}"))

# Also count in table cells
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    inline_imgs += len(run._r.findall(f".//{qn('a:blip')}"))

table_count = len(doc.tables)
print(f"[IMAGES] inline image blips = {inline_imgs}")
print(f"[TABLES] count = {table_count}")

# ── 6. Heading presence ───────────────────────────────────────────────────────
all_text = "\n".join(p.text for p in doc.paragraphs)
# Also harvest text from table cells
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            all_text += "\n" + "\n".join(p.text for p in cell.paragraphs)

headings = [
    "RINGKASAN MATERI KULIAH",
    "Bagian A",
    "Bagian B",
    "Bagian C",
    "Bagian D",
    "Bagian E",
    "Bagian F",
]
for h in headings:
    found = h in all_text
    print(f"[HEADING] {h!r} present = {found}")

# ── 7. Identitas block ────────────────────────────────────────────────────────
ident_keys = ["Nama", "NIM", "Mata Kuliah", "Pertemuan"]
for k in ident_keys:
    print(f"[IDENTITAS] {k!r} present = {k in all_text}")

# ── 8. Bagian D word count ────────────────────────────────────────────────────
# Collect text between "Bagian D" and "Bagian E"
m_d = re.search(r"Bagian D(.+?)Bagian E", all_text, re.DOTALL | re.IGNORECASE)
if m_d:
    d_text = m_d.group(1)
    d_words = len(d_text.split())
    print(f"[BAGIAN_D] word count = {d_words}  (target 150–250)  snippet: {d_text[:120]!r}")
else:
    print("[BAGIAN_D] section NOT found between markers")

# ── 9. Bagian E question/list count ──────────────────────────────────────────
m_e = re.search(r"Bagian E(.+?)Bagian F", all_text, re.DOTALL | re.IGNORECASE)
if m_e:
    e_text = m_e.group(1)
    # Count lines that look like numbered items or questions
    lines = [l.strip() for l in e_text.split("\n") if l.strip()]
    q_lines = [l for l in lines if re.match(r"^\d+[.)]\s|^\?|.*\?$", l)]
    print(f"[BAGIAN_E] total lines = {len(lines)}  question/list lines = {len(q_lines)}")
    print(f"           snippet: {e_text[:300]!r}")
else:
    print("[BAGIAN_E] section NOT found between markers")

# ── 10. Cornell table detection ───────────────────────────────────────────────
# Cornell table is a 2-col table with headers Cue / Notes
cornell_found = False
for i, table in enumerate(doc.tables):
    if len(table.columns) == 2:
        header_texts = [c.text.strip() for c in table.rows[0].cells] if table.rows else []
        if any("cue" in t.lower() or "catatan" in t.lower() or "ringkasan" in t.lower()
               for t in header_texts):
            cornell_found = True
            print(f"[CORNELL] table {i} looks like Cornell: headers={header_texts}")
    elif len(table.columns) >= 2:
        # Count rows — large tables are likely Cornell body
        if len(table.rows) >= 10:
            header_texts = [c.text.strip() for c in table.rows[0].cells] if table.rows else []
            print(f"[CORNELL?] table {i} has {len(table.rows)} rows, {len(table.columns)} cols, headers={header_texts[:4]}")

if not cornell_found:
    # Just dump table shapes
    for i, table in enumerate(doc.tables):
        header_texts = [c.text.strip()[:30] for c in table.rows[0].cells] if table.rows else []
        print(f"[TABLE {i}] cols={len(table.columns)} rows={len(table.rows)} headers={header_texts}")

print("\nDone.")
