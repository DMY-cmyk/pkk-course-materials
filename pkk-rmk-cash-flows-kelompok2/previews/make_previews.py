# One-shot preview generator for Phase 1 Visual Companion (documented Python
# fallback; the production exhibit pipeline is Rust in Phase 3.5).
import subprocess
from pathlib import Path

from PIL import Image, ImageChops
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Mm, Pt

ROOT = Path(__file__).resolve().parent.parent
PDF = ROOT / "sources/textbook-chapter/Sage_Chapter_13_Kelompok_2.pdf"
PREV = ROOT / "previews"
FIG = PREV / "_fig"
FIG.mkdir(parents=True, exist_ok=True)

MEMBERS = [
    ("122501039", "Satriyo Nugroho"),
    ("122501048", "Mario Da Costa"),
    ("122501067", "Amelda Putri Zhany Wiguna"),
    ("122501078", "Ahmad Ramadhan"),
    ("122501084", "Nida Nur Cahyati"),
    ("122501094", "Priska Putri Parungky"),
]

# ---------------------------------------------------------------- exhibits
def render_page(page: int) -> Path:
    out = FIG / f"page-{page}"
    png = Path(str(out) + f"-{page:02d}.png")
    if not png.exists():
        subprocess.run(
            ["pdftoppm", "-png", "-r", "240", "-f", str(page), "-l", str(page),
             str(PDF), str(out)],
            check=True,
        )
    return png

def crop_region(page: int, top: float, bottom: float) -> Image.Image:
    img = Image.open(render_page(page)).convert("RGB")
    w, h = img.size
    region = img.crop((0, int(h * top), w, int(h * bottom)))
    # trim whitespace
    bg = Image.new("RGB", region.size, (255, 255, 255))
    bbox = ImageChops.difference(region, bg).getbbox()
    if bbox:
        pad = 8
        bbox = (max(bbox[0] - pad, 0), max(bbox[1] - pad, 0),
                min(bbox[2] + pad, region.width), min(bbox[3] + pad, region.height))
        region = region.crop(bbox)
    return region

def stack(parts: list[Image.Image]) -> Image.Image:
    w = max(p.width for p in parts)
    h = sum(p.height for p in parts)
    out = Image.new("RGB", (w, h), (255, 255, 255))
    y = 0
    for p in parts:
        out.paste(p, ((w - p.width) // 2, y))
        y += p.height
    return out

ex33 = FIG / "exhibit-13-03.png"
if not ex33.exists():
    stack([crop_region(8, 0.755, 0.935), crop_region(9, 0.06, 0.46)]).save(ex33)
ex310 = FIG / "exhibit-13-10.png"
crop_region(20, 0.742, 0.90).save(ex310)  # body only; caption supplied in Word

# ---------------------------------------------------------------- docx base
def base_doc(font: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = sec.bottom_margin = Mm(25.4)
    sec.left_margin = sec.right_margin = Mm(25.4)
    normal = doc.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    for h, size in (("Heading 1", 14), ("Heading 2", 12)):
        st = doc.styles[h]
        st.font.name = font
        st.font.size = Pt(size)
        st.font.bold = True
        from docx.shared import RGBColor
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.line_spacing = 1.5
    return doc

def center(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    return p

def cover_block(doc):
    center(doc, "RINGKASAN MATERI KULIAH (RMK)", bold=True, size=14)
    center(doc, "Statement of Cash Flows", bold=True, size=14)
    center(doc, "Wolk, Dodd & Rozycki — Accounting Theory, 9th ed., Chapter 13")
    center(doc, "Mata Kuliah: Pelaporan Keuangan Korporat (MNK202)")
    center(doc, "Kelompok 2", bold=True)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    hdr[0].paragraphs[0].add_run("NIM").bold = True
    hdr[1].paragraphs[0].add_run("Nama").bold = True
    for nim, nama in MEMBERS:
        c = tbl.add_row().cells
        c[0].text, c[1].text = nim, nama
    doc.add_paragraph()

def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True

def picture(doc, path, width_in):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))

S5_P1 = (
    "SFAS No. 95 mengizinkan dua cara menyajikan arus kas dari aktivitas operasi. "
    "Metode langsung (direct method) melaporkan arus kas secara literal mengikuti "
    "klasifikasi laporan laba rugi: kas yang diterima dari pelanggan, kas yang dibayarkan "
    "kepada pemasok dan karyawan, bunga yang diterima dan dibayar, serta pajak penghasilan "
    "yang dibayar. Sebaliknya, metode tidak langsung (indirect atau reconciliation method) "
    "berangkat dari laba akrual, lalu menyesuaikannya dengan pos-pos nonkas — depresiasi, "
    "keuntungan pelepasan aset, serta perubahan akun modal kerja — hingga diperoleh kas "
    "neto dari aktivitas operasi."
)
S5_P2 = (
    "FASB menyukai metode langsung karena informasi yang dilaporkan lebih kaya. Namun "
    "dalam exposure draft maupun standar finalnya, FASB mengakui argumen biaya: tidak semua "
    "perusahaan menata catatan akuntansinya sehingga data arus kas literal tersedia. Yang "
    "tidak disebut FASB, dalam praktik metode tidak langsung kerap memuat angka penyeimbang "
    "(plug number) agar laporan tetap balance. Mayoritas besar perusahaan Amerika memilih "
    "metode tidak langsung — preferensi yang tampak digerakkan oleh biaya penyusunan. Survei "
    "McEnroe atas 282 responden (analis keuangan, penasihat investasi, akademisi, dan akuntan) "
    "menemukan 56% menyukai metode langsung dan hanya 44% menyukai metode tidak langsung."
)
S5_P3 = (
    "Kedua metode menghasilkan angka cash flow from operations (CFO) yang sama; yang berbeda "
    "adalah informasi dalam perjalanan menuju angka itu. Trade-off-nya dapat dirumuskan tajam: "
    "metode langsung mudah dipahami tetapi sulit disusun, metode tidak langsung mudah disusun "
    "tetapi sulit dipahami. Jika metode langsung digunakan, SFAS No. 95 tetap mewajibkan skedul "
    "rekonsiliasi laba bersih ke kas operasi — sebagaimana diilustrasikan pada Exhibit 13.3 — "
    "sehingga metode rekonsiliasi pada hakikatnya selalu hadir, sendiri atau sebagai suplemen."
)
EX33_CAP = ("Exhibit 13.3 — Indirect or Reconciliation Method of Presenting Net Cash Flows "
            "From Operating Activities (Wolk, Dodd & Rozycki, 2017, Ch. 13)")

FCF_P1 = (
    "Free cash flow (FCF) adalah analog tingkat-perusahaan dari arus kas yang digunakan dalam "
    "keputusan penganggaran modal. Mulford dan Comiskey menegaskan bahwa kata “free” merujuk "
    "pada tiadanya klaim yang lebih senior: kas yang benar-benar bebas digunakan tanpa mengurangi "
    "kemampuan perusahaan menghasilkan kas berikutnya. Mengikuti entity theory, fokusnya adalah "
    "arus kas kepada perusahaan (cash flow to the firm), sehingga FCF didefinisikan sebagai NOPLAT "
    "dikurangi investasi pada operating invested capital. Konsekuensinya, beban bunga — sebuah "
    "beban pendanaan — tidak termasuk dalam FCF, dan kas operasi diperlakukan sebagai bagian dari "
    "modal kerja operasi neto, setara piutang dan persediaan."
)
FCF_P2 = (
    "FCF tidak dapat dibaca langsung dari SCF; ia harus dikonstruksi. Exhibit 13.10 memperlihatkan "
    "jembatannya: mulai dari CFO, tambahkan kembali beban bunga setelah pajak untuk membersihkan "
    "unsur pendanaan yang “mengontaminasi” CFO, sesuaikan perubahan kas operasi sebagai bagian "
    "dari invested capital, lalu perhitungkan arus kas investasi neto. Untuk tahun 2005, ABC Company "
    "menghasilkan FCF $332: CFO $527 disesuaikan bunga setelah pajak $26, perubahan kas operasi, dan "
    "investasi neto $(277). Pembacaan ini menunjukkan mengapa FCF lebih bersih sebagai dasar valuasi: "
    "nilai intrinsik perusahaan adalah FCF masa depan yang didiskontokan pada biaya modal rata-rata "
    "tertimbang (WACC)."
)
EX310_CAP = ("Exhibit 13.10 — Computing Free Cash Flow From the SCF for ABC Company "
             "(Wolk, Dodd & Rozycki, 2017, Ch. 13)")

def add_s5(doc, img_width=None):
    doc.add_heading("5. Metode Langsung vs Metode Tidak Langsung", level=1)
    doc.add_paragraph(S5_P1)
    doc.add_heading("5.1 Preferensi FASB dan praktik yang berlawanan", level=2)
    doc.add_paragraph(S5_P2)
    if img_width:
        picture(doc, ex33, img_width)
        caption(doc, EX33_CAP)
    doc.add_paragraph(S5_P3)

# ------------------------------------------------------- Class 1: fonts
for opt, font in (("option-a-calibri", "Calibri"), ("option-b-aptos", "Aptos")):
    d = base_doc(font)
    cover_block(d)
    add_s5(d)
    out = PREV / "font-template" / f"{opt}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    d.save(out)

# ------------------------------------------- Class 2: exhibit placement
for opt, width in (("option-a-full-width", 6.25), ("option-b-inset", 4.6)):
    d = base_doc("Calibri")
    cover_block(d)
    add_s5(d, img_width=width)
    out = PREV / "exhibit-placement" / f"{opt}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    d.save(out)

# ---------------------------------------------- Class 3: section layout
def add_fcf(doc, with_table: bool):
    doc.add_heading("12. Free Cash Flow", level=1)
    doc.add_paragraph(FCF_P1)
    picture(doc, ex310, 5.8)
    caption(doc, EX310_CAP)
    doc.add_paragraph(FCF_P2)
    if with_table:
        doc.add_heading("Ringkasan empat ukuran kinerja", level=2)
        rows = [
            ("Ukuran", "Apa yang ditangkap", "Keterbatasan utama"),
            ("Laba bersih", "Kinerja akrual ringkas", "Banyak akrual nonkas; rawan manajemen laba"),
            ("CFO", "Kas dari operasi; kualitas laba", "Terkontaminasi bunga; rawan misklasifikasi"),
            ("CFO − CFI", "Kemampuan neto menghasilkan kas", "Bergantung legitimasi investasi"),
            ("FCF", "Kas bebas untuk penyedia modal; dasar valuasi", "Tidak tersedia langsung dari SCF"),
        ]
        tbl = doc.add_table(rows=0, cols=3)
        tbl.style = "Table Grid"
        for i, row in enumerate(rows):
            cells = tbl.add_row().cells
            for c, text in zip(cells, row):
                r = c.paragraphs[0].add_run(text)
                if i == 0:
                    r.bold = True

for opt, with_table in (("option-a-prose", False), ("option-b-prose-plus-table", True)):
    d = base_doc("Calibri")
    cover_block(d)
    add_fcf(d, with_table)
    out = PREV / "section-layout" / f"{opt}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    d.save(out)

print("Previews generated.")
