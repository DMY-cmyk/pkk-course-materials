"""Validate the built RMK docx against the hard gates. Prints JSON to stdout."""
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Mm, Pt

MEMBERS = ["122501039", "122501048", "122501067", "122501078", "122501084", "122501094"]
NAMES = ["Satriyo Nugroho", "Mario Da Costa", "Amelda Putri Zhany Wiguna",
         "Ahmad Ramadhan", "Nida Nur Cahyati", "Priska Putri Parungky"]
CAPTION_MARKERS = [f"Exhibit 13.{n}" for n in range(1, 12)]
EQUATION_MARKERS = ["(13.1)", "(13.2)"]


def all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def page_count(path):
    try:
        import win32com.client  # type: ignore
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        d = word.Documents.Open(str(Path(path).resolve()), ReadOnly=True)
        n = d.ComputeStatistics(2)  # wdStatisticPages
        d.Close(False)
        word.Quit()
        return n
    except Exception as e:  # noqa: BLE001 — report, don't crash
        return {"error": str(e)}


def main(path):
    doc = Document(path)
    sec = doc.sections[0]
    text = all_text(doc)
    normal = doc.styles["Normal"]
    result = {
        "a4": abs(sec.page_width - Mm(210)) < Mm(0.5) and abs(sec.page_height - Mm(297)) < Mm(0.5),
        "font": normal.font.name == "Calibri",
        "size_12": normal.font.size == Pt(12),
        "line_spacing_1_5": normal.paragraph_format.line_spacing == 1.5,
        "images": len(doc.inline_shapes),
        "images_ok": len(doc.inline_shapes) >= 8,
        "captions_missing": [m for m in CAPTION_MARKERS if m not in text],
        "equations_missing": [m for m in EQUATION_MARKERS if m not in text],
        "identity_missing": [m for m in MEMBERS + NAMES if m not in text],
        "page_count": page_count(path),
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1]))
