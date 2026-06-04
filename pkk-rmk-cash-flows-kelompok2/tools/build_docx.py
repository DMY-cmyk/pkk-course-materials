"""python-docx bridge for the RMK pipeline.

Documented substitution for the docx skill (unavailable on this Windows host).
Invoked by crates/rmk-build. Phase 2 ships --smoke only; the full
markdown+manifest assembly lands in Phase 4.
"""
import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Mm, Pt

ROOT = Path(__file__).resolve().parent.parent

FONT = "Calibri"          # Phase 1 decision D3
FONT_SIZE_PT = 12
LINE_SPACING = 1.5
FULL_COLUMN_IN = 6.25     # Phase 1 decision D4


def base_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)  # A4
    sec.top_margin = sec.bottom_margin = Mm(25.4)
    sec.left_margin = sec.right_margin = Mm(25.4)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(FONT_SIZE_PT)
    normal.paragraph_format.line_spacing = LINE_SPACING
    normal.paragraph_format.space_after = Pt(6)
    return doc


def smoke() -> int:
    """Produce a stub docx proving A4 / 1.5 / Calibri 12 / image embedding."""
    test_image = ROOT / "previews" / "_fig" / "exhibit-13-10.png"
    if not test_image.exists():
        print(f"smoke: test image missing: {test_image}", file=sys.stderr)
        return 1
    doc = base_doc()
    doc.add_paragraph("Smoke test — A4, 1.5 spacing, Calibri 12 pt.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(test_image), width=Inches(FULL_COLUMN_IN))
    out = ROOT / "output" / "_smoke.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    # verify round-trip
    check = Document(str(out))
    sec = check.sections[0]
    # Word stores page size in twips, so EMU round-trips within ~0.02 mm.
    assert abs(sec.page_width - Mm(210)) < Mm(0.05), "not A4 width"
    assert abs(sec.page_height - Mm(297)) < Mm(0.05), "not A4 height"
    assert check.styles["Normal"].font.name == FONT, "font mismatch"
    assert check.styles["Normal"].font.size == Pt(FONT_SIZE_PT), "size mismatch"
    assert check.styles["Normal"].paragraph_format.line_spacing == LINE_SPACING
    assert check.inline_shapes, "no embedded image"
    print(f"smoke OK: {out}")
    return 0


def parse_md_runs(s: str):
    """Minimal **bold**/*italic* splitter for table cells.

    Returns list of (text, bold, italic) tuples.
    """
    results = []
    rest = s
    while rest:
        if rest.startswith("**"):
            inner = rest[2:]
            end = inner.find("**")
            if end != -1:
                results.append((inner[:end], True, False))
                rest = inner[end + 2:]
                continue
        if rest.startswith("*"):
            inner = rest[1:]
            end = inner.find("*")
            if end != -1:
                results.append((inner[:end], False, True))
                rest = inner[end + 1:]
                continue
        # find next '*' — treat everything up to it as plain text
        m = re.search(r"\*", rest[1:])
        if m:
            cut = m.start() + 1
        else:
            cut = len(rest)
        results.append((rest[:cut], False, False))
        rest = rest[cut:]
    return results


def _apply_font(run_obj, bold: bool, italic: bool):
    run_obj.bold = bold
    run_obj.italic = italic
    run_obj.font.name = FONT
    run_obj.font.size = Pt(FONT_SIZE_PT)


def _set_heading_style(doc: Document, level: int):
    """Ensure Heading 1/2 uses Calibri, correct size, black, 1.5 spacing."""
    style_name = f"Heading {level}"
    try:
        style = doc.styles[style_name]
    except KeyError:
        return
    font_size = 14 if level == 1 else 12
    style.font.name = FONT
    style.font.size = Pt(font_size)
    style.font.bold = True
    style.font.color.rgb = None  # inherit / black
    style.paragraph_format.line_spacing = LINE_SPACING


def render_spec(spec_path: str, output: str) -> int:
    with open(spec_path, encoding="utf-8") as fh:
        blocks = json.load(fh)

    doc = base_doc()
    _set_heading_style(doc, 1)
    _set_heading_style(doc, 2)

    for block in blocks:
        btype = block["type"]

        if btype == "heading":
            level = block["level"]
            text = block["text"]
            p = doc.add_heading(text, level=level)
            p.paragraph_format.line_spacing = LINE_SPACING
            # Ensure the run font is applied
            for run in p.runs:
                run.font.name = FONT
                run.font.size = Pt(14 if level == 1 else 12)
                run.font.bold = True

        elif btype == "para":
            p = doc.add_paragraph()
            for run_data in block["runs"]:
                r = p.add_run(run_data["text"])
                _apply_font(r, run_data["bold"], run_data["italic"])

        elif btype == "center":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run_data in block["runs"]:
                r = p.add_run(run_data["text"])
                _apply_font(r, run_data["bold"], run_data["italic"])

        elif btype == "image":
            img_path = ROOT / block["path"]
            width_in = block["width_in"]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run()
            run.add_picture(str(img_path), width=Inches(width_in))

        elif btype == "caption":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(block["text"])
            r.italic = True
            r.font.name = FONT
            r.font.size = Pt(FONT_SIZE_PT)

        elif btype == "table":
            rows = block["rows"]
            if not rows:
                continue
            ncols = max(len(row) for row in rows)
            # Pad short rows
            padded = [row + [""] * (ncols - len(row)) for row in rows]
            tbl = doc.add_table(rows=len(padded), cols=ncols)
            tbl.style = "Table Grid"
            for r_idx, row_data in enumerate(padded):
                for c_idx, cell_text in enumerate(row_data):
                    cell = tbl.cell(r_idx, c_idx)
                    cell_para = cell.paragraphs[0]
                    cell_para.paragraph_format.line_spacing = LINE_SPACING
                    for part_text, bold, italic in parse_md_runs(cell_text):
                        r = cell_para.add_run(part_text)
                        _apply_font(r, bold, italic)

        elif btype == "equation":
            # Label paragraph (left-aligned, no space after)
            label_p = doc.add_paragraph()
            label_p.paragraph_format.space_after = Pt(0)
            lr = label_p.add_run(block["label"])
            _apply_font(lr, False, False)
            # Equation text (centered, bold)
            eq_p = doc.add_paragraph()
            eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            er = eq_p.add_run(block["text"])
            _apply_font(er, True, False)

        else:
            raise ValueError(f"Unknown block type: {btype!r}")

    out_path = Path(output)
    if not out_path.is_absolute():
        out_path = ROOT / out_path
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"built {out_path}")
    return 0


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--smoke", action="store_true", help="baseline smoke test")
    ap.add_argument("--spec", metavar="PATH", help="path to build-spec.json")
    ap.add_argument("--output", metavar="PATH", help="output .docx path")
    args = ap.parse_args()
    if args.smoke:
        return smoke()
    if args.spec:
        if not args.output:
            print("--output required with --spec", file=sys.stderr)
            return 1
        return render_spec(args.spec, args.output)
    print("full build not yet implemented (Phase 4)", file=sys.stderr)
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
