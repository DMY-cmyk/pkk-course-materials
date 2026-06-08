"""
build_docx.py — assembles output/01079_Dzaki Muhammad Yusfian_RMK Pert. 9.docx
from content/*.md + assets/. python-docx is the documented Python exception in
this otherwise-Rust pipeline (see README: proven K2 typography conventions;
docx-rs would need exact-18pt spacing, hanging indents, captioned tables and
footer fields re-proven from scratch).
Styling layer adapted from "Kelompok 2 Pasca UTS/output/build_docx.py".
"""
import glob
import os
import re
import tomllib

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT = "Times New Roman"
STUDENT = "Dzaki Muhammad Yusfian"
NIM = "1225 01079"
OUT_NAME = "01079_Dzaki Muhammad Yusfian_RMK Pert. 9.docx"

# --- markdown parsing (K2 lineage, + ### and @table) -----------------------

def parse_inline_runs(text):
    runs = []
    pattern = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], False, False))
        token = m.group(0)
        if token.startswith('**'):
            runs.append((token[2:-2], True, False))
        else:
            runs.append((token[1:-1], False, True))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], False, False))
    return [r for r in runs if r[0]]


IMAGE_RE = re.compile(r'^!\[(.+?)\]\((.+?)\)$')
TABLE_RE = re.compile(r'^@table\((.+?)\)$')


def split_caption(caption):
    if " | " in caption:
        title, source = caption.split(" | ", 1)
        return title.strip(), source.strip()
    return caption.strip(), None


def parse_blocks(md_text):
    blocks = []
    in_refs = False
    buf = []

    def flush():
        nonlocal buf
        if buf:
            blocks.append(("ref" if in_refs else "para", " ".join(buf)))
            buf = []

    for raw in md_text.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            flush()
        elif line.startswith("### "):
            flush()
            blocks.append(("subheading", line[4:].strip()))
        elif line.startswith("## "):
            flush()
            text = line[3:].strip()
            blocks.append(("heading", text))
            in_refs = "referensi" in text.lower() or "daftar pustaka" in text.lower()
        elif line.startswith("# "):
            flush()
        elif TABLE_RE.match(stripped):
            flush()
            blocks.append(("table", TABLE_RE.match(stripped).group(1)))
        elif line.startswith("!["):
            m = IMAGE_RE.match(stripped)
            if m:
                flush()
                blocks.append(("image", (m.group(1), m.group(2))))
            else:
                buf.append(stripped)
        elif line.startswith("- "):
            flush()
            blocks.append(("bullet", line[2:].strip()))
        else:
            buf.append(stripped)
    flush()
    return blocks

# --- docx styling (K2 lineage) ----------------------------------------------

def _style_run(run, font_size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rFonts.set(qn('w:cs'), FONT)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _add_para(doc, runs, font_size=12, bold=False,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
              space_before_pt=0, space_after_pt=6,
              left_indent_cm=None, hanging=False, italic_all=False):
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    if left_indent_cm is not None:
        pf.left_indent = Cm(left_indent_cm)
    if hanging:
        pf.left_indent = Cm(0.75)
        pf.first_line_indent = Cm(-0.75)
    for text, b, i in runs:
        run = para.add_run(text)
        _style_run(run, font_size=font_size, bold=bold or b, italic=i or italic_all)
    return para


def add_blank(doc):
    return _add_para(doc, [("", False, False)], space_after_pt=0)


def add_rule(doc):
    """Thin horizontal rule under an empty paragraph (front-matter divider)."""
    para = _add_para(doc, [("", False, False)], space_after_pt=10)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def add_page_number_footer(section):
    """Bottom-right PAGE field in the footer."""
    para = section.footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run()
    _style_run(run, font_size=11)
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)


CAPTION_PREFIX_RE = re.compile(r'^((?:Gambar|Tabel) \d+\.)\s*(.*)$')


def _add_caption(doc, caption, space_after_pt=12):
    title, source = split_caption(caption)
    m = CAPTION_PREFIX_RE.match(title)
    runs = ([(m.group(1) + " ", True, False), (m.group(2), False, False)]
            if m else [(title, False, False)])
    _add_para(doc, runs, font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER,
              space_before_pt=0, space_after_pt=0 if source else space_after_pt)
    if source:
        _add_para(doc, [(source, False, True)], font_size=11,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before_pt=0, space_after_pt=space_after_pt)


def add_image_with_caption(doc, img_path, caption):
    doc.add_picture(img_path, width=Cm(14.5))
    pic = doc.paragraphs[-1]
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.space_before = Pt(6)
    pic.paragraph_format.space_after = Pt(6)
    _add_caption(doc, caption)


def add_table_from_toml(doc, toml_path):
    with open(toml_path, 'rb') as f:
        spec = tomllib.load(f)
    _add_caption(doc, spec["caption"], space_after_pt=4)  # caption ABOVE table
    n_cols = len(spec["header"])
    table = doc.add_table(rows=1 + len(spec["rows"]), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(spec["header"]):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _style_run(run, font_size=11, bold=True)
    for i, row in enumerate(spec["rows"], start=1):
        for j, val in enumerate(row["cells"]):
            cell = table.rows[i].cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(14)
            for text, b, it in parse_inline_runs(val):
                _style_run(para.add_run(text), font_size=11, bold=b, italic=it)
    for j, w in enumerate(spec.get("widths_cm", [])):
        for row in table.rows:
            row.cells[j].width = Cm(w)
    add_blank(doc)

# --- build -------------------------------------------------------------------

def build():
    here = os.path.dirname(os.path.abspath(__file__))
    root = os.path.normpath(os.path.join(here, "..", ".."))
    content_dir = os.path.join(root, "content")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Cm(3))
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(12)
    add_page_number_footer(section)

    # Concise front matter (Layout B) — from content/00_front_matter.md:
    # each non-blank line = centered paragraph; first line bold 14pt.
    fm_path = os.path.join(content_dir, "00_front_matter.md")
    with open(fm_path, encoding="utf-8") as f:
        fm_lines = [ln.strip() for ln in f if ln.strip() and not ln.startswith("#")]
    for i, ln in enumerate(fm_lines):
        _add_para(doc, parse_inline_runs(ln),
                  font_size=14 if i == 0 else 12, bold=(i == 0),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=2)
    add_rule(doc)

    # Body: every content/*.md except 00, sorted by filename
    files = sorted(glob.glob(os.path.join(content_dir, "*.md")))
    for path in files:
        if os.path.basename(path).startswith("00_"):
            continue
        with open(path, encoding="utf-8") as f:
            md = f.read()
        for kind, payload in parse_blocks(md):
            if kind == "image":
                caption, rel = payload
                add_image_with_caption(
                    doc, os.path.normpath(os.path.join(content_dir, rel)), caption)
            elif kind == "table":
                add_table_from_toml(
                    doc, os.path.normpath(os.path.join(content_dir, payload)))
            elif kind == "heading":
                _add_para(doc, parse_inline_runs(payload), font_size=13, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          space_before_pt=12, space_after_pt=6)
            elif kind == "subheading":
                _add_para(doc, parse_inline_runs(payload), font_size=12, bold=True,
                          italic_all=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          space_before_pt=8, space_after_pt=4)
            elif kind == "bullet":
                _add_para(doc, [("• ", False, False)] + parse_inline_runs(payload),
                          left_indent_cm=0.75)
            elif kind == "ref":
                _add_para(doc, parse_inline_runs(payload), hanging=True)
            else:
                _add_para(doc, parse_inline_runs(payload))

    out_path = os.path.join(root, "output", OUT_NAME)
    doc.save(out_path)
    print(f"Saved: {out_path}")
    print(f"Size:  {os.path.getsize(out_path):,} bytes")


if __name__ == "__main__":
    build()
