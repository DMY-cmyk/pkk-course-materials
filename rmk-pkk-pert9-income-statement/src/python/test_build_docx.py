import os, subprocess, sys
import pytest
from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.normpath(os.path.join(HERE, "..", ".."))
OUT = os.path.join(ROOT, "output", "01079_Dzaki Muhammad Yusfian_RMK Pert. 9.docx")

sys.path.insert(0, HERE)
from build_docx import parse_inline_runs, parse_blocks, split_caption  # noqa: E402


def test_inline_runs_bold_italic():
    assert parse_inline_runs("a **b** *c*") == [
        ("a ", False, False), ("b", True, False), (" ", False, False), ("c", False, True)]


def test_parse_blocks_subheading_and_table():
    md = "## I. Judul\n\n### Sub Bagian\n\npara satu\n\n@table(../assets/tables/tabel1_definisi.toml)\n\n- butir\n"
    kinds = [k for k, _ in parse_blocks(md)]
    assert kinds == ["heading", "subheading", "para", "table", "bullet"]


def test_split_caption():
    t, s = split_caption("Tabel 1. Judul | Sumber: diolah dari Wolk et al. (2017)")
    assert t.startswith("Tabel 1.") and s.startswith("Sumber:")


@pytest.fixture(scope="module")
def built():
    subprocess.run([sys.executable, os.path.join(HERE, "build_docx.py")],
                   check=True, cwd=ROOT)
    return Document(OUT)


def test_output_exists_with_exact_name(built):
    assert os.path.exists(OUT)


def test_page_setup(built):
    # A4 + 3 cm margins. python-docx stores page size in twips, so an exact
    # EMU match to Cm(21.0) is impossible (round-trips to 11906 twips); compare
    # at 0.1 cm precision instead.
    s = built.sections[0]
    assert round(s.page_width.cm, 1) == 21.0
    assert round(s.page_height.cm, 1) == 29.7
    assert round(s.left_margin.cm, 1) == 3.0
    assert round(s.right_margin.cm, 1) == 3.0


def test_front_matter_first_line(built):
    first = built.paragraphs[0].text
    assert "RINGKASAN MATERI KULIAH" in first and "PERTEMUAN 9" in first.upper()


def test_six_images_and_four_tables(built):
    assert len(built.inline_shapes) == 6
    assert len(built.tables) == 4


def test_identity_present(built):
    head = "\n".join(p.text for p in built.paragraphs[:6])
    assert "Dzaki Muhammad Yusfian" in head and "01079" in head
